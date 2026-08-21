import shutil
import logging
import cv2
import datetime
import sys
import pytesseract
import re
import json
import os
import functools
import docx.text
import numpy as np
import tkinter as tk
import win32com.client as win32
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from openpyxl import load_workbook
from PIL import Image, ImageEnhance
from tkinter import messagebox
from copy import deepcopy
from docx.shared import RGBColor
from docx.enum.text import WD_COLOR_INDEX
from typing import Any, List, Optional, Tuple

# Указываем путь к tesseract
pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'

# Настройка логирования (вывод только в консоль)
logging.basicConfig(level=logging.INFO, format="%(asctime)s - %(levelname)s - %(message)s")
logger = logging.getLogger(__name__)

# Определяем порядок сортировки локусов (в верхнем регистре)
LOCUS_ORDER = [
    "D3S1358", "VWA", "D16S539", "CSF1PO", "TPOX", "YINDEL", "AMEL", 
    "D8S1179", "D21S11", "D18S51", "DYS391", "D2S441", "D19S433", 
    "TH01", "FGA", "D22S1045", "D5S818", "D13S317", "D7S820", "SE33", 
    "D10S1248", "D1S1656", "D12S391", "D2S1338", "SRY", "PENTA E",
    "D6S1043", "RS759551978", "RS771783753", "RS199815934", "PENTA D", "D4S2366"
]

# Декоратор для автоматического перестроения кэша win32com при ошибках CLSIDToClassMap. Используется для функций, работающих с win32com.
def rebuild_win32com_cache_on_error(func):
    @functools.wraps(func)
    def wrapper(*args, **kwargs):
        try:
            return func(*args, **kwargs)
        except AttributeError as e:
            if "CLSIDToClassMap" in str(e):
                logger.warning("Обнаружена ошибка кэша win32com. Автоматическое перестроение кэша...")
                try:
                    # Перестраиваем кэш win32com
                    import win32com.client
                    win32com.client.gencache.is_readonly = False
                    win32com.client.gencache.Rebuild()
                    logger.info("Кэш win32com успешно перестроен. Повторная попытка выполнения функции...")
                    
                    # Повторяем вызов функции после перестроения кэша
                    return func(*args, **kwargs)
                except Exception as rebuild_error:
                    logger.error(f"Ошибка при перестроении кэша win32com: {rebuild_error}")
                    
                    # Если не удалось исправить ошибку, предлагаем пользователю выбор
                    logger.error("Не удалось автоматически исправить ошибку кэша win32com.")
                    response = input("Применить эвристический метод? (д/н): ").strip().lower()
                    
                    if response == 'д':
                        logger.info("Пользователь выбрал применение эвристического метода. Продолжаем выполнение...")
                        # Возвращаем специальное значение, указывающее на использование эвристики
                        return "USE_HEURISTIC"
                    else:
                        logger.info("Пользователь отказался от эвристического метода. Завершение программы...")
                        # Завершаем программу
                        sys.exit(1)
            # Если это другая AttributeError, пробрасываем её дальше
            raise
    return wrapper

#def test_win32com_error_decorator():
#    """
#    Тестовая функция для проверки работы декоратора rebuild_win32com_cache_on_error.
#    Имитирует реальное поведение win32com с возможностью "исправления" после перестроения кэша.
#    """
#    logger.info("=== НАЧАЛО ТЕСТИРОВАНИЯ ДЕКОРАТОРА win32com ===")
#    
#    # Счетчик вызовов для имитации "исправления" после перестроения кэша
#    call_count = 0
#    
#    @rebuild_win32com_cache_on_error
#    def simulated_win32com_operation():
#        nonlocal call_count
#        call_count += 1
#        
#        if call_count == 1:
#            # Первый вызов - имитируем ошибку кэша
#            logger.info("Первый вызов: искусственное создание ошибки CLSIDToClassMap...")
#            raise AttributeError("module 'win32com.gen_py.00020905-0000-0000-C000-000000000046x0x8x7' has no attribute 'CLSIDToClassMap'")
#        else:
#            # Последующие вызовы - имитируем успешное выполнение после перестроения кэша
#            logger.info(f"Вызов #{call_count}: операция выполнена успешно после перестроения кэша")
#            return "OPERATION_SUCCESS"
#    
#    try:
#        logger.info("Запуск тестовой операции с декоратором...")
#        result = simulated_win32com_operation()
#        logger.info(f"Результат выполнения: {result}")
#        
#        if result == "OPERATION_SUCCESS":
#            logger.info("ТЕСТ ПРОЙДЕН: Декоратор успешно перестроил кэш и операция выполнена")
#        elif result == "USE_HEURISTIC":
#            logger.info("ТЕСТ ЧАСТИЧНО ПРОЙДЕН: Декоратор применил эвристический метод (это fallback)")
#        else:
#            logger.warning(f"ТЕСТ С НЕОЖИДАННЫМ РЕЗУЛЬТАТОМ: {result}")
#            
#    except Exception as e:
#        logger.error(f"ТЕСТ НЕ ПРОЙДЕН: Необработанное исключение: {e}")
#        return False
#    
#    logger.info("=== КОНЕЦ ТЕСТИРОВАНИЯ ДЕКОРАТОРА win32com ===")
#    return True

def setup_wrap_text_debug_logger():
    """
    Создает специальный логгер для диагностики проблем с обработкой WRAP_TEXT.
    Не создает файл, а использует основной логгер.
    
    :return: Объект логгера для диагностики
    """
    # Создаем логгер с уникальным именем
    timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S_%f")
    debug_logger_name = f'wrap_text_debug_{timestamp}'
    
    # Создаем логгер
    debug_logger = logging.getLogger(debug_logger_name)
    
    # Устанавливаем уровень - такой же как у основного логгера
    debug_logger.setLevel(logging.DEBUG)
    
    # Убираем распространение на родительские логгеры, чтобы не дублировать сообщения
    debug_logger.propagate = False
    
    # Добавляем только NullHandler, чтобы не выводить сообщения
    # Но мы будем использовать основной логгер для важных сообщений
    debug_logger.addHandler(logging.NullHandler())
    
    return debug_logger, None  # Возвращаем None вместо пути к файлу

def diagnose_file_access_issues(file_path):
    """
    Диагностирует возможные проблемы с доступом к файлу.
    Выводит диагностическую информацию в лог.
    
    :param file_path: Путь к файлу для диагностики
    """
    logger.debug("[FILE_DIAG] === ДЕТАЛЬНАЯ ДИАГНОСТИКА ДОСТУПА К ФАЙЛУ ===")
    
    try:
        # Анализ пути
        logger.debug(f"[FILE_DIAG] Исходный путь: '{file_path}'")
        logger.debug(f"[FILE_DIAG] Нормализованный путь: '{os.path.normpath(file_path)}'")
        logger.debug(f"[FILE_DIAG] Абсолютный путь: '{os.path.abspath(file_path)}'")
        
        # Проверка компонентов пути
        drive, path = os.path.splitdrive(file_path)
        logger.debug(f"[FILE_DIAG] Диск: '{drive}'")
        logger.debug(f"[FILE_DIAG] Путь без диска: '{path}'")
        
        # Проверка на недопустимые символы
        invalid_chars = '<>:"|?*'
        has_invalid_chars = any(c in file_path for c in invalid_chars)
        logger.debug(f"[FILE_DIAG] Путь содержит недопустимые символы: {has_invalid_chars}")
        if has_invalid_chars:
            for char in invalid_chars:
                if char in file_path:
                    logger.warning(f"[FILE_DIAG] Найден потенциально проблемный символ: '{char}'")
        
        # Проверка длины пути (Windows имеет ограничение в 260 символов)
        if len(file_path) > 260:
            logger.warning(f"[FILE_DIAG] Путь превышает лимит Windows (260 символов): {len(file_path)}")
        
        # Пробуем различные операции с файлом
        if os.path.exists(file_path):
            logger.debug("[FILE_DIAG] ✓ Файл существует")
            
            try:
                stat_info = os.stat(file_path)
                logger.debug(f"[FILE_DIAG] Размер файла: {stat_info.st_size} байт")
                logger.debug(f"[FILE_DIAG] Последняя модификация: {datetime.datetime.fromtimestamp(stat_info.st_mtime)}")
            except Exception as stat_error:
                logger.warning(f"[FILE_DIAG] Ошибка получения информации о файле: {stat_error}")
        else:
            logger.error("[FILE_DIAG] ✗ Файл не существует")
            
            # Проверяем существование родительской папки
            parent_dir = os.path.dirname(file_path)
            if os.path.exists(parent_dir):
                logger.debug(f"[FILE_DIAG] ✓ Родительская папка существует: {parent_dir}")
            else:
                logger.error(f"[FILE_DIAG] ✗ Родительская папка не существует: {parent_dir}")
        
    except Exception as e:
        logger.error(f"[FILE_DIAG] Ошибка при диагностике файла: {e}")

# Показываем сообщение об ошибке в графическом окне и логируем её
def show_error_message(message):
    logger.error(message)
    root = tk.Tk()
    root.withdraw()  # Скрываем главное окно Tkinter
    messagebox.showerror("ОШИБКА", message)
    root.destroy()

def show_critical_message_and_exit(message):
    """
    Показывает критическое сообщение и завершает программу после ожидания ввода.
    """
    logger.error(message)
    root = tk.Tk()
    root.withdraw()
    
    # Создаем кастомное диалоговое окно
    dialog = tk.Toplevel(root)
    dialog.title("ВНИМАНИЕ!")
    dialog.geometry("500x150")
    dialog.transient(root)
    dialog.grab_set()
    
    # Центрируем окно
    dialog.update_idletasks()
    x = (dialog.winfo_screenwidth() - dialog.winfo_width()) // 2
    y = (dialog.winfo_screenheight() - dialog.winfo_height()) // 2
    dialog.geometry(f"+{x}+{y}")
    
    # Добавляем сообщение
    label = tk.Label(dialog, text=message, wraplength=480, justify="center", font=("Arial", 10))
    label.pack(pady=20)
    
    # Добавляем поле для ввода
    entry_label = tk.Label(dialog, text="Введите 'выход' для завершения программы:")
    entry_label.pack(pady=5)
    
    entry = tk.Entry(dialog, width=30)
    entry.pack(pady=5)
    entry.focus_set()
    
    def check_exit():
        if entry.get().strip().lower() == 'выход':
            dialog.destroy()
            root.destroy()
            sys.exit(1)
        else:
            entry.delete(0, tk.END)
    
    # Кнопка подтверждения
    button = tk.Button(dialog, text="Подтвердить", command=check_exit)
    button.pack(pady=10)
    
    # Обработка нажатия Enter
    dialog.bind('<Return>', lambda event: check_exit())
    
    root.mainloop()

# Загружаем словарь замен из JSON-файла в указанной папке
def load_switch_dictionary(templates_folder):
    """Загружает словарь замен из JSON-файла в указанной папке."""
    dictionary_path = Path(templates_folder) / "switch_dictionary.json"
    try:
        if not dictionary_path.exists():
            logger.warning(f"Файл словаря не найден: {dictionary_path}")
            return None
        
        with open(dictionary_path, 'r', encoding='utf-8') as f:
            data = json.load(f)
            if not data:
                logger.warning("Словарь замен пуст")
            else:
                logger.info(f"Загружен словарь замен. Ключи: {list(data.keys())}")
            return data
            
    except json.JSONDecodeError as e:
        logger.error(f"Ошибка при чтении JSON-файла: {str(e)}")
        return None
    except Exception as e:
        logger.error(f"Неожиданная ошибка при загрузке словаря: {str(e)}")
        return None

# Функции для работы с изображениями
def resize_image(image, max_size=1200):
    """Уменьшает размер изображения до максимального размера с сохранением качества."""
    logger.info(f"Resizing image to max size {max_size}")
    height, width = image.shape[:2]
    if max(height, width) > max_size:
        scale = max_size / max(height, width)
        new_width = int(width * scale)
        new_height = int(height * scale)
        return cv2.resize(image, (new_width, new_height), interpolation=cv2.INTER_LINEAR)
    return image

def increase_contrast(image):
    """Увеличивает контрастность изображения без потери деталей."""
    logger.info("Increasing image contrast")
    lab = cv2.cvtColor(image, cv2.COLOR_BGR2LAB)
    l_channel, a, b = cv2.split(lab)
    clahe = cv2.createCLAHE(clipLimit=3.0, tileGridSize=(8, 8))
    cl = clahe.apply(l_channel)
    enhanced_lab = cv2.merge((cl, a, b))
    return cv2.cvtColor(enhanced_lab, cv2.COLOR_LAB2BGR)

def auto_crop_image(image_path, debug=False):
    """
    Автоматическая обработка и обрезка изображения.
    """
    logger.info(f"Auto cropping image: {image_path}")
    try:
        # Загружаем изображение через Pillow (поддерживает кириллицу в пути)
        image_pil = Image.open(image_path)
        # Конвертируем в numpy array для работы с OpenCV
        image = np.array(image_pil)
        image = cv2.cvtColor(image, cv2.COLOR_RGB2BGR)  # PIL -> OpenCV (RGB -> BGR)
    except Exception as e:
        logger.error(f"Ошибка: не удалось загрузить изображение. Проверьте путь к файлу: {e}")
        return None

    # Увеличение контрастности
    image = increase_contrast(image)
    # Уменьшение размера изображения
    image = resize_image(image)
    original_height, original_width = image.shape[:2]
    logger.info(f"Размер обработанного изображения: {original_width}x{original_height}")

    # Преобразование в градации серого
    gray = cv2.cvtColor(image, cv2.COLOR_BGR2GRAY)
    # Размытие для удаления шума
    gray = cv2.GaussianBlur(gray, (5, 5), 0)

    # Обнаружение границ с помощью Canny
    logger.info("Применение детектора границ Canny...")
    edges = cv2.Canny(gray, 30, 100)

    # Поиск контуров
    logger.info("Поиск контуров...")
    contours, _ = cv2.findContours(edges, cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_SIMPLE)
    if len(contours) == 0:
        logger.error("Ошибка: контуры не найдены.")
        return None

    # Фильтрация маленьких контуров
    filtered_contours = [cnt for cnt in contours if cv2.contourArea(cnt) > 100]
    if len(filtered_contours) == 0:
        logger.error("Ошибка: после фильтрации контуры отсутствуют.")
        return None

    # Выбор ограничивающих прямоугольников
    bounding_rects = [cv2.boundingRect(cnt) for cnt in filtered_contours]
    rects = [(x, y, w, h) for x, y, w, h in bounding_rects if 0.3 < w / h < 3]
    if len(rects) == 0:
        logger.error("Ошибка: не найдено подходящих прямоугольников.")
        return None

    # Отладочная информация
    if debug:
        logger.debug(f"Найдено {len(contours)} контуров.")
        for i, cnt in enumerate(filtered_contours):
            area = cv2.contourArea(cnt)
            x, y, w, h = cv2.boundingRect(cnt)
            logger.debug(f"Контур {i}: площадь={area}, размер=({w}x{h}), соотношение сторон={w/h:.2f}")

    # Объединение контуров
    all_contours = np.concatenate(filtered_contours)
    x, y, w, h = cv2.boundingRect(all_contours)

    # Обрезка изображения
    cropped_image = image[y:y + h, x:x + w]
    logger.info(f"Обрезанное изображение: {w}x{h}")

    if debug:
        cv2.imshow("AUTO CROP", cropped_image)
        cv2.waitKey(0)
        cv2.destroyAllWindows()

    return cropped_image

def interactive_crop(image):
    """Позволяет пользователю выбрать область обрезки."""
    print("Выберите область обрезки (нажмите Enter, когда закончите).")
    roi = cv2.selectROI("INTERACTIVE CROP", image, fromCenter=False)
    cv2.destroyWindow("INTERACTIVE CROP")
    if sum(roi) == 0:  # Проверка, была ли выбрана область
        print("Область не выбрана. Используется автоматическая обрезка.")
        return None
    x, y, w, h = roi
    cropped_image = image[int(y):int(y+h), int(x):int(x+w)]
    return cropped_image

def adjust_brightness(image, brightness_factor):
    """Увеличивает яркость изображения."""
    enhancer = ImageEnhance.Brightness(image)
    return enhancer.enhance(brightness_factor)

def adjust_contrast(image, contrast_factor):
    """Увеличивает контрастность изображения."""
    enhancer = ImageEnhance.Contrast(image)
    return enhancer.enhance(contrast_factor)

def load_image_with_pillow(image_path):
    """
    Загружает изображение с помощью Pillow и конвертирует его в numpy array.
    """
    try:
        image = Image.open(image_path)
        return np.array(image)  # Конвертируем в numpy array
    except Exception as e:
        raise ValueError(f"Не удалось загрузить изображение: {image_path}. Ошибка: {e}")

def auto_adjust_brightness_contrast(image_path, brightness_factor=1.5, contrast_factor=1.5):
    """
    Автоматически настраивает яркость и контрастность изображения.
    """
    # Загружаем изображение с помощью Pillow
    image_np = load_image_with_pillow(image_path)
    # Преобразуем из RGB (Pillow) в BGR (OpenCV)
    image_bgr = cv2.cvtColor(image_np, cv2.COLOR_RGB2BGR)
    # Разделяем изображение на каналы (B, G, R)
    channels = cv2.split(image_bgr)
    # Нормализуем каждый канал отдельно
    normalized_channels = []
    for channel in channels:
        normalized = cv2.normalize(channel, None, alpha=0, beta=255, norm_type=cv2.NORM_MINMAX)
        normalized_channels.append(normalized)
    # Объединяем каналы обратно в цветное изображение
    result_bgr = cv2.merge(normalized_channels)
    # Преобразуем обратно в RGB
    result_rgb = cv2.cvtColor(result_bgr, cv2.COLOR_BGR2RGB)
    # Конвертируем в формат PIL для совместимости с docx
    result_image = Image.fromarray(result_rgb)
    # Увеличиваем яркость и контрастность
    result_image = adjust_brightness(result_image, brightness_factor)
    result_image = adjust_contrast(result_image, contrast_factor)
    return result_image

# Функция для сортировки данных по локусам
def sort_data_by_locus(data):
    """
    Сортирует данные по локусам в порядке, определённом в LOCUS_ORDER.
    Локусы, не указанные в LOCUS_ORDER, сортируются по алфавиту.
    При этом сохраняется оригинальный регистр названий локусов.
    """
    def get_locus_index(locus):
        # Приводим локус к верхнему регистру для сравнения
        locus_upper = locus.strip().upper()
        try:
            return LOCUS_ORDER.index(locus_upper)
        except ValueError:
            # Если локус не найден в LOCUS_ORDER, возвращаем большое число,
            # чтобы он оказался в конце списка
            return len(LOCUS_ORDER) + ord(locus_upper[0])  # Сортируем по алфавиту
    
    # Сортируем данные по локусам (без учёта регистра), но сохраняем оригинальный регистр
    return sorted(data, key=lambda x: get_locus_index(x[3]))

# Функция для проверки обязательных полей
def check_card_data(card_data):
    required_fields = ["НОМ", "1", "2", "3", "4", "5", "6", "7"]  # Список обязательных полей
    for field in required_fields:
        if field not in card_data or card_data[field] is None:
            raise ValueError(f"Ошибка: в Карте отсутствует значение для кода '{field}'.")

# Читает данные из Excel-файла и возвращает их в виде словаря.
def read_card(file_path):
    logger.info(f"Чтение данных из Excel-файла: {file_path}")
    data = {}
    try:
        wb = load_workbook(file_path)
        sheet = wb.active
        for row in sheet.iter_rows(min_row=2, values_only=True):
            code = str(row[0]).strip()  # Преобразуем код в строку и убираем лишние пробелы
            value = row[2]
            if code == "ДБИ" and value is not None and value.strip():
                value = f"с {value.strip()} г."
            data[code] = value
        logger.info("Данные успешно загружены из Excel-файла.")
    except Exception as e:
        logger.error(f"Ошибка при чтении Excel-файла: {e}")
        raise
    return data

# Копирует шаблон Word на основе данных из карты.
def copy_template(card_data):
    logger.info("Копирование шаблона Word...")
    template_folder = Path(card_data["3"])  # Путь к папке с шаблонами
    template_name = card_data["1"]  # Название шаблона из строки с кодом 1
    
    # Проверяем, есть ли файл с расширением .doc или .docx
    for ext in [".doc", ".docx"]:
        template_path = template_folder / f"{template_name}{ext}"
        if template_path.exists():
            output_path = Path(card_data["2"]) / f"{card_data['НОМ']}-26{ext}"
            shutil.copy(template_path, output_path)
            logger.info(f"Шаблон скопирован в: {output_path}")
            return output_path
    
    # Если файл не найден
    logger.error(f"Файл шаблона '{template_name}' не найден в папке {template_folder}.")
    raise FileNotFoundError(f"Файл шаблона '{template_name}' не найден в папке {template_folder}.")

# Секция для вставки фрагментов из paragraphs.docx
def read_paragraphs_from_word(template_folder):
    """
    Читает абзацы из Word-файла, находящегося в папке с шаблонами.
    Сохраняет все форматирование исходного документа.
    Поддерживает подблоки, обозначенные числами (например, "1", "2").
    :param template_folder: Путь к папке с шаблонами (из строки Карты с кодом 3)
    :return: Словарь, где ключ — код (ОБ, ПАН), а значение — список элементов документа
             или словарь подблоков (если они есть) с сохраненным форматированием
    """
    paragraphs = {}
    doc_path = Path(template_folder) / "paragraphs.docx"  # Путь к paragraphs.docx
    
    if not doc_path.exists():
        raise FileNotFoundError(f"Файл paragraphs.docx не найден в папке {template_folder}.")
    
    try:
        doc = Document(doc_path)
        
        current_key = None
        current_elements = []
        current_sub_key = None
        
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            
            # Проверяем, является ли строка ключом (например, "ОБ:" или "PG+:")
            if text.endswith(':') and len(text.split()) == 1:
                # Сохраняем предыдущий ключ/подблок
                if current_key is not None:
                    if current_sub_key is not None:
                        if not isinstance(paragraphs.get(current_key), dict):
                            paragraphs[current_key] = {}
                        paragraphs[current_key][current_sub_key] = current_elements.copy()
                    elif current_elements:
                        paragraphs[current_key] = current_elements.copy()
                    current_elements.clear()
                    current_sub_key = None
                
                current_key = text[:-1]  # Убираем двоеточие
                
            elif current_key is not None and re.match(r'^\d+$', text):
                # Это подблок (например, "1" или "2")
                if current_sub_key is not None:
                    if not isinstance(paragraphs.get(current_key), dict):
                        paragraphs[current_key] = {}
                    paragraphs[current_key][current_sub_key] = current_elements.copy()
                    current_elements.clear()
                current_sub_key = text
                
            else:
                if current_key is not None:
                    current_elements.append(('paragraph', paragraph))
        
        # Обрабатываем таблицы в документе
        for table in doc.tables:
            if current_key is not None:
                current_elements.append(('table', table))
        
        # Добавляем последний ключ/подблок
        if current_key is not None:
            if current_sub_key is not None:
                if not isinstance(paragraphs.get(current_key), dict):
                    paragraphs[current_key] = {}
                paragraphs[current_key][current_sub_key] = current_elements.copy()
            elif current_elements:
                paragraphs[current_key] = current_elements.copy()
    
    except Exception as e:
        raise ValueError(f"Ошибка при чтении Word-файла: {e}")
    
    return paragraphs

def _paragraph_has_section_break(paragraph_element: Optional[Any]) -> bool:
    """Проверяет, содержит ли абзац разрыв раздела (<w:sectPr>)."""
    if paragraph_element is None:
        return False

    pPr = paragraph_element.find(qn('w:pPr'))
    if pPr is None:
        return False

    return pPr.find(qn('w:sectPr')) is not None

def _remove_section_break_from_paragraph(paragraph_element: Optional[Any]) -> bool:
    """Удаляет разрыв раздела (<w:sectPr>) из абзаца, если он есть."""
    if paragraph_element is None:
        return False

    pPr = paragraph_element.find(qn('w:pPr'))
    if pPr is None:
        return False

    sectPr = pPr.find(qn('w:sectPr'))
    if sectPr is None:
        return False

    pPr.remove(sectPr)
    return True

def _detach_section_break_from_paragraph(paragraph_element: Optional[Any]) -> Optional[Any]:
    """
    Извлекает разрыв раздела (<w:sectPr>) из абзаца и возвращает его как элемент.
    Сам абзац остаётся без этого разрыва.
    """
    if paragraph_element is None:
        return None

    pPr = paragraph_element.find(qn('w:pPr'))
    if pPr is None:
        return None

    sectPr = pPr.find(qn('w:sectPr'))
    if sectPr is None:
        return None

    pPr.remove(sectPr)
    return sectPr

def _get_or_create_paragraph_properties(paragraph_element: Any) -> Any:
    """Возвращает <w:pPr> абзаца, при необходимости создаёт его."""
    pPr = paragraph_element.find(qn('w:pPr'))
    if pPr is None:
        pPr = OxmlElement('w:pPr')
        paragraph_element.insert(0, pPr)
    return pPr

def _attach_section_break_to_paragraph(paragraph_element: Optional[Any], sectPr: Optional[Any]) -> None:
    """Добавляет разрыв раздела (<w:sectPr>) в указанный абзац."""
    if paragraph_element is None or sectPr is None:
        return

    pPr = _get_or_create_paragraph_properties(paragraph_element)

    existing_sectPr = pPr.find(qn('w:sectPr'))
    if existing_sectPr is not None:
        pPr.remove(existing_sectPr)

    pPrChange = pPr.find(qn('w:pPrChange'))
    if pPrChange is not None:
        pPr.insert(list(pPr).index(pPrChange), sectPr)
    else:
        pPr.append(sectPr)

def _remove_page_breaks_from_paragraph(paragraph_element: Optional[Any]) -> bool:
    """
    Удаляет из абзаца принудительные разрывы страницы:
    - <w:pageBreakBefore />;
    - <w:br w:type="page"/>;
    - <w:lastRenderedPageBreak/>.
    """
    if paragraph_element is None:
        return False

    removed = False

    pPr = paragraph_element.find(qn('w:pPr'))
    if pPr is not None:
        page_break_before = pPr.find(qn('w:pageBreakBefore'))
        if page_break_before is not None:
            pPr.remove(page_break_before)
            removed = True

    for br in paragraph_element.findall('.//' + qn('w:br')):
        if br.get(qn('w:type')) == 'page':
            br_parent = br.getparent()
            if br_parent is not None:
                br_parent.remove(br)
                removed = True

    for rendered_break in paragraph_element.findall('.//' + qn('w:lastRenderedPageBreak')):
        rendered_parent = rendered_break.getparent()
        if rendered_parent is not None:
            rendered_parent.remove(rendered_break)
            removed = True

    return removed

def _paragraph_has_content(paragraph_element: Optional[Any]) -> bool:
    """Проверяет, есть ли в абзаце полезное содержимое."""
    if paragraph_element is None:
        return False

    for t in paragraph_element.iter(qn('w:t')):
        if t.text and t.text.strip():
            return True

    content_tags = (
        'w:drawing',
        'w:pict',
        'w:object',
        'w:bookmarkStart',
        'w:commentReference',
    )

    for tag in content_tags:
        if paragraph_element.findall('.//' + qn(tag)):
            return True

    return False

def _clear_insert_marker(paragraph: Any) -> None:
    """Удаляет служебные метки INSERT и DYNAMIC_PAN_BLOCKS из абзаца."""
    try:
        if not hasattr(paragraph, 'text'):
            return

        new_text = re.sub(r'\{INSERT_[^}]+\}', '', paragraph.text)
        new_text = re.sub(r'\[DYNAMIC_PAN_BLOCKS\]', '', new_text)

        if new_text != paragraph.text:
            paragraph.text = new_text
    except Exception as e:
        logger.warning(f"Не удалось очистить служебную метку в абзаце: {e}")

def _cleanup_dynamic_placeholder(paragraph: Any) -> None:
    """
    Удаляет метку [DYNAMIC_PAN_BLOCKS].
    Если абзац после удаления метки пуст и не содержит разрыв раздела, абзац удаляется.
    Если абзац содержит разрыв раздела, он сохраняется, чтобы не потерять структуру раздела.
    """
    if "[DYNAMIC_PAN_BLOCKS]" not in paragraph.text:
        return

    paragraph.text = paragraph.text.replace("[DYNAMIC_PAN_BLOCKS]", "")

    if paragraph.text.strip():
        return

    if _paragraph_has_section_break(paragraph._element):
        logger.debug(
            "Плейсхолдер [DYNAMIC_PAN_BLOCKS] содержит разрыв раздела. "
            "Абзац сохранён для сохранения структуры документа."
        )
        return

    parent = paragraph._element.getparent()
    if parent is None:
        logger.warning(
            "Не удалось определить родительский элемент для удаления "
            "плейсхолдера [DYNAMIC_PAN_BLOCKS]."
        )
        return

    parent.remove(paragraph._element)

def _remove_empty_paragraphs_after_element(parent: Any, after_element: Optional[Any], max_remove: int = 100) -> int:
    """
    Удаляет пустые абзацы сразу после указанного элемента.
    Останавливается, если встречает абзац с разрывом раздела, содержимым или другой значимый элемент.
    """
    if parent is None or after_element is None:
        return 0

    removed = 0
    sibling = after_element.getnext()

    while sibling is not None and removed < max_remove:
        next_sibling = sibling.getnext()

        if sibling.tag == qn('w:sectPr'):
            break

        if sibling.tag == qn('w:p'):
            if _paragraph_has_section_break(sibling):
                break

            if _paragraph_has_content(sibling):
                break

            parent.remove(sibling)
            removed += 1
        else:
            break

        sibling = next_sibling

    if removed:
        logger.debug(f"Удалено пустых абзацев после вставленного блока: {removed}")

    return removed

def insert_word_elements(doc: Any, elements: List[Tuple[str, Any]], placeholder_paragraph: Any) -> None:
    """
    Вставляет элементы Word с сохранением форматирования в указанный параграф-заполнитель.

    Важные изменения:
    - не используются doc.add_paragraph() и doc.add_table(), чтобы не создавать
      лишние пустые абзацы и таблицы в конце документа;
    - из вставляемых абзацев удаляются разрывы раздела и разрывы страницы;
    - разрыв раздела из абзаца-плейсхолдера сохраняется и переносится
      в последний вставленный абзац.
    """
    parent = placeholder_paragraph._p.getparent()
    if parent is None:
        logger.error("Не удалось определить родительский элемент для плейсхолдера INSERT. Вставка пропущена.")
        return

    index = parent.index(placeholder_paragraph._p)

    # Если плейсхолдер сам содержит разрыв раздела, сохраняем его.
    placeholder_sectPr = _detach_section_break_from_paragraph(placeholder_paragraph._p)

    if not elements:
        _clear_insert_marker(placeholder_paragraph)

        if placeholder_sectPr is not None:
            _attach_section_break_to_paragraph(placeholder_paragraph._p, placeholder_sectPr)
            logger.debug(
                "Список вставляемых элементов пуст. "
                "Разрыв раздела сохранён в абзаце-плейсхолдере."
            )
        elif not placeholder_paragraph.text.strip():
            parent.remove(placeholder_paragraph._p)

        return

    # Удаляем заполнитель до вставки новых элементов.
    parent.remove(placeholder_paragraph._p)

    last_inserted_paragraph: Optional[Any] = None

    for elem_type, elem_content in elements:
        if elem_type == 'paragraph':
            new_p = deepcopy(elem_content._p)

            if _remove_section_break_from_paragraph(new_p):
                logger.debug("Из вставляемого абзаца удалён разрыв раздела.")

            if _remove_page_breaks_from_paragraph(new_p):
                logger.debug("Из вставляемого абзаца удалён разрыв страницы.")

            parent.insert(index, new_p)
            index += 1
            last_inserted_paragraph = new_p

        elif elem_type == 'table':
            new_tbl = deepcopy(elem_content._tbl)
            parent.insert(index, new_tbl)
            index += 1

            # После таблицы разрыв раздела можно сохранить только отдельным абзацем.
            last_inserted_paragraph = None

        else:
            logger.warning(f"Неизвестный тип элемента для вставки: {elem_type}")

    if placeholder_sectPr is not None:
        if last_inserted_paragraph is not None:
            _attach_section_break_to_paragraph(last_inserted_paragraph, placeholder_sectPr)
            logger.debug("Разрыв раздела из плейсхолдера перенесён в последний вставленный абзац.")
        else:
            empty_p = OxmlElement('w:p')
            _attach_section_break_to_paragraph(empty_p, placeholder_sectPr)
            parent.insert(index, empty_p)
            logger.debug(
                "Разрыв раздела из плейсхолдера сохранён в новом пустом абзаце, "
                "так как последним вставленным элементом была не абзац-структура."
            )

# Обрабатывает параграфы и удаляет абзацы с метками INSERT_XXX, если значение пусто
def preprocess_paragraphs(doc_path, paragraphs, card_data):
    doc = Document(doc_path)
    paragraphs_to_remove = []

    def process_paragraph(paragraph):
        """
        Обрабатывает один параграф. 
        Возвращает True, если параграф нужно удалить (пустое значение в карте).
        """
        text = paragraph.text
        found_placeholders = re.findall(r'\{INSERT_([^}]+)\}', text)
        if not found_placeholders:
            return False
        
        # Проверяем, нужно ли удалять параграф (если хотя бы один ключ пустой)
        for inner in found_placeholders:
            parts = inner.rsplit('_', 1)
            if len(parts) == 2 and parts[0] in card_data:
                key = parts[0]
            else:
                key = inner
            
            if key in card_data and not card_data[key]:
                return True # Помечаем на удаление
        
        # Если не удаляем, обрабатываем вставки
        for inner in found_placeholders:
            parts = inner.rsplit('_', 1)
            if len(parts) == 2 and parts[0] in card_data:
                key, sub_key = parts[0], parts[1]
            else:
                key, sub_key = inner, None
                
            placeholder = f"{{INSERT_{inner}}}"
            
            if key not in card_data or not card_data[key]:
                # Если ключа нет в карте или он пустой, удаляем метку из текста,
                # чтобы не осталось пустых строк после финальной очистки
                paragraph.text = paragraph.text.replace(placeholder, "")
                logger.debug(f"Метка {placeholder} удалена, так как ключ '{key}' отсутствует или пуст.")
                continue 
                
            map_value = str(card_data[key]).strip()
            block = paragraphs.get(map_value)
            if block is None:
                logger.warning(f"Для значения '{map_value}' (ключ {key}) не найден блок в paragraphs. Метка {placeholder} будет удалена.")
                paragraph.text = paragraph.text.replace(placeholder, "")
                continue
            
            elements_to_insert = None
            if sub_key is not None:
                if isinstance(block, dict) and sub_key in block:
                    elements_to_insert = block[sub_key]
                else:
                    logger.warning(f"Подблок '{sub_key}' не найден в блоке '{map_value}'. Метка {placeholder} будет удалена.")
                    paragraph.text = paragraph.text.replace(placeholder, "")
                    continue
            else:
                if isinstance(block, list):
                    elements_to_insert = block
                else:
                    logger.warning(f"Блок '{map_value}' содержит подблоки, но в шаблоне использован {placeholder} без номера. Метка будет удалена.")
                    paragraph.text = paragraph.text.replace(placeholder, "")
                    continue
            
            logger.info(f"Вставка блока '{map_value}' (подблок '{sub_key}') вместо {placeholder}")
            insert_word_elements(doc, elements_to_insert, paragraph)
            return False # Параграф уже удален внутри insert_word_elements
            
        # Если после всех замен параграф стал пустым, помечаем его на удаление
        if not paragraph.text.strip():
            return True
            
        return False

    # Обрабатываем все параграфы в документе
    for paragraph in doc.paragraphs:
        if process_paragraph(paragraph):
            paragraphs_to_remove.append(paragraph)

    # Удаляем помеченные абзацы (в обратном порядке)
    for paragraph in reversed(paragraphs_to_remove):
        p = paragraph._element
        p.getparent().remove(p)
    
    # Обрабатываем таблицы
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                cell_paragraphs_to_remove = []
                for paragraph in cell.paragraphs:
                    if process_paragraph(paragraph):
                        cell_paragraphs_to_remove.append(paragraph)
                
                # Удаляем пустые абзацы в ячейках
                for paragraph in reversed(cell_paragraphs_to_remove):
                    p = paragraph._element
                    p.getparent().remove(p)
    
    doc.save(doc_path)

# Заполнение таблицы RT
def process_rt_table(doc_path, card_data):
    """
    Полная обработка таблицы RT с:
    - сохранением форматирования
    - проверкой наличия маркера
    - безопасной обработкой отсутствующих таблиц
    """
    logger.info("Начало обработки таблицы RT")
    doc = Document(doc_path)
    tables_with_marker = 0  # Счётчик найденных таблиц
    
    # Подготовка данных
    # Проверяем наличие объектов в карте
    ob_exists = bool(card_data.get("ОБ"))
    ob2_exists = bool(card_data.get("ОБ2"))
    ob3_exists = bool(card_data.get("ОБ3"))
    
    # Формируем список объектов, которые нужно оставить
    objects_to_keep = ["ОБ"]
    if ob2_exists:
        objects_to_keep.append("ОБ2")
    if ob3_exists:
        objects_to_keep.append("ОБ3")
    
    # Формируем фразу для KV
    kv_parts = ["KV"]
    if ob2_exists:
        kv_parts.append("KV-2")
    if ob3_exists:
        kv_parts.append("KV-3")
    kv_text = ", ".join(kv_parts) + " (отрицательный контроль при выделении ДНК)"
    
    # Поиск таблиц
    for table in doc.tables:
        marker_found = False
        header_cell = None
        
        # Поиск маркера в таблице
        for row in table.rows:
            for cell in row.cells:
                if "[INSERT_RT_DATA]" in cell.text:
                    header_cell = cell
                    marker_found = True
                    break
            if marker_found:
                break
        
        if not marker_found:
            continue  # Пропускаем таблицы без маркера
        
        tables_with_marker += 1
        logger.info(f"Обработка таблицы #{tables_with_marker} с маркером [INSERT_RT_DATA]")
        
        # Обработка таблицы
        # Удаляем метку из заголовка
        if header_cell.paragraphs:
            for paragraph in header_cell.paragraphs:
                if paragraph.runs:
                    for run in paragraph.runs:
                        run.text = run.text.replace("[INSERT_RT_DATA]", "")
        
        # Поиск ячейки с {KV}
        kv_cell = None
        for row in table.rows:
            for cell in row.cells:
                if "{KV}" in cell.text:
                    kv_cell = cell
                    break
            if kv_cell:
                break
        
        # Сбор информации о строках объектов
        rows_info = []
        for row in table.rows:
            row_text = " ".join(cell.text.strip() for cell in row.cells)
            
            if "Объект {ОБ}" in row_text:
                rows_info.append(("ОБ", row))
            elif "Объект {ОБ2}" in row_text:
                rows_info.append(("ОБ2", row))
            elif "Объект {ОБ3}" in row_text:
                rows_info.append(("ОБ3", row))
        
        # Удаление лишних строк
        for obj_code, row in reversed(rows_info):
            if obj_code not in objects_to_keep:
                try:
                    table._tbl.remove(row._tr)
                    logger.info(f"Удалена строка для {obj_code}")
                except ValueError:
                    logger.warning(f"Строка для {obj_code} уже удалена")
        
        # Заполнение данных
        for obj_code, row in rows_info:
            if obj_code in objects_to_keep:
                obj_num = "" if obj_code == "ОБ" else obj_code[-1]
                
                for cell in row.cells:
                    if cell.paragraphs:
                        for paragraph in cell.paragraphs:
                            if paragraph.runs:
                                for run in paragraph.runs:
                                    original_text = run.text
                                    new_text = original_text.replace(
                                        f"{{{obj_code}}}", 
                                        str(card_data.get(obj_code, "")))
                                    
                                    rt_codes = ["ДМ", "КМ", "YХ", "ИД"]
                                    for rt_code in rt_codes:
                                        full_code = f"{rt_code}{obj_num}"
                                        new_text = new_text.replace(
                                            f"{{{full_code}}}", 
                                            str(card_data.get(full_code, "")))
                                    
                                    run.text = new_text
        
        # Обработка KV
        if kv_cell:
            if kv_cell.paragraphs and kv_cell.paragraphs[0].runs:
                first_run = kv_cell.paragraphs[0].runs[0]
                font = first_run.font
                
                # Сохраняем форматирование
                formatting = {
                    'bold': font.bold,
                    'italic': font.italic,
                    'underline': font.underline,
                    'size': font.size,
                    'name': font.name,
                    'color': font.color.rgb if font.color else None
                }
                
                # Заменяем текст
                first_run.text = kv_text
                
                # Восстанавливаем форматирование
                for attr, value in formatting.items():
                    if value is not None:
                        setattr(font, attr, value)
                
                # Удаляем лишние runs
                for run in kv_cell.paragraphs[0].runs[1:]:
                    kv_cell.paragraphs[0]._p.remove(run._r)
    
    # Финализация
    if tables_with_marker == 0:
        logger.warning("В документе нет таблиц с маркером [INSERT_RT_DATA]")
    else:
        logger.info(f"Успешно обработано таблиц: {tables_with_marker}")
    
    doc.save(doc_path)
    logger.info("Обработка завершена")

# Замена меток по словарю
def process_switch_tags(doc_path, card_data, switch_dict):
    """
    Обрабатывает все метки вида {PREFIX_SWITCH:word} в документе Word.
    Сохраняет форматирование и регистр, добавляет логирование.
    """
    logger.info("Начало обработки switch-меток в документе")
    
    try:
        doc = Document(doc_path)
        total_replacements = 0
        
        def process_text(text):
            nonlocal total_replacements
            new_text = text
            # Ищем все метки вида {PREFIX_SWITCH:word}
            pattern = r'\{(\w+)_SWITCH:([^}]+)\}'
            matches = re.findall(pattern, text)
            for switch_type, original_word in matches:
                try:
                    # Логируем найденную метку
                    logger.info(f"[SWITCH] Найдена метка: {{{switch_type}_SWITCH:{original_word}}}")
                    # Определяем, нужно ли множественное число
                    use_plural = False
                    if switch_type == "OB2":
                        ob2_value = card_data.get('ОБ2')
                        use_plural = bool(ob2_value) if ob2_value is not None else False
                        logger.debug(f"Значение ОБ2: '{ob2_value}', use_plural: {use_plural}")
                    elif switch_type == "PAN2":
                        pan2_value = card_data.get('ПАН2')
                        if pan2_value is not None:
                            pan2_str = str(pan2_value).strip()
                            # Множественное число, если значение не пустое и не прочерк
                            use_plural = pan2_str != "" and pan2_str != "-"
                        else:
                            use_plural = False
                        logger.info(f"[SWITCH] Значение ПАН2: '{pan2_value}', use_plural: {use_plural}")
                    # Получаем варианты замены из словаря
                    replacement_dict = switch_dict.get(f"{switch_type}_SWITCH", {})
                    # ИСПРАВЛЕНИЕ: применяем strip() к original_word для защиты от пробелов
                    lookup_word = original_word.strip().lower()
                    logger.debug(f"[SWITCH] Ищем слово '{lookup_word}' в словаре {switch_type}_SWITCH")
                    replacements = replacement_dict.get(lookup_word)
                    if not replacements or len(replacements) < 2:
                        logger.warning(f"[SWITCH] Не найдена замена для слова '{lookup_word}' в словаре {switch_type}_SWITCH. Доступные ключи: {list(replacement_dict.keys())}")
                        continue
                    # Выбираем нужную форму слова
                    replacement = replacements[1] if use_plural else replacements[0]
                    logger.info(f"[SWITCH] Выбрана замена: '{original_word}' -> '{replacement}' (use_plural={use_plural})")
                    # Сохраняем регистр исходного слова
                    if original_word.strip().istitle():
                        replacement = replacement.title()
                    elif original_word.strip().isupper():
                        replacement = replacement.upper()
                    # Заменяем метку в тексте
                    old_text = new_text
                    new_text = new_text.replace(
                        f"{{{switch_type}_SWITCH:{original_word}}}", 
                        replacement
                    )
                    if old_text != new_text:
                        total_replacements += 1
                        logger.info(f"[SWITCH] Замена выполнена: '{old_text}' -> '{new_text}'")
                except Exception as e:
                    logger.error(f"[SWITCH] Ошибка при обработке метки {switch_type}_SWITCH:{original_word}: {str(e)}", exc_info=True)
                    continue
            return new_text
        
        # Обрабатываем все параграфы
        for paragraph in doc.paragraphs:
            original_text = paragraph.text
            new_text = process_text(original_text)
            if new_text != original_text:
                paragraph.text = new_text
        
        # Обрабатываем таблицы
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        original_text = paragraph.text
                        new_text = process_text(original_text)
                        if new_text != original_text:
                            paragraph.text = new_text
        
        # Обрабатываем колонтитулы
        for section in doc.sections:
            for paragraph in section.header.paragraphs:
                original_text = paragraph.text
                new_text = process_text(original_text)
                if new_text != original_text:
                    paragraph.text = new_text
            
            for paragraph in section.footer.paragraphs:
                original_text = paragraph.text
                new_text = process_text(original_text)
                if new_text != original_text:
                    paragraph.text = new_text
        
        logger.info(f"Обработка switch-меток завершена. Всего замен: {total_replacements}")
        doc.save(doc_path)
    
    except Exception as e:
        logger.error(f"Критическая ошибка при обработке switch-меток: {str(e)}")
        raise

def clean_text_from_none_and_spaces(text, max_allowed_spaces=3):
    """
    Универсальная функция для очистки текста от меток с None и лишних пробелов.
    Заменяет последовательности пробелов длиной меньше max_allowed_spaces на один пробел,
    но сохраняет более длинные последовательности (предполагая, что это намеренное форматирование).
    
    :param text: Исходный текст с метками
    :param max_allowed_spaces: Максимальное количество пробелов, которые считаются "случайными" и подлежат очистке
    :return: Очищенный текст
    """
    logger.debug(f"Очистка текста от None и лишних пробелов: '{text}' (макс. пробелов: {max_allowed_spaces})")
    
    # Шаг 1: Заменяем последовательности пробелов короче max_allowed_spaces на один пробел
    # Используем функцию замены с условием
    def replace_short_spaces(match):
        spaces = match.group(0)
        if len(spaces) < max_allowed_spaces:
            return ' '
        else:
            return spaces  # Сохраняем длинные последовательности пробелов
    
    # Применяем замену только к последовательностям пробелов (не табуляциям и другим whitespace)
    cleaned_text = re.sub(r' +', replace_short_spaces, text)
    
    # Шаг 2: Убираем пробелы в начале и конце строки
    cleaned_text = cleaned_text.strip()
    
    # Шаг 3: Убираем пробелы перед знаками препинания (только одиночные пробелы)
    cleaned_text = re.sub(r'\s+([.,)])', r'\1', cleaned_text)
    cleaned_text = re.sub(r'([(])\s+', r'\1', cleaned_text)
    
    logger.debug(f"Текст после очистки: '{cleaned_text}'")
    return cleaned_text

def build_highlighted_alleles_text(alleles_str, reference_alleles_set):
    """
    Формирует список частей текста с выделением фона для несовпадающих аллелей.
    Выделение применяется ТОЛЬКО если НИ ОДНА из аллелей рабочего профиля не совпадает с образцом.
    
    :param alleles_str: строка с аллелями, разделёнными запятой (например "10,14" или "-")
    :param reference_alleles_set: множество аллелей образца (строки)
    :return: список кортежей (текст, highlight_color), где highlight_color может быть None (без выделения)
    """
    logger.debug(f"Формирование выделенного текста для аллелей '{alleles_str}', образец: {reference_alleles_set}")
    
    # Если в образце нет данных — не выделяем ничего
    if not reference_alleles_set:
        logger.debug("В образце нет данных для этого локуса — выделение отключено")
        return [(alleles_str, None)]
    
    # Если прочерк – не выделяем
    if alleles_str == "-" or not alleles_str.strip():
        logger.debug("Обнаружен прочерк – возвращаем без выделения")
        return [("-", None)]
    
    # Разбиваем по запятой, убираем лишние пробелы
    alleles = [a.strip() for a in alleles_str.split(",") if a.strip()]
    if not alleles:
        return [("-", None)]
    
    # Проверяем, есть ли хотя бы одно совпадение
    has_common = any(allele in reference_alleles_set for allele in alleles)
    
    if has_common:
        # Если есть хотя бы одна общая аллель — не выделяем ничего
        logger.debug(f"Есть общая аллель с образцом — выделение не применяется")
        return [(alleles_str, None)]
    else:
        # Нет общих аллелей — выделяем все аллели
        logger.debug(f"Нет общих аллелей — все аллели будут выделены")
        result = []
        for i, allele in enumerate(alleles):
            result.append((allele, WD_COLOR_INDEX.GRAY_25))
            if i < len(alleles) - 1:
                result.append((",", None))
        return result

# Функция для замены меток в Word
def replace_in_doc(doc_path, replacements):
    """ Заменяет метки в шаблоне Word """
    doc = Document(doc_path)
    found_keys = set()
    
    def replace_text(element):
        nonlocal found_keys
        if not hasattr(element, "runs"):
            return

        # Собираем полный текст параграфа
        full_text = "".join(run.text for run in element.runs)
        if not full_text:  # Если параграф пустой, пропускаем
            return

        # Создаем карту форматирования для каждого символа
        char_formatting = []
        current_pos = 0
        for run in element.runs:
            for _ in run.text:
                char_formatting.append({
                    'bold': run.bold,
                    'italic': run.italic,
                    'underline': run.underline,
                    'font': run.font.name,
                    'size': run.font.size
                })
            current_pos += len(run.text)

        # Создаем новый текст с заменами
        new_text = full_text
        positions = []  # Список кортежей (позиция, длина_старого_текста, новый_текст)

        # Сначала обрабатываем TIME
        if "{TIME}" in new_text:
            current_time = datetime.datetime.now().strftime("%H.%M")
            start_pos = new_text.find("{TIME}")
            if start_pos != -1:
                positions.append((start_pos, len("{TIME}"), current_time))
                found_keys.add("TIME")

        # Затем обрабатываем остальные метки
        for key, value in replacements.items():
            patterns = [f"{{{key}}}", f"{{ {key} }}", f"{{{key} }}", f"{{ {key}}}", f"{{  {key}}}"]
            for pattern in patterns:
                start_pos = 0
                while True:
                    pos = new_text.find(pattern, start_pos)
                    if pos == -1:
                        break
                    positions.append((pos, len(pattern), str(value)))
                    found_keys.add(key)
                    start_pos = pos + len(pattern)

            # Обработка новых меток {KEY:?, {}}
            pattern = f"{{{key}:?, {{}}}}"
            start_pos = 0
            while True:
                pos = new_text.find(pattern, start_pos)
                if pos == -1:
                    break
                # Если значение есть → подставляем ", значение", иначе пустую строку
                replacement = f", {value}" if value else ""
                positions.append((pos, len(pattern), replacement))
                found_keys.add(key)
                start_pos = pos + len(pattern)

        # Сортируем позиции в обратном порядке, чтобы замены не влияли на последующие позиции
        positions.sort(reverse=True)

        # Выполняем замены и обновляем карту форматирования
        for pos, old_len, new_text_piece in positions:
            new_text = new_text[:pos] + new_text_piece + new_text[pos + old_len:]
            # Обновляем карту форматирования
            format_at_pos = char_formatting[pos]
            char_formatting[pos:pos + old_len] = [format_at_pos] * len(new_text_piece)

        new_text = re.sub(r'\{[^\}]*\}', '', new_text)

        # Очищаем от лишних пробелов после всех замен
        new_text = clean_text_from_none_and_spaces(new_text)

        # Очищаем существующие runs
        for run in element.runs:
            run._element.getparent().remove(run._element)

        # Создаем новые runs с правильным форматированием
        current_format = None
        current_text = ""
        
        for i, char in enumerate(new_text):
            char_format = char_formatting[i] if i < len(char_formatting) else char_formatting[-1]
            
            if current_format != char_format:
                if current_text:
                    run = element.add_run(current_text)
                    if current_format:
                        run.bold = current_format['bold']
                        run.italic = current_format['italic']
                        run.underline = current_format['underline']
                        if current_format['font']:
                            run.font.name = current_format['font']
                        if current_format['size']:
                            run.font.size = current_format['size']
                current_text = char
                current_format = char_format
            else:
                current_text += char

        # Добавляем последний run
        if current_text:
            run = element.add_run(current_text)
            if current_format:
                run.bold = current_format['bold']
                run.italic = current_format['italic']
                run.underline = current_format['underline']
                if current_format['font']:
                    run.font.name = current_format['font']
                if current_format['size']:
                    run.font.size = current_format['size']

    # Обрабатываем все элементы документа
    for paragraph in doc.paragraphs:
        replace_text(paragraph)
        
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replace_text(paragraph)
                    
    for section in doc.sections:
        for paragraph in section.header.paragraphs:
            replace_text(paragraph)
        for paragraph in section.footer.paragraphs:
            replace_text(paragraph)

        # Обработка номеров страниц
        footer = section.footer
        for paragraph in footer.paragraphs:
            if "стр. [СТ] из [КОЛ]" in paragraph.text:
                paragraph.text = paragraph.text.replace("стр. [СТ] из [КОЛ]", "")
                run = paragraph.add_run("стр. ")
                
                fldChar1 = OxmlElement('w:fldChar')
                fldChar1.set(qn('w:fldCharType'), 'begin')
                instrText = OxmlElement('w:instrText')
                instrText.set(qn('xml:space'), 'preserve')
                instrText.text = 'PAGE \* Arabic'
                fldChar2 = OxmlElement('w:fldChar')
                fldChar2.set(qn('w:fldCharType'), 'end')
                
                run._r.append(fldChar1)
                run._r.append(instrText)
                run._r.append(fldChar2)
                
                run = paragraph.add_run(" из ")
                
                fldChar3 = OxmlElement('w:fldChar')
                fldChar3.set(qn('w:fldCharType'), 'begin')
                instrText2 = OxmlElement('w:instrText')
                instrText2.set(qn('xml:space'), 'preserve')
                instrText2.text = 'NUMPAGES \* Arabic'
                fldChar4 = OxmlElement('w:fldChar')
                fldChar4.set(qn('w:fldCharType'), 'end')
                
                run._r.append(fldChar3)
                run._r.append(instrText2)
                run._r.append(fldChar4)

    # Проверяем, все ли метки были заменены
    missing_keys = set(replacements.keys()) - found_keys
    if missing_keys:
        logging.warning(f"Предупреждение: следующие метки не найдены в документе: {missing_keys}")

    doc.save(doc_path)

# ...
def generate_dynamic_pan_blocks(doc_path: str, card_data: dict) -> None:
    """
    Генерирует динамические блоки для ПАН2, ПАН3, ПАН4.
    Ищет в документе эталонные параграфы с [IMAGE_1_3] и {2ИЭ},
    клонирует их XML и вставляет на место метки [DYNAMIC_PAN_BLOCKS].

    При клонировании удаляются:
    - разрывы раздела (<w:sectPr>);
    - разрывы страницы (<w:pageBreakBefore />, <w:br w:type="page"/>);
    - служебные маркеры <w:lastRenderedPageBreak/>.
    """
    logger.info("Начало генерации динамических блоков ПАН...")
    doc = Document(doc_path)

    # 1. Определяем количество блоков для генерации (по 3 на каждый заполненный ПАН)
    pan_keys = ["ПАН2", "ПАН3", "ПАН4"]
    blocks_to_generate = 0
    for key in pan_keys:
        val = card_data.get(key)
        # Считаем заполненным, если значение есть и это не пустота/прочерк/None
        if val is not None and str(val).strip() not in ("", "-", "None"):
            blocks_to_generate += 3

    if blocks_to_generate == 0:
        logger.info("Значения ПАН2/ПАН3/ПАН4 отсутствуют. Генерация блоков не требуется.")

        # Удаляем метку-плейсхолдер, чтобы она не осталась в финальном документе
        for p in doc.paragraphs:
            _cleanup_dynamic_placeholder(p)

        doc.save(doc_path)
        return

    logger.info(
        f"Обнаружены заполненные ПАН. "
        f"Необходимо сгенерировать {blocks_to_generate} блоков (картинок)."
    )

    # 2. Ищем эталонные параграфы и метку для вставки
    image_1_3_para = None
    expert_para = None
    placeholder_para = None

    for p in doc.paragraphs:
        if "[IMAGE_1_3]" in p.text:
            image_1_3_para = p
        elif image_1_3_para and not expert_para and "{2ИЭ}" in p.text:
            # Берем первый параграф после [IMAGE_1_3], содержащий {2ИЭ}
            expert_para = p

        if "[DYNAMIC_PAN_BLOCKS]" in p.text:
            placeholder_para = p

    if not image_1_3_para or not expert_para:
        logger.warning(
            "Не найдены эталонные параграфы с [IMAGE_1_3] и {2ИЭ}. "
            "Генерация пропущена."
        )

        if placeholder_para:
            _cleanup_dynamic_placeholder(placeholder_para)

        doc.save(doc_path)
        return

    if not placeholder_para:
        logger.warning("Не найдена метка [DYNAMIC_PAN_BLOCKS] в документе. Вставлять некуда.")
        doc.save(doc_path)
        return

    # 3. Клонируем XML и вставляем
    parent = placeholder_para._element.getparent()
    if parent is None:
        logger.error(
            "Не удалось определить родительский элемент для плейсхолдера "
            "[DYNAMIC_PAN_BLOCKS]. Вставка пропущена."
        )
        _cleanup_dynamic_placeholder(placeholder_para)
        doc.save(doc_path)
        return

    insert_index = parent.index(placeholder_para._element)

    # Если плейсхолдер сам содержит разрыв раздела, сохраняем его,
    # чтобы не потерять границу разделов при удалении плейсхолдера.
    placeholder_sectPr = _detach_section_break_from_paragraph(placeholder_para._element)
    if placeholder_sectPr is not None:
        logger.debug(
            "Из плейсхолдера [DYNAMIC_PAN_BLOCKS] извлечён разрыв раздела "
            "для последующего переноса в последний вставленный абзац."
        )

    current_image_num = 4
    last_inserted_paragraph: Optional[Any] = None

    for _ in range(blocks_to_generate):
        # Клонируем параграф с картинкой
        new_img_p = deepcopy(image_1_3_para._element)

        # Обязательно удаляем разрыв раздела из клона, если он туда попал
        if _remove_section_break_from_paragraph(new_img_p):
            logger.debug("Из клонированного параграфа изображения удалён разрыв раздела.")

        # Удаляем возможные принудительные разрывы страницы
        if _remove_page_breaks_from_paragraph(new_img_p):
            logger.debug("Из клонированного параграфа изображения удалён разрыв страницы.")

        # Ищем все текстовые узлы <w:t> внутри параграфа
        found_and_replaced = False
        for t in new_img_p.iter(qn('w:t')):
            if t.text and "[IMAGE_1_3]" in t.text:
                t.text = t.text.replace("[IMAGE_1_3]", f"[IMAGE_1_{current_image_num}]")
                t.set(qn('xml:space'), 'preserve')
                found_and_replaced = True
                logger.debug(f"Успешно заменено на [IMAGE_1_{current_image_num}]")

        if not found_and_replaced:
            logger.warning(
                f"Не удалось найти текст '[IMAGE_1_3]' внутри клонированного параграфа "
                f"для картинки 1_{current_image_num}. "
                "Проверь, не разбит ли текст на части в шаблоне Word."
            )

        parent.insert(insert_index, new_img_p)
        insert_index += 1

        # Клонируем параграф с экспертом (текст {2ИЭ} заменится позже в replace_in_doc)
        new_exp_p = deepcopy(expert_para._element)

        # Обязательно удаляем разрыв раздела из клона, если он туда попал
        if _remove_section_break_from_paragraph(new_exp_p):
            logger.debug("Из клонированного параграфа эксперта удалён разрыв раздела.")

        # Удаляем возможные принудительные разрывы страницы
        if _remove_page_breaks_from_paragraph(new_exp_p):
            logger.debug("Из клонированного параграфа эксперта удалён разрыв страницы.")

        parent.insert(insert_index, new_exp_p)
        insert_index += 1

        last_inserted_paragraph = new_exp_p
        current_image_num += 1

    # 4. Возвращаем разрыв раздела плейсхолдера, если он был
    if placeholder_sectPr is not None:
        if last_inserted_paragraph is not None:
            _attach_section_break_to_paragraph(last_inserted_paragraph, placeholder_sectPr)
            logger.debug(
                "Разрыв раздела из плейсхолдера перенесён в последний вставленный абзац."
            )
        else:
            empty_p = OxmlElement('w:p')
            _attach_section_break_to_paragraph(empty_p, placeholder_sectPr)
            parent.insert(insert_index, empty_p)
            last_inserted_paragraph = empty_p
            logger.debug(
                "Разрыв раздела из плейсхолдера сохранён в новом пустом абзаце, "
                "так как не найдено последнего абзаца для переноса."
            )

    # 5. Удаляем параграф с плейсхолдером
    parent.remove(placeholder_para._element)

    # 6. Безопасно удаляем пустые абзацы, которые могли остаться сразу после вставки
    removed_empty = _remove_empty_paragraphs_after_element(parent, last_inserted_paragraph)
    if removed_empty:
        logger.debug(
            f"После генерации ПАН-блоков удалено пустых абзацев: {removed_empty}."
        )

    logger.info(
        f"Успешно сгенерировано {blocks_to_generate} пар параграфов для ПАН "
        f"(картинки 1_4 ... 1_{current_image_num - 1})."
    )

    doc.save(doc_path)

# Функция для вставки изображений
def insert_images(doc_path, images_folder, card_data, brightness_factor=1.5, contrast_factor=1.5):
    """
    Вставляет изображения в документ Word с настройкой яркости и контрастности.
    Яркость и контрастность регулируются только для обычных изображений.
    Обычные изображения получают черную рамку толщиной 0.5 пт.
    """
    doc = Document(doc_path)
    images_folder = Path(images_folder)
    
    # Проверяем наличие обычных изображений
    ordinary_images = [img for img in images_folder.glob("*.*") 
                     if img.stem not in {"1", "1_2", "1_3", "1_4", "1_5", "1_6", "1_7", "1_8", "1_9", "1_10", "1_11", "1_12", "1_13"} 
                     and img.suffix.lower() in {".jpg", ".jpeg", ".png"}]
    if not ordinary_images:
        raise FileNotFoundError(f"Не найдено ни одного обычного изображения в папке {images_folder}")
    
    # Проверяем наличие специальных изображений, если это требуется
    if (card_data.get("1") != "Кость СВО нет пригодной ДНК") and (card_data.get("1") != "СВО_кость_нет результата"):
    #if card_data.get("1") != "СВО_кость_нет результата":
        special_images_found = any(img_path.exists() 
                                for img_name in ["1", "1_2", "1_3", "1_4", "1_5", "1_6", "1_7", "1_8", "1_9", "1_10", "1_11", "1_12", "1_13"] 
                                for ext in [".jpg", ".jpeg", ".png"] 
                                for img_path in [images_folder / f"{img_name}{ext}"])
        if not special_images_found:
            raise FileNotFoundError(f"Не найдено ни одного специального изображения (1, 1_2, 1_3) в папке {images_folder}")
    
    # Определяем высоту изображения в зависимости от значения в поле "1"
    card_value = card_data.get("1", "").strip()
    if card_value == "СВО_Молов_образец_родственники":
        image_height = Inches(1.8)  # Высота для "СВО_Молов_образец_родственники"
    elif card_value in ["СВО_Молов_образец_прямая идентификация", "СВО_Молов БВП_образец_прямая идентификация"]:
        image_height = Inches(1.8)  # Высота для "СВО_Молов_образец_прямая идентификация"
    elif card_value == "СВО_Ростов_образец_родственники":
        image_height = Inches(2.0)  # Высота для "СВО_Ростов_образец_родственники"
    elif card_value == "СВО_Ростов_образец_прямая идентификация":
        image_height = Inches(2.0)  # Высота для "СВО_Ростов_образец_прямая идентификация"
    elif card_value == "СВО_кость":
        image_height = Inches(2.1)  # Высота для "СВО_кость"
    elif card_value == "СВО_кость_нет результата":
        image_height = Inches(2.2)  # Высота для "СВО_кость_нет результата"
    elif card_value == "ЭКЦ":
        image_height = Inches(2.1)  # Высота для "ЭКЦ"
    else:
        image_height = Inches(1.8)  # Значение по умолчанию
    
    # Счетчик для подписей к фото
    photo_counter = 1

    # Вставляем обычные изображения (не 1, 1_2, 1_3)
    ordinary_images = [img for img in images_folder.glob("*.*") 
                    if img.stem not in {"1", "1_2", "1_3", "1_4", "1_5", "1_6", "1_7", "1_8", "1_9", "1_10", "1_11", "1_12", "1_13"} 
                    and img.suffix.lower() in {".jpg", ".jpeg", ".png"}]
    if not ordinary_images:
        raise FileNotFoundError(f"Не найдено ни одного обычного изображения в папке {images_folder}")

    # Сортируем изображения по имени для предсказуемого порядка
    ordinary_images_sorted = sorted(ordinary_images, key=lambda x: x.name)

    # Находим все таблицы с меткой [IMAGE_OTHER] в документе
    tables_with_marker = []
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if "[IMAGE_OTHER]" in paragraph.text:
                        tables_with_marker.append((table, row, cell, paragraph))
                        break

    if not tables_with_marker:
        raise ValueError("В документе не найдена таблица с меткой [IMAGE_OTHER]")

    # Обрабатываем каждое обычное изображение
    for image_idx, image_path in enumerate(ordinary_images_sorted):
        print(f"Обработка обычного изображения ({image_idx+1}/{len(ordinary_images_sorted)}): {image_path}")
        try:
            # Автоматическая обрезка
            cropped_image = auto_crop_image(str(image_path))
            if cropped_image is None:
                raise ValueError("Автоматическая обрезка не удалась.")
            print("Обрезка успешно выполнена.")
            
            # Показываем результат автоматической обрезки
            cv2.imshow("AUTO CROP RESULT", cropped_image)
            cv2.waitKey(0)
            cv2.destroyAllWindows()
            
            # Предлагаем пользователю выбрать вариант обрезки
            user_choice = input("Необходимо:\n1. Дополнительно обрезать имеющееся изображение;\n2. Самостоятельно обрезать оригинальное изображение.\nЛюбая клавиша для продолжения.\n")
            
            if user_choice == '1':
                print("Запуск интерактивного режима...")
                interactive_cropped_image = interactive_crop(cropped_image)
                if interactive_cropped_image is not None:
                    cropped_image = interactive_cropped_image
                print("Интерактивная обрезка завершена.")
            elif user_choice == '2':
                print("Загрузка оригинального изображения...")
                original_image_np = load_image_with_pillow(str(image_path))
                original_image = cv2.cvtColor(original_image_np, cv2.COLOR_RGB2BGR)
                
                print("Применение обработки к оригинальному изображению...")
                original_image = increase_contrast(original_image)
                original_image = resize_image(original_image)
                
                print("Запуск интерактивного режима для оригинального изображения...")
                interactive_cropped_image = interactive_crop(original_image)
                if interactive_cropped_image is not None:
                    cropped_image = interactive_cropped_image
                else:
                    cropped_image = original_image
                print("Интерактивная обрезка оригинального изображения завершена.")
            else:
                print("Продолжение без изменений.")
            
            # Конвертируем обрезанное изображение в формат Pillow для сохранения
            print("Конвертация изображения в формат Pillow...")
            cropped_image_pil = Image.fromarray(cv2.cvtColor(cropped_image, cv2.COLOR_BGR2RGB))
            
            # Сохраняем временное изображение после обрезки
            temp_cropped_path = images_folder / f"temp_cropped_image_{image_idx}.jpg"
            cropped_image_pil.save(temp_cropped_path, format="JPEG", quality=100)
            if not temp_cropped_path.exists():
                raise FileNotFoundError(f"Не удалось сохранить временный файл: {temp_cropped_path}")
            print(f"Временное изображение после обрезки сохранено: {temp_cropped_path}")
            
            # Применяем настройку яркости и контрастности
            print("Применение настроек яркости и контрастности...")
            adjusted_image = auto_adjust_brightness_contrast(temp_cropped_path, brightness_factor, contrast_factor)
            
            # Сохраняем временное изображение в формате JPEG
            temp_image_path = images_folder / f"temp_adjusted_image_{image_idx}.jpg"
            adjusted_image.save(temp_image_path, format="JPEG", quality=100)
            if not temp_image_path.exists():
                raise FileNotFoundError(f"Не удалось сохранить временный файл: {temp_image_path}")
            print(f"Временное изображение после настройки сохранено: {temp_image_path}")
            
            # Вставка в документ
            print("Вставка изображения в документ...")
            for table, row, cell, paragraph in tables_with_marker:
                if image_idx > 0:  # Для второго и последующих изображений добавляем новую строку
                    new_row = table.add_row()
                    # Копируем форматирование из исходной строки
                    for i, src_cell in enumerate(row.cells):
                        if i < len(new_row.cells):
                            # Копируем ВСЕ свойства ячейки
                            new_cell = new_row.cells[i]
                            src_tcPr = src_cell._tc.get_or_add_tcPr()
                            new_tcPr = new_cell._tc.get_or_add_tcPr()
                            
                            # Удаляем старые свойства
                            for child in list(new_tcPr):
                                new_tcPr.remove(child)
                            
                            # Копируем все свойства из исходной ячейки
                            for child in src_tcPr:
                                new_tcPr.append(deepcopy(child))
                            
                            # Добавляем отступы заново
                            tcMar = OxmlElement('w:tcMar')
                            for margin_name in ['top', 'left', 'bottom', 'right']:
                                margin = OxmlElement(f'w:{margin_name}')
                                margin.set(qn('w:w'), '20')
                                margin.set(qn('w:type'), 'dxa')
                                tcMar.append(margin)
                            new_tcPr.append(tcMar)
                    
                    # Переносим метку в новую строку
                    new_row.cells[0].text = "[IMAGE_OTHER]"
                    current_row = new_row
                    current_cell = new_row.cells[0]
                    current_paragraph = current_cell.paragraphs[0]
                else:
                    current_row = row
                    current_cell = cell
                    current_paragraph = paragraph
                
                current_paragraph.text = ""
                run = current_paragraph.add_run()
                
                # Добавляем изображение с нужной высотой
                picture = run.add_picture(str(temp_image_path), height=image_height)
                
                # Устанавливаем ширину ячейки
                image_width = picture.width.pt
                padding = 10
                current_cell.width = Pt(image_width + padding)
                
                # Добавляем черную границу изображения
                props = picture._inline.graphic.graphicData.pic.spPr
                ln = OxmlElement('a:ln')
                ln.set('w', str(int(0.5 * 12700)))  # толщина 0.5 пт
                solidFill = OxmlElement('a:solidFill')
                srgbClr = OxmlElement('a:srgbClr')
                srgbClr.set('val', '000000')
                solidFill.append(srgbClr)
                ln.append(solidFill)
                props.append(ln)

                # Устанавливаем границы и отступы для ячейки
                tc = current_cell._tc
                tcPr = tc.get_or_add_tcPr()

                # Удаляем все существующие границы и отступы
                for element in tcPr.xpath('.//*[self::w:tcBorders or self::w:tcMar]'):
                    tcPr.remove(element)

                # Устанавливаем отступы
                tcMar = OxmlElement('w:tcMar')
                for margin_name in ['top', 'left', 'bottom', 'right']:
                    margin = OxmlElement(f'w:{margin_name}')
                    margin.set(qn('w:w'), '20')
                    margin.set(qn('w:type'), 'dxa')
                    tcMar.append(margin)
                tcPr.append(tcMar)

                # Финализируем настройки
                tc.set(qn('w:tcFitText'), 'false')

                # Добавляем подпись "Фото X" в соседнюю ячейку
                if len(current_row.cells) > 1:  # Проверяем, что есть соседняя ячейка
                    neighbor_cell = current_row.cells[1]
                    # Сбрасываем все возможные отступы ячейки
                    neighbor_cell.paragraphs[0].paragraph_format.space_before = Pt(0)
                    neighbor_cell.paragraphs[0].paragraph_format.space_after = Pt(0)
                    neighbor_cell.text = f"\n\n\n\n\n\n\n\n\nФото {photo_counter}"
                    # Форматирование текста подписи
                    for paragraph in neighbor_cell.paragraphs:
                        paragraph.alignment = 0  # Выравнивание по левому краю
                        # Убираем все отступы
                        paragraph.paragraph_format.left_indent = Pt(0)
                        paragraph.paragraph_format.first_line_indent = Pt(0)
                        # Форматирование шрифта
                        for run in paragraph.runs:
                            run.font.name = 'Times New Roman'
                            run.font.size = Pt(12)
                    photo_counter += 1

                print(f"Изображение {image_path.name} вставлено в строку {image_idx+1}")
            
            # Удаляем временные файлы
            print("Удаление временных файлов...")
            if temp_image_path.exists():
                temp_image_path.unlink()
            if temp_cropped_path.exists():
                temp_cropped_path.unlink()
            print("Временные файлы удалены.")
        
        except Exception as e:
            print(f"Ошибка при обработке изображения {image_path}: {e}")
            continue
    
    # Вставляем специальные изображения (1, 1_2, 1_3) без регулировки яркости и контрастности
    special_images = {
        "1": "[IMAGE_1]",
        "1_2": "[IMAGE_1_2]",
        "1_3": "[IMAGE_1_3]"
    }

    # Определяем, какие условия выполняются
    card_value = card_data.get("1", "").strip()
    logger.info(f"Анализ значения поля '1': '{card_value}' для определения настраиваемой ширины изображений 1_2 и 1_3.")
    
    # Список префиксов, при совпадении с началом card_value включаем настраиваемую ширину
    custom_width_prefixes = ("СВО_Молов_образец_", "СВО_Молов БВП_образец_", "СВО_Ростов_образец_", "ЭКЦ_")
    use_custom_width_for_1_2_3 = (
        card_value.startswith(custom_width_prefixes) or card_value == "ЭКЦ"
    )
    
    logger.info(
        f"Результат проверки настраиваемой ширины для '{card_value}': "
        f"{use_custom_width_for_1_2_3}. "
        f"Префиксы для проверки: {custom_width_prefixes}."
    )

    for image_name, placeholder in special_images.items():
        image_path = images_folder / f"{image_name}.jpg"
        if not image_path.exists():  # Если файл не найден, пробуем другие расширения
            for ext in [".jpg", ".jpeg", ".png"]:
                image_path = images_folder / f"{image_name}{ext}"
                if image_path.exists():
                    break
        if image_path.exists():
            logger.info(f"Обработка специального изображения: {image_path}")
            try:
                # Загружаем изображение без регулировки яркости и контрастности
                image = Image.open(image_path)
                # Сохраняем временное изображение в формате JPEG
                temp_image_path = images_folder / "temp_special_image.jpg"
                image.save(temp_image_path, format="JPEG", quality=100)
                if not temp_image_path.exists():
                    raise FileNotFoundError(f"Не удалось сохранить временный файл: {temp_image_path}")
                logger.info(f"Временное изображение сохранено: {temp_image_path}")

                # Устанавливаем ширину изображения в зависимости от условия
                if image_name in ["1_2", "1_3"] and use_custom_width_for_1_2_3:
                    image_width = Inches(3.3)  # значение ширины для 1_2 и 1_3
                    logger.info(f"Установлена настраиваемая ширина для изображения {image_name}: {image_width}")
                else:
                    image_width = Inches(6.4)  # стандартная ширина
                    logger.info(f"Установлена стандартная ширина для изображения {image_name}: {image_width}")

                # Вставка в документ
                for p in doc.paragraphs:
                    if placeholder in p.text:
                        p.text = p.text.replace(placeholder, "")
                        run = p.add_run()
                        run.add_picture(str(temp_image_path), width=image_width)
                        logger.info(f"Изображение вставлено в метку {placeholder}")
                        break

                # Удаляем временное изображение после вставки
                if temp_image_path.exists():
                    temp_image_path.unlink()
                    logger.info(f"Временный файл {temp_image_path} удалён.")
            except Exception as e:
                logger.error(f"Ошибка при обработке изображения {image_path}: {e}")
    
    # Динамический поиск дополнительных картинок ПАН (1_4, 1_5 и т.д.)
    logger.info("Поиск динамических изображений ПАН (1_4, 1_5...)...")
    dynamic_placeholders = set()
    for p in doc.paragraphs:
        matches = re.findall(r'\[IMAGE_1_(\d+)\]', p.text)
        for m in matches:
            num = int(m)
            if num >= 4:
                dynamic_placeholders.add(f"1_{num}")
                
    for image_name in sorted(list(dynamic_placeholders)):
        placeholder = f"[IMAGE_1_{image_name.split('_')[1]}]"
        image_path = images_folder / f"{image_name}.jpg"
        if not image_path.exists():
            for ext in [".jpg", ".jpeg", ".png"]:
                image_path = images_folder / f"{image_name}{ext}"
                if image_path.exists():
                    break
                    
        if image_path.exists():
            logger.info(f"Обработка динамического изображения: {image_path}")
            try:
                image = Image.open(image_path)
                temp_image_path = images_folder / f"temp_dynamic_{image_name}.jpg"
                image.save(temp_image_path, format="JPEG", quality=100)
                
                # Используем ту же логику ширины, что и для 1_2 / 1_3
                if use_custom_width_for_1_2_3:
                    image_width = Inches(3.3)
                else:
                    image_width = Inches(6.4)
                    
                for p in doc.paragraphs:
                    if placeholder in p.text:
                        p.text = p.text.replace(placeholder, "")
                        run = p.add_run()
                        run.add_picture(str(temp_image_path), width=image_width)
                        logger.info(f"Динамическое изображение вставлено в метку {placeholder}")
                        break
                        
                if temp_image_path.exists():
                    temp_image_path.unlink()
            except Exception as e:
                logger.error(f"Ошибка при обработке динамического изображения {image_path}: {e}")
        else:
            logger.warning(f"Файл для {image_name} не найден в папке. Метка {placeholder} останется пустой.")

    doc.save(doc_path)
    print("Все изображения обработаны и вставлены.")

# Функция для поиска файла .txt в целевой папке
def find_txt_file(folder_path):
    """
    Ищет файл(ы) .txt в указанной папке, которые заканчиваются на "Genotypes Table".
    Если найдено несколько файлов, соответствующих этому правилу, позволяет пользователю выбрать нужный файл.
    Если найден только один файл, соответствующий правилу, он автоматически выбирается.
    Если файлов, соответствующих правилу, нет, возвращает None.
    """
    # Ищем все файлы с расширением .txt в целевой папке
    txt_files = list(Path(folder_path).glob("*.txt"))
    logger.info(f"Найдены файлы .txt: {[f.name for f in txt_files]}")

    # Фильтруем файлы, оставляя только те, которые заканчиваются на "Genotypes Table"
    filtered_files = [file for file in txt_files if file.name.endswith("Genotypes Table.txt")]
    logger.info(f"Отфильтрованные файлы: {[f.name for f in filtered_files]}")

    # Если файлов, соответствующих правилу, нет
    if not filtered_files:
        if txt_files:
            logger.warning(f"Файлы .txt найдены, но ни один не соответствует правилу (не заканчивается на 'Genotypes Table').")
        else:
            logger.warning("Файлы .txt не найдены.")
        return None  # Возвращаем None, если файл не найден или не соответствует правилу
    
    # Если найден только один файл, соответствующий правилу
    if len(filtered_files) == 1:
        logger.info(f"Найден файл, соответствующий правилу: {filtered_files[0].name}")
        return filtered_files[0]  # Возвращаем единственный найденный файл
    
    # Если найдено несколько файлов .txt, соответствующих правилу, спрашиваем пользователя
    logger.info(f"Найдено несколько файлов .txt, заканчивающихся на 'Genotypes Table': {[f.name for f in filtered_files]}")
    print(f"В папке {folder_path} найдено несколько файлов .txt, заканчивающихся на 'Genotypes Table':")
    for i, txt_file in enumerate(filtered_files, start=1):
        print(f"{i}. {txt_file.name}")
    
    while True:
        try:
            choice = int(input("Выберите номер файла для загрузки: "))
            if 1 <= choice <= len(filtered_files):
                selected_file = filtered_files[choice - 1]
                logger.info(f"Выбран файл: {selected_file.name}")
                return selected_file
            else:
                print("Некорректный выбор. Пожалуйста, введите номер из списка.")
        except ValueError:
            print("Пожалуйста, введите числовое значение.")

# Функция для чтения и фильтрации данных из .txt файла
def read_and_filter_txt_data(file_path, require_suffix=True):
    """
    Читает данные из txt файла и возвращает все профили.
    :param require_suffix: если True, оставляем только профили, оканчивающиеся на -26.
                           если False, возвращаем все профили, исключая служебные строки.
    """
    logger.info(f"Начало чтения и фильтрации данных из файла: {file_path} (require_suffix={require_suffix})")
    # Значения, которые нужно исключить (в разных вариантах регистра)
    exclude_values = {"AL", "Al", "al", "K+", "К+", "K-", "К-", "KV", "kv", "KF", "kf"}
    
    with open(file_path, 'r', encoding='utf-8') as file:
        lines = file.readlines()
    
    # Убираем заголовок и разбиваем строки на столбцы
    data = [line.strip().split('\t') for line in lines[1:]]
    
    # Фильтруем данные
    filtered_data = []
    skipped_rows = 0
    
    for line in data:
        if not line or not line[0].strip():
            skipped_rows += 1
            continue
        
        profile_name = line[0].strip()
        
        # Пропускаем служебные строки (начинающиеся с #)
        if profile_name.startswith('#'):
            logger.debug(f"Пропущена служебная строка: {profile_name}")
            skipped_rows += 1
            continue
        
        if require_suffix:
            # Старая логика: только профили, оканчивающиеся на -26
            if profile_name.endswith("-26"):
                filtered_data.append(line)
            else:
                skipped_rows += 1
        else:
            # Новая логика: исключаем строки, где имя профиля входит в exclude_values
            if profile_name.upper() in exclude_values:
                logger.debug(f"Пропущена строка с именем профиля из списка исключений: {profile_name}")
                skipped_rows += 1
                continue
            # Также исключаем строки, где имя профиля не содержит цифр (эвристика)
            # Это помогает отсеять случайные служебные значения
            if not any(char.isdigit() for char in profile_name):
                logger.debug(f"Пропущена строка с именем профиля без цифр: {profile_name}")
                skipped_rows += 1
                continue
            filtered_data.append(line)
    
    if skipped_rows > 0:
        logger.info(f"Пропущено строк: {skipped_rows} (служебные или не соответствующие критериям)")
    
    if not filtered_data:
        if require_suffix:
            error_msg = "В файле не найдено ни одного профиля, оканчивающегося на '-26'."
        else:
            error_msg = "В файле не найдено ни одного профиля (после исключения служебных строк)."
        logger.error(error_msg)
        raise ValueError(error_msg)
    
    # Группируем данные по имени профиля
    profiles = {}
    for line in filtered_data:
        sample_name = line[0].strip()
        if sample_name not in profiles:
            profiles[sample_name] = []
        profiles[sample_name].append(line)
    
    # Проверяем на дублирующиеся локусы в каждом профиле
    for sample_name, profile_data in profiles.items():
        sample_loci = set()
        for line in profile_data:
            marker = line[3].strip()
            if marker in sample_loci:
                error_msg = f"Для объекта {sample_name} найдены дублирующиеся локусы. Проверьте данные. Marker: {marker}"
                logger.error(error_msg)
                raise ValueError(error_msg)
            sample_loci.add(marker)
    
    logger.info(f"Найдены профили: {list(profiles.keys())}")
    
    # Проверка на многоаллельные профили
    logger.info("Проверка на многоаллельные профили...")
    multi_allele_detected = False
    multi_allele_info = {}
    
    for profile_name, profile_data in profiles.items():
        logger.info(f"Проверка профиля: {profile_name}")
        
        for line in profile_data:
            marker = line[3].strip()
            alleles = []
            for i in range(5, 9):
                if i < len(line) and line[i] and line[i].strip():
                    allele_value = line[i].strip()
                    if allele_value not in exclude_values:
                        alleles.append(allele_value)
            
            if len(alleles) > 2:
                logger.warning(f"Обнаружен многоаллельный профиль в локусе {marker}: {len(alleles)} аллели")
                multi_allele_detected = True
                if profile_name not in multi_allele_info:
                    multi_allele_info[profile_name] = []
                if len(alleles) == 3:
                    multi_allele_info[profile_name].append((marker, 3, "Обнаружена 3 аллель"))
                elif len(alleles) == 4:
                    multi_allele_info[profile_name].append((marker, 4, "Обнаружены 3 и 4 аллели"))
    
    # Если найдены многоаллельные профили, запрашиваем подтверждение
    if multi_allele_detected:
        logger.info("Обнаружены многоаллельные профили, запрашиваем подтверждение пользователя")
        
        # Формируем сообщение для пользователя
        message = "ВНИМАНИЕ! Обнаружены многоаллельные профили:\n\n"
        
        for profile_name, loci_info in multi_allele_info.items():
            message += f"Профиль: {profile_name}\n"
            for marker, count, description in loci_info:
                message += f"  - {description} в локусе {marker}!\n"
            message += "\n"
        
        message += "Проверьте фореграмму и подтвердите наличие трёхаллельного профиля в указанном локусе.\n"
        message += "В случае ошибки переанализируйте фореграмму и поместите в папку данных новый файл txt.\n\n"
        message += "д - подтверждение профиля, н - завершение работы программы"
        
        # Выводим сообщение в консоль
        print("\n" + "="*80)
        print(message)
        print("="*80)
        
        # Запрашиваем подтверждение
        while True:
            try:
                response = input("\nВведите ваш выбор (д/н): ").strip().lower()
                if response in ['д', 'да', 'y', 'yes']:
                    logger.info("Пользователь подтвердил наличие многоаллельных профилей")
                    print("Подтверждение получено. Продолжаем работу...")
                    break
                elif response in ['н', 'нет', 'n', 'no']:
                    logger.info("Пользователь отказался от подтверждения многоаллельных профилей")
                    print("Завершение работы программы по требованию пользователя...")
                    import sys
                    sys.exit(0)
                else:
                    print("Некорректный ввод. Пожалуйста, введите 'д' для подтверждения или 'н' для завершения.")
            except KeyboardInterrupt:
                logger.warning("Пользователь прервал ввод с помощью Ctrl+C")
                print("\nПрервано пользователем. Завершение работы...")
                import sys
                sys.exit(0)
    
    logger.info("Проверка на многоаллельные профили завершена.")
    
    return profiles

def count_profiles_in_txt(file_path):
    """
    Подсчитывает количество уникальных имён профилей в .txt файле.
    Игнорирует служебные строки (заголовок, строки с #HEADER:, строки с именами из exclude_values).
    Возвращает количество профилей.
    """
    logger.info(f"Подсчёт количества профилей в файле: {file_path}")
    exclude_values = {"AL", "Al", "al", "K+", "К+", "K-", "К-", "KV", "kv", "KF", "kf"}
    
    try:
        with open(file_path, 'r', encoding='utf-8') as file:
            lines = file.readlines()
        
        if len(lines) < 2:
            logger.warning(f"Файл содержит менее 2 строк, профилей не найдено")
            return 0
        
        profile_names = set()
        # Пропускаем первую строку (заголовок)
        for line in lines[1:]:
            line = line.strip()
            if not line:
                continue
            # Пропускаем служебные строки с #HEADER:
            if line.startswith("#HEADER:"):
                continue
            parts = line.split('\t')
            if parts and parts[0].strip():
                profile_name = parts[0].strip()
                # Исключаем строки, где имя профиля входит в exclude_values
                if profile_name.upper() in exclude_values:
                    logger.debug(f"Исключён служебный профиль: {profile_name}")
                    continue
                # Исключаем строки, начинающиеся с #
                if profile_name.startswith('#'):
                    continue
                # Дополнительная эвристика: имя профиля должно содержать хотя бы одну цифру
                if not any(char.isdigit() for char in profile_name):
                    logger.debug(f"Исключён профиль без цифр: {profile_name}")
                    continue
                profile_names.add(profile_name)
        
        logger.info(f"Найдено уникальных профилей: {len(profile_names)} -> {profile_names}")
        return len(profile_names)
    except Exception as e:
        logger.error(f"Ошибка при подсчёте профилей в файле {file_path}: {e}")
        raise

def check_y_rs_loci(txt_data):
    """
    Проверяет, пусты ли значения аллелей всех локусов, которые в названии содержат "Y", "y", "RS" или "rs".
    Если хотя бы один локус содержит значение, отличное от пробела или нуля, возвращает False.
    """
    logger.info("Начало проверки наличия значений в локусах, содержащих Y/RS")
    # Локусы, которые нужно проверить (все в верхнем регистре для регистронезависимости)
    loci_to_check_upper = ["Y", "RS"]
    
    for line in txt_data:
        marker = line[3].strip().upper()  # Приводим к верхнему регистру
        # Проверяем, содержит ли название локуса "Y" или "RS" (в верхнем регистре)
        if any(substring in marker for substring in loci_to_check_upper):
            logger.debug(f"Найден локус с Y/RS: {marker}")
            # Проверяем значения аллелей (столбцы 5 и 6, индексы 5 и 6)
            allele1 = line[5].strip() if len(line) > 5 and line[5] else ""
            allele2 = line[6].strip() if len(line) > 6 and line[6] else ""
            # Если хотя бы одно значение не пустое, возвращаем False
            if allele1 or allele2:
                logger.warning(f"Обнаружено значение в локусе {marker}: аллели '{allele1}', '{allele2}'")
                return False
            else:
                logger.debug(f"Локус {marker} не содержит значений аллелей")
    logger.info("Проверка завершена: все Y/RS локусы пусты")
    return True

def check_ol_alleles_in_profiles(profiles_data):
    """
    Проверяет, содержат ли профили, которые будут вставлены в таблицу,
    аллели со значением 'OL' (off ladder, регистронезависимо).
    При обнаружении выдаёт предупреждение, ожидает нажатия Enter
    и завершает программу с кодом 1.
    """
    logger.info("Проверка профилей на наличие аллелей OL...")
    ol_found = False
    error_details = []

    for profile_name, lines in profiles_data.items():
        for line in lines:
            # marker находится в столбце с индексом 3
            marker = line[3].strip()
            # аллели могут быть в столбцах 5,6,7,8 (индексы 5,6,7,8)
            for idx in range(5, min(len(line), 9)):
                allele = line[idx].strip()
                if allele.upper() == "OL":
                    ol_found = True
                    # для пользователя: номер аллели (1-4)
                    allele_num = idx - 4
                    error_details.append(
                        f"Профиль: {profile_name}, локус: {marker}, "
                        f"аллель {allele_num}: '{allele}'"
                    )
                    logger.warning(
                        f"Обнаружен OL аллель в профиле {profile_name}, "
                        f"локусе {marker}: значение '{allele}'"
                    )

    if ol_found:
        msg = (
            "ВНИМАНИЕ! Обнаружены аллели со значением 'OL'"
            "в профилях, предназначенных для вставки в таблицу:\n"
            + "\n".join(error_details)
            + "\n\nНеобходимо исправить данные в файле .txt и перезапустить программу."
            + "\n\nНажмите Enter для завершения работы..."
        )
        logger.error("Обнаружены аллели OL. Завершение программы.")
        print("\n" + "=" * 80)
        print(msg)
        print("=" * 80)
        input()
        import sys
        sys.exit(1)
    else:
        logger.info("Аллелей OL не обнаружено.")

# Функция для вставки данных в таблицу Word
def insert_table_data(doc_path, profiles_data, card_data, custom_headers=None, skip_comment=False,
                      reference_profile_data=None, reference_profile_name=None, is_tobolsk=False):
    """
    Вставляет данные из .txt файла в таблицу Word с несколькими столбцами.
    Добавляет комментарий после таблицы в зависимости от условий.
    
    :param custom_headers: Словарь с кастомными заголовками для профилей {profile_name: header_text}
    :param skip_comment: Если True, комментарий после таблицы не добавляется
    :param reference_profile_data: Данные образца (словарь локус->множество аллелей) для сравнения и выделения
    :param reference_profile_name: Имя профиля-образца (для него выделение не применяется)
    :param is_tobolsk: Флаг, что обрабатывается тобольский образец (включает подсчёт несовпадений)
    """
    logger.info("Начало вставки данных в таблицу с несколькими столбцами")
    
    # Определяем множество исключаемых значений (используется в цикле)
    exclude_values = {"AL", "Al", "al", "K+", "К+", "K-", "К-", "KV", "kv", "KF", "kf"}
    
    doc = Document(doc_path)
    
    # Получаем список всех профилей
    profile_names = list(profiles_data.keys())
    logger.info(f"Обрабатываем профили: {profile_names}")
    
    # Блок проверки наличия y и rs при Amel X (оставляем без изменений)
    logger.info("Проверка наличия значений в Y/RS локусах для всех профилей, оканчивающихся на -26")

    # Отбираем профили, которые нужно проверить (только с суффиксом -26)
    profiles_to_check = {name: data for name, data in profiles_data.items() if name.endswith("-26")}

    if not profiles_to_check:
        logger.warning("Не найдено ни одного профиля, оканчивающегося на -26. Проверка Y/RS пропущена.")
    else:
        logger.info(f"Будут проверены профили: {list(profiles_to_check.keys())}")
        
        for profile_name, profile_data in profiles_to_check.items():
            logger.debug(f"Проверка профиля: {profile_name}")

            # Поиск Amel/Amelogenin в профиле (регистронезависимо)
            amel_line = None
            for line in profile_data:
                marker = line[3].strip().upper()
                if marker == "AMEL" or marker == "AMELOGENIN" or marker.startswith("AMEL"):
                    amel_line = line
                    logger.debug(f"Найден локус Amel/Amelogenin: {marker}")
                    break

            if amel_line is not None:
                alleles = []
                for idx in [5, 6]:
                    if idx < len(amel_line) and amel_line[idx] and amel_line[idx].strip():
                        alleles.append(amel_line[idx].strip().upper())
                logger.debug(f"Профиль {profile_name}, Amel аллели: {alleles}")

                has_y = "Y" in alleles

                if has_y:
                    logger.debug(f"Профиль {profile_name} имеет Amel Y, пропускаем проверку Y/RS локусов")
                    continue
                else:
                    logger.debug(f"Профиль {profile_name} имеет Amel без Y, проверяем Y/RS локусы")
                    if not check_y_rs_loci(profile_data):
                        error_msg = (f"Обнаружены значения в локусах, содержащих 'Y', 'y', 'RS' или 'rs', "
                                    f"при этом в Amel отсутствует Y (профиль {profile_name}). Выполнение остановлено.")
                        logger.error(error_msg)
                        show_error_message(error_msg)
                        sys.exit(1)
                    else:
                        logger.debug(f"Профиль {profile_name} прошел проверку: значения в Y/RS локусах отсутствуют")
            else:
                logger.warning(f"В профиле {profile_name} не найден локус Amel. Проверка Y/RS локусов будет выполнена, так как Y в Amel не обнаружен.")
                if not check_y_rs_loci(profile_data):
                    error_msg = (f"Обнаружены значения в локусах, содержащих 'Y', 'y', 'RS' или 'rs', "
                                f"при этом локус Amel отсутствует в профиле {profile_name}. Выполнение остановлено.")
                    logger.error(error_msg)
                    show_error_message(error_msg)
                    sys.exit(1)

    logger.info("Проверка Y/RS локусов для всех профилей с -26 завершена успешно")

    # Если передан словарь кастомных заголовков, используем его
    if custom_headers is None:
        custom_headers = {}
    
    logger.info(f"Кастомные заголовки: {custom_headers}")
    
    # Получаем все уникальные локусы из всех профилей
    all_loci = set()
    for profile_data in profiles_data.values():
        for line in profile_data:
            all_loci.add(line[3].strip())
    
    # Сортируем локусы
    sorted_loci = sort_data_by_locus([["", "", "", locus] for locus in all_loci])
    sorted_loci = [line[3] for line in sorted_loci]
    logger.info(f"Уникальные отсортированные локусы: {sorted_loci}")
    
    # Определяем ширину таблицы и тип заголовка на основе значения из карты
    card_value = card_data.get("1", "").strip()
    logger.info(f"Определение параметров таблицы для значения карты: '{card_value}'")
    
    # Установка ширины таблицы
    if card_value in [
        "СВО_Молов_образец_родственники", 
        "СВО_Ростов_образец_родственники", 
        "СВО_Молов_образец_прямая идентификация",
        "СВО_Молов БВП_образец_прямая идентификация",
        "СВО_Ростов_образец_прямая идентификация", 
        "ЭКЦ", 
        "СВО_кость"
    ]:
        table_width = 5.0
        logger.info("Установлена ширина таблицы: 5.0 дюймов")
    elif card_value.startswith("ЭКЦ_образец_"):
        table_width = 6.0
        logger.info("Установлена ширина таблицы: 6.0 дюймов для ЭКЦ образца")
    elif card_value.startswith("Тобольск_образец_"):
        table_width = 6.0
        logger.info("Установлена ширина таблицы: 6.0 дюймов для тобольского образца")
    else:
        table_width = 7.0
        logger.info("Установлена ширина таблицы: 7.0 дюймов (по умолчанию)")
    
    # Определение формата заголовка
    if card_value in [
        "СВО_Молов_образец_родственники",
        "СВО_Ростов_образец_родственники",
        "СВО_Молов_образец_прямая идентификация",
        "СВО_Молов БВП_образец_прямая идентификация",
        "СВО_Ростов_образец_прямая идентификация",
        "ЭКЦ"
    ]:
        header_template = "Объект {НОМ}\n(образец {ТО} {ФР} {ИР} {ОР}, {ДР} г.р.)"
        logger.info("Используется стандартный формат заголовка")
    elif card_value == "СВО_кость":
        header_template = "Объект {ОБ}\nобразец биологического материала\nот трупа № {ТР}"
        logger.info("Используется специальный формат заголовка для кости")
    else:
        header_template = "Объект {НОМ}\n(образец {ТО} {ФР} {ИР} {ОР}, {ДР} г.р.)"
        logger.info("Используется стандартный формат заголовка (по умолчанию)")
    
    # Ищем таблицу с меткой [TABLE_DATA]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if "[TABLE_DATA]" in cell.text:
                    logger.info("Найдена таблица с меткой [TABLE_DATA]")
                    
                    # Очищаем таблицу, оставляя только строку с меткой
                    while len(table.rows) > 1:
                        table._tbl.remove(table.rows[1]._tr)
                    
                    # Очищаем ячейку с меткой
                    cell.text = ""
                    
                    # Добавляем заголовки для каждого профиля
                    header_row = table.rows[0]
                    
                    # Убедимся, что в заголовочной строке достаточно ячеек
                    needed_columns = len(profile_names) + 1
                    current_columns = len(header_row.cells)
                    
                    # Добавляем недостающие ячейки в заголовочную строку
                    if current_columns < needed_columns:
                        for _ in range(needed_columns - current_columns):
                            # Создаем новую ячейку с помощью низкоуровневого метода
                            tc = OxmlElement('w:tc')
                            header_row._tr.append(tc)
                    
                    # Заполняем заголовки
                    header_cells = header_row.cells
                    for i, profile_name in enumerate(profile_names):
                        if i < len(header_cells) - 1:
                            # Проверяем, есть ли кастомный заголовок для этого профиля
                            if profile_name in custom_headers:
                                header_text = custom_headers[profile_name]
                                logger.info(f"Используется кастомный заголовок для профиля {profile_name}: '{header_text}'")
                            else:
                                # Стандартная логика формирования заголовка
                                clean_profile_name = profile_name
                                if clean_profile_name.endswith("-26"):
                                    clean_profile_name = clean_profile_name[:-3]  # Удаляем последние 3 символа ("-26")
                                logger.info(f"Преобразование профиля: '{profile_name}' -> '{clean_profile_name}'")

                                # Проверка, что номер не стал пустым после удаления суффикса
                                if not clean_profile_name.strip():
                                    logger.warning(f"Пустой номер после удаления суффикса для профиля: {profile_name}")
                                    clean_profile_name = profile_name  # Вернуть оригинальное имя как fallback

                                # Формируем заголовок в зависимости от типа экспертизы
                                logger.info(f"Формирование заголовка для профиля: {clean_profile_name}")

                                if card_value == "СВО_кость":
                                    header_text = header_template
                                    data_for_header = {
                                        'ОБ': card_data.get('ОБ'),
                                        'ТР': card_data.get('ТР')
                                    }
                                else:
                                    header_text = header_template
                                    data_for_header = {
                                        'НОМ': clean_profile_name,
                                        'ТО': card_data.get('ТО'),
                                        'ФР': card_data.get('ФР'),
                                        'ИР': card_data.get('ИР'), 
                                        'ОР': card_data.get('ОР'),
                                        'ДР': card_data.get('ДР')
                                    }

                                # Заменяем метки в заголовке
                                for key, value in data_for_header.items():
                                    if value is not None and value != '':
                                        header_text = header_text.replace(f"{{{key}}}", str(value))
                                        logger.debug(f"Замена {{{key}}} -> '{value}'")
                                    else:
                                        header_text = header_text.replace(f"{{{key}}}", "")
                                        logger.debug(f"Удаление метки {{{key}}} (значение: {value})")

                                # Очищаем от лишних пробелов
                                header_text = clean_text_from_none_and_spaces(header_text)
                                logger.info(f"Сформирован заголовок: '{header_text}'")
                            
                            # Очищаем ячейку перед добавлением форматированного текста
                            header_cells[i+1].text = ""
                            
                            # Разделяем текст на части для разного форматирования
                            lines = header_text.split('\n')
                            
                            # Добавляем первую строку жирным шрифтом
                            if lines:
                                p = header_cells[i+1].paragraphs[0]
                                run = p.add_run(lines[0])
                                run.bold = True
                                run.font.name = 'Times New Roman'
                                run.font.size = Pt(12)
                                
                                # Добавляем остальные строки обычным шрифтом
                                if len(lines) > 1:
                                    for line in lines[1:]:
                                        run = p.add_run('\n' + line)
                                        run.bold = True
                                        run.font.name = 'Times New Roman'
                                        run.font.size = Pt(12)
                    
                    # Инициализация счётчика несовпадений (только для тобольского режима)
                    if is_tobolsk and reference_profile_data is not None:
                        mismatch_counts = {}
                        logger.info("Включён подсчёт несовпадающих локусов для тобольского образца")
                    
                    # Добавляем данные для каждого локуса
                    for locus in sorted_loci:
                        new_row = table.add_row()
                        
                        # Убедимся, что в новой строке достаточно ячеек
                        if len(new_row.cells) < needed_columns:
                            for _ in range(needed_columns - len(new_row.cells)):
                                tc = OxmlElement('w:tc')
                                new_row._tr.append(tc)
                        
                        new_row.cells[0].text = locus
                        
                        # Заполняем ячейки для каждого профиля
                        for i, profile_name in enumerate(profile_names):
                            profile_data = profiles_data[profile_name]
                            alleles = "-"
                            
                            # Ищем локус в профиле
                            for line in profile_data:
                                if line[3].strip() == locus:
                                    # Читаем все аллели (до 4)
                                    alleles_list = []
                                    for col_index in range(5, 9):  # Столбцы 5-8: Allele 1-4
                                        if col_index < len(line) and line[col_index] and line[col_index].strip():
                                            allele_value = line[col_index].strip()
                                            # Исключаем специальные значения
                                            if allele_value not in exclude_values:
                                                alleles_list.append(allele_value)
                                    alleles = ",".join(alleles_list) if alleles_list else "-"
                                    break
                            
                            target_cell = new_row.cells[i+1]
                            
                            # Логика выделения аллелей для тобольского режима
                            if reference_profile_data is not None and profile_name != reference_profile_name:
                                # Получаем аллели образца для этого локуса
                                reference_alleles_set = set()
                                # Ищем локус в данных образца
                                for ref_line in reference_profile_data:
                                    if ref_line[3].strip() == locus:
                                        ref_alleles = []
                                        for col_idx in range(5, 9):
                                            if col_idx < len(ref_line) and ref_line[col_idx] and ref_line[col_idx].strip():
                                                av = ref_line[col_idx].strip()
                                                if av not in exclude_values:
                                                    ref_alleles.append(av)
                                        reference_alleles_set = set(ref_alleles)
                                        break
                                logger.debug(f"Сравнение локуса {locus} для профиля {profile_name}: "
                                             f"аллели рабочего '{alleles}', аллели образца {reference_alleles_set}")
                                
                                # Подсчёт несовпадений для тобольского режима
                                if is_tobolsk and reference_alleles_set and alleles != "-":
                                    current_alleles = [a.strip() for a in alleles.split(',') if a.strip()]
                                    # Проверяем, есть ли хотя бы одна общая аллель
                                    has_common = any(a in reference_alleles_set for a in current_alleles)
                                    if not has_common:
                                        mismatch_counts[profile_name] = mismatch_counts.get(profile_name, 0) + 1
                                        logger.debug(f"Несовпадение по локусу {locus} для профиля {profile_name}: "
                                                     f"аллели {alleles} vs образец {reference_alleles_set}")
                                
                                # Формируем текст с выделением фона
                                highlighted_parts = build_highlighted_alleles_text(alleles, reference_alleles_set)
                                
                                # Очищаем ячейку и вставляем runs с выделением фона
                                target_cell.text = ""
                                paragraph = target_cell.paragraphs[0]
                                for text, highlight in highlighted_parts:
                                    run = paragraph.add_run(text)
                                    if highlight is not None:
                                        run.font.highlight_color = highlight
                                    # Для обычного текста (без выделения) фон не устанавливаем
                                    run.font.name = 'Times New Roman'
                                    run.font.size = Pt(12)
                                logger.debug(f"Ячейка для {profile_name}, локус {locus} заполнена с выделением фона")

                            else:
                                # Обычная вставка (без выделения)
                                target_cell.text = alleles
                                # Для обычной вставки тоже можно задать шрифт, но это не обязательно
                                for paragraph in target_cell.paragraphs:
                                    for run in paragraph.runs:
                                        run.font.name = 'Times New Roman'
                                        run.font.size = Pt(12)
                    
                    # Показ предупреждений для тобольского режима
                    if is_tobolsk and mismatch_counts:
                        logger.info(f"Результаты подсчёта несовпадений: {mismatch_counts}")
                        for profile_name, count in mismatch_counts.items():
                            if 1 <= count <= 3:
                                # Склонение слова «локус»
                                if count % 10 == 1 and count % 100 != 11:
                                    word = "локус"
                                elif count % 10 in [2,3,4] and count % 100 not in [12,13,14]:
                                    word = "локуса"
                                else:
                                    word = "локусов"
                                msg = f"Не совпадает всего {count} {word}! Программа будет продолжена, но проверьте возможность родства."
                                logger.info(f"Тобольский образец: {msg} (профиль {profile_name})")
                                # Показываем графическое окно предупреждения
                                root = tk.Tk()
                                root.withdraw()
                                tk.messagebox.showwarning("Предупреждение", msg)
                                root.destroy()
                    
                    # Убедимся, что границы таблицы видны
                    for row in table.rows:
                        for cell in row.cells:
                            tc = cell._tc
                            tcPr = tc.get_or_add_tcPr()
                            tcBorders = OxmlElement('w:tcBorders')
                            for border_name in ['top', 'left', 'bottom', 'right']:
                                border = OxmlElement(f'w:{border_name}')
                                border.set(qn('w:val'), 'single')
                                border.set(qn('w:sz'), '4')
                                border.set(qn('w:space'), '0')
                                border.set(qn('w:color'), '000000')
                                tcBorders.append(border)
                            tcPr.append(tcBorders)
                    
                    # Устанавливаем общую ширину таблицы
                    logger.info(f"Установка общей ширины таблицы: {table_width} дюймов")
                    tbl = table._tbl
                    tblPr = tbl.tblPr
                    if tblPr is None:
                        tblPr = OxmlElement('w:tblPr')
                        tbl.insert(0, tblPr)
                    
                    # Удаляем существующие настройки ширины таблицы, если они есть
                    existing_tblW = tblPr.find(qn('w:tblW'))
                    if existing_tblW is not None:
                        tblPr.remove(existing_tblW)
                        logger.info("Удалены существующие настройки ширины таблицы")
                    
                    # Устанавливаем новую ширину таблицы
                    tblW = OxmlElement('w:tblW')
                    tblW.set(qn('w:w'), str(int(table_width * 1440)))  # Конвертируем дюймы в twips
                    tblW.set(qn('w:type'), 'dxa')
                    tblPr.append(tblW)
                    logger.info(f"Установлена общая ширина таблицы: {int(table_width * 1440)} twips")
                    
                    # Устанавливаем ширину для каждого столбца
                    # Первый столбец (с локусами) - уже, остальные - одинаковой ширины
                    if table.rows and len(table.rows[0].cells) > 0:
                        # Определяем общее количество столбцов
                        total_columns = len(table.rows[0].cells)
                        logger.info(f"Общее количество столбцов: {total_columns}")
                        
                        # Устанавливаем ширину для каждого столбца
                        for col_idx in range(total_columns):
                            # Для первого столбца устанавливаем меньшую ширину
                            if col_idx == 0:
                                width_value = Inches(1.5)  # Ширина первого столбца
                                logger.info(f"Столбец {col_idx}: ширина = 1.5 дюймов")
                            else:
                                # Для остальных столбцов вычисляем равную ширину
                                width_value = Inches((table_width - 1.5) / (total_columns - 1))
                                logger.info(f"Столбец {col_idx}: ширина = {(table_width - 1.5) / (total_columns - 1):.2f} дюймов")
                            
                            # Устанавливаем ширину для всех ячеек в столбце
                            for row_idx, row in enumerate(table.rows):
                                if col_idx < len(row.cells):
                                    row.cells[col_idx].width = width_value
                                    logger.debug(f"Установлена ширина ячейки [{row_idx}][{col_idx}]: {width_value}")
                    
                    # Выравниваем текст во всех столбцах (кроме первого) по центру
                    for row in table.rows:
                        for i, cell in enumerate(row.cells):
                            if i > 0:  # Все столбцы кроме первого (с локусами)
                                for paragraph in cell.paragraphs:
                                    paragraph.alignment = 1  # Центрирование
                    
                    logger.info("Завершена настройка ширины таблицы и столбцов")
                    
                    # Проверяем наличие прочерков и F во всех ячейках таблицы
                    has_dash = False
                    has_f = False

                    # ДЛЯ ЭКЦ_ОБРАЗЕЦ_: дополнительная проверка по столбцам
                    has_dash_main_profile = False  # Прочерк в основном профиле (проанализированном в экспертизе)
                    has_dash_ekc_sample = False   # Прочерк в столбце ЭКЦ образца

                    for row in table.rows[1:]:  # Пропускаем заголовок
                        for cell_idx, cell in enumerate(row.cells[1:], start=1):  # Пропускаем первый столбец (локусы)
                            cell_text = cell.text.strip()
                            if "-" in cell_text:
                                has_dash = True
                                
                                # Дополнительная проверка для ЭКЦ образцов
                                if card_value.startswith("ЭКЦ_образец_") and len(profile_names) >= 2:
                                    # Определяем, в каком столбце находится прочерк
                                    if cell_idx == 1:  # Первый столбец после локусов - основной профиль
                                        has_dash_main_profile = True
                                        logger.info(f"Найден прочерк в основном профиле (столбец {cell_idx})")
                                    elif cell_idx == 2:  # Второй столбец - ЭКЦ образец
                                        has_dash_ekc_sample = True
                                        logger.info(f"Найден прочерк в ЭКЦ образце (столбец {cell_idx})")
                                
                            if "f" in cell_text:
                                has_f = True

                    logger.info(f"Результат анализа ячеек: has_dash={has_dash}, has_f={has_f}")
                    if card_value.startswith("ЭКЦ_образец_"):
                        logger.info(f"Для ЭКЦ образца: has_dash_main_profile={has_dash_main_profile}, has_dash_ekc_sample={has_dash_ekc_sample}")

                    # Формирование комментария (только если skip_comment == False)
                    if not skip_comment:
                        comment = None

                        # ЛОГИКА ДЛЯ ЭКЦ_ОБРАЗЕЦ_
                        if card_value.startswith("ЭКЦ_образец_"):
                            logger.info("Применение специальной логики комментария для ЭКЦ образца")
                            
                            if has_dash_main_profile and has_dash_ekc_sample:
                                comment = " (-) - отсутствие продуктов амплификации или данные не предоставлены."
                                logger.info("Сформирован комментарий: оба столбца содержат прочерки")
                            elif has_dash_main_profile:
                                comment = " (-) - отсутствие продуктов амплификации или удовлетворительных результатов."
                                logger.info("Сформирован комментарий: только основной профиль содержит прочерки")
                            elif has_dash_ekc_sample:
                                comment = " (-) - данные не предоставлены."
                                logger.info("Сформирован комментарий: только ЭКЦ образец содержит прочерки")
                            # Если есть F, добавляем к существующему комментарию
                            if has_f and comment:
                                # Заменяем точку на точку с запятой и добавляем часть про F
                                comment = comment.rstrip('.') + "; f - любой аллель в диапазоне данного локуса."
                                logger.info("Добавлена информация про f к комментарию")
                            elif has_f and not comment:
                                comment = " f - любой аллель в диапазоне данного локуса."
                                logger.info("Сформирован комментарий только про f")

                        # ЛОГИКА ДЛЯ ОСТАЛЬНЫХ СЛУЧАЕВ
                        else:
                            logger.info("Применение стандартной логики комментария")
                            
                            # Проверка условия для СВО_кость с деградацией генетического материала
                            is_svo_kost = card_data.get("1") == "СВО_кость"
                            vyv_value = card_data.get("ВЫВ", "")
                            has_degradation = vyv_value and "деградации генетического материала" in vyv_value
                            
                            logger.info(f"Проверка условий: is_svo_kost={is_svo_kost}, has_degradation={has_degradation}")
                            logger.debug(f"Значение ВЫВ: {vyv_value}")
                            
                            if has_dash and has_f:
                                if is_svo_kost and has_degradation:
                                    # Особый случай: СВО_кость + прочерк + f + деградация
                                    comment = "(-) - данный локус не представлен в используемой панели, либо нет продуктов амплификации, либо устойчивые данные не получены; (f) – любой аллель в диапазоне данного локуса."
                                    logger.info("Сформирован комментарий для СВО_кость с деградацией (прочерк + f)")
                                elif card_data.get("1") in ["СВО_Молов_образец_родственники", 
                                                        "СВО_Молов_образец_прямая идентификация",
                                                        "СВО_Молов БВП_образец_прямая идентификация",
                                                        "СВО_Ростов_образец_родственники", 
                                                        "СВО_Ростов_образец_прямая идентификация", 
                                                        "ЭКЦ"]:
                                    comment = " (-) - прочерк означает отсутствие продуктов амплификации; f - любой аллель в диапазоне данного локуса."
                                elif is_svo_kost:
                                    comment = " (-) - отсутствие продуктов амплификации или удовлетворительных результатов; f - любой аллель в диапазоне данного локуса."
                                else:
                                    comment = " (-) - прочерк означает отсутствие продуктов амплификации; f - любой аллель в диапазоне данного локуса."
                            elif has_f:
                                if is_svo_kost and has_degradation:
                                    # Особый случай: СВО_кость + f + деградация (без прочерка)
                                    comment = "(f) – любой аллель в диапазоне данного локуса."
                                    logger.info("Сформирован комментарий для СВО_кость с деградацией (только f)")
                                else:
                                    comment = " f - любой аллель в диапазоне данного локуса."
                            elif has_dash:
                                if is_svo_kost and has_degradation:
                                    # Особый случай: СВО_кость + прочерк + деградация (без f)
                                    comment = "(-) - данный локус не представлен в используемой панели, либо нет продуктов амплификации, либо устойчивые данные не получены."
                                    logger.info("Сформирован комментарий для СВО_кость с деградацией (только прочерк)")
                                elif card_data.get("1") in ["СВО_Молов_образец_родственники", 
                                                        "СВО_Молов_образец_прямая идентификация",
                                                        "СВО_Молов БВП_образец_прямая идентификация",
                                                        "СВО_Ростов_образец_родственники", 
                                                        "СВО_Ростов_образец_прямая идентификация", 
                                                        "ЭКЦ"]:
                                    comment = " (-) - прочерк означает отсутствие продуктов амплификации."
                                elif is_svo_kost:
                                    comment = " (-) - отсутствие продуктов амплификации или удовлетворительных результатов."

                        # Добавление комментария
                        if comment:
                            p = doc.add_paragraph()
                            run = p.add_run("Примечание:")
                            run.underline = True
                            run.font.size = Pt(10)
                            run.font.name = 'Times New Roman'
                            p.add_run(f" {comment}").font.size = Pt(10)
                            # Перемещаем примечание сразу после таблицы
                            table_element = table._tbl
                            parent_element = table_element.getparent()
                            parent_element.insert(parent_element.index(table_element) + 1, p._element)
                            logger.info("Комментарий добавлен после таблицы")
                    
                    break

    # Сохраняем документ
    doc.save(doc_path)
    logger.info("Данные успешно вставлены в таблицу и документ сохранён")

# Функция для сравнения образца с ЭКЦ
def process_ekc_sample_comparison(card_data, working_materials_folder, original_txt_path):
    """
    Обрабатывает сравнение с ЭКЦ образцом - единый процесс от начала до конца.
    Возвращает кортеж (успех_обработки, временный_путь_к_файлу, нужно_остановиться, профили_данные, кастомные_заголовки, имя_совпавшего_профиля)
    """
    logger.info("Начало обработки сравнения с ЭКЦ образцом...")

    # Список служебных значений, которые не являются именами профилей
    exclude_values = {"AL", "Al", "al", "K+", "К+", "K-", "К-", "KV", "kv", "KF", "kf"}

    template_value = card_data.get("1", "").strip()
    
    if not template_value.startswith("ЭКЦ_образец_") and template_value != "Тобольск_образец_женщина":
        logger.info("ЭКЦ образец не требуется - пропускаем обработку")
        profiles_data = read_and_filter_txt_data(original_txt_path)  # обычный вызов
        return True, original_txt_path, False, profiles_data, {}, None
    
    logger.info(f"Обнаружен образец для обработки и сравнения: {template_value}")
    
    temp_txt_path = original_txt_path
    
    try:
        # === 1. ПОДГОТОВКА ПУТЕЙ ===
        templates_folder = Path(card_data["3"])
        sample_txt_path = templates_folder / f"{template_value}.txt"
        
        if not sample_txt_path.exists():
            raise FileNotFoundError(f"Файл образца {sample_txt_path} не найден")
        
        # Количество профилей в файле образца должно быть равно 1
        profile_count = count_profiles_in_txt(sample_txt_path)
        if profile_count != 1:
            error_msg = f"Файл образца {sample_txt_path} содержит {profile_count} профилей (ожидается ровно 1). Пожалуйста, проверьте файл и исправьте вручную."
            logger.error(error_msg)
            raise ValueError(error_msg)
        logger.info("Проверка количества профилей в файле образца пройдена: ровно 1 профиль")
        
        # === 2. ПОЛУЧЕНИЕ ЗАГОЛОВКА ПРОФИЛЯ ОБРАЗЦА И НОМЕРА ОБРАЗЦА ===
        with open(sample_txt_path, 'r', encoding='utf-8') as f:
            sample_lines = f.readlines()
        
        if not sample_lines:
            raise ValueError("Файл-образец пуст")
        
        # Ищем заголовок в последней строке
        last_line = sample_lines[-1].strip()
        sample_profile_title = None
        
        # Проверяем, является ли последняя строка заголовком с префиксом
        if last_line.startswith("#HEADER:"):
            sample_profile_title = last_line.replace("#HEADER:", "").strip()
            # Убираем строку заголовка из данных
            sample_data_lines = sample_lines[:-1]
        else:
            # Если префикса нет, используем последнюю строку как есть (для обратной совместимости)
            sample_profile_title = last_line
            sample_data_lines = sample_lines[:-1]
            logger.warning("В файле образца не найден префикс #HEADER: для заголовка")
        
        logger.info(f"Заголовок профиля образца: '{sample_profile_title}'")
        
        # === 2.1. ИЗВЛЕЧЕНИЕ НОМЕРА ОБРАЗЦА ИЗ ДАННЫХ ФАЙЛА ===
        sample_real_profile_name = None
        if sample_data_lines and len(sample_data_lines) > 1:  # Пропускаем заголовок
            for line_idx, line in enumerate(sample_data_lines[1:], start=2):
                parts = line.strip().split('\t')
                if len(parts) == 0:
                    logger.debug(f"Строка {line_idx} пуста, пропущена")
                    continue
                
                candidate_name = parts[0].strip()
                if not candidate_name:
                    logger.debug(f"Строка {line_idx}: имя профиля пустое, пропущена")
                    continue
                
                # Исключаем служебные строки (начинаются с #)
                if candidate_name.startswith('#'):
                    logger.debug(f"Строка {line_idx}: служебная строка '{candidate_name}', пропущена")
                    continue
                
                # Исключаем значения из списка исключений (AL, K+ и т.д.)
                if candidate_name.upper() in exclude_values:
                    logger.debug(f"Строка {line_idx}: имя '{candidate_name}' входит в список исключений, пропущена")
                    continue
                
                # Эвристика: имя профиля должно содержать хотя бы одну цифру
                if not any(char.isdigit() for char in candidate_name):
                    logger.debug(f"Строка {line_idx}: имя '{candidate_name}' не содержит цифр, пропущена")
                    continue
                
                # Валидное имя найдено
                sample_real_profile_name = candidate_name
                logger.info(f"Извлечено имя образца из данных файла (строка {line_idx}): '{sample_real_profile_name}'")
                break

        if not sample_real_profile_name:
            # Если не нашли в данных, пытаемся извлечь из заголовка или использовать fallback
            logger.warning("Не удалось извлечь номер образца из данных файла (после проверки валидности). Используем альтернативные методы...")
            
            # Пытаемся найти любое имя (не обязательно -26) в заголовке
            import re
            # Ищем последовательность цифр, возможно, с дефисом и цифрами (например "123-25" или "456")
            match = re.search(r'(\d+(?:-\d+)?)', sample_profile_title)
            if match:
                sample_real_profile_name = match.group(1)
                logger.info(f"Извлечено имя образца из заголовка: '{sample_real_profile_name}'")
            else:
                # Fallback: берём первое валидное имя из файла образца (без ограничения суффикса)
                for line in sample_data_lines[1:]:
                    parts = line.strip().split('\t')
                    if len(parts) > 0:
                        candidate = parts[0].strip()
                        if candidate and not candidate.startswith('#') and candidate.upper() not in exclude_values:
                            if any(char.isdigit() for char in candidate):
                                sample_real_profile_name = candidate
                                logger.info(f"Использовано fallback-имя образца: '{sample_real_profile_name}'")
                                break

        if not sample_real_profile_name:
            raise ValueError("Не удалось определить имя ЭКЦ образца из файла (нет валидного имени профиля). Проверьте содержимое файла образца.")
        
        # === 3. СОЗДАНИЕ ВРЕМЕННОГО ФАЙЛА ТОЛЬКО С НУЖНЫМИ ПРОФИЛЯМИ ===
        temp_txt_path = working_materials_folder / f"temp_combined_{card_data['НОМ']}.txt"
        logger.info(f"Создание объединённого профиля: основной профиль из {original_txt_path} (только -26) + образец из {sample_txt_path} -> {temp_txt_path}")

        # Читаем рабочий файл
        with open(original_txt_path, 'r', encoding='utf-8') as f:
            working_lines = f.readlines()

        if not working_lines:
            raise ValueError("Рабочий файл пуст")

        # Читаем файл образца
        with open(sample_txt_path, 'r', encoding='utf-8') as f:
            sample_lines = f.readlines()

        if not sample_lines:
            raise ValueError("Файл образца пуст")

        # Отбираем из рабочего файла только строки, где имя профиля оканчивается на -26
        main_profile_lines = [working_lines[0]]  # заголовок
        for line in working_lines[1:]:
            parts = line.strip().split('\t')
            if parts and parts[0].strip().endswith("-26"):
                main_profile_lines.append(line)
            else:
                logger.debug(f"Пропущена строка из рабочего файла (не -26): {line[:50]}")

        if len(main_profile_lines) <= 1:
            raise ValueError("В рабочем файле не найдено строк с профилем, оканчивающимся на -26")

        # Отбираем из файла образца строки, относящиеся к его профилю (исключая служебные)
        sample_profile_lines = [sample_lines[0]]  # заголовок
        exclude_values = {"AL", "Al", "al", "K+", "К+", "K-", "К-", "KV", "kv", "KF", "kf"}

        for line in sample_lines[1:]:
            line_stripped = line.strip()
            if not line_stripped:
                continue
            if line_stripped.startswith("#HEADER:"):
                continue
            parts = line_stripped.split('\t')
            if parts and parts[0].strip():
                profile_name = parts[0].strip()
                if profile_name.upper() in exclude_values:
                    logger.debug(f"Пропущена служебная строка из файла образца: {profile_name}")
                    continue
                if profile_name.startswith('#'):
                    continue
                # Для файла образца мы берём все неслужебные строки (предполагается 1 профиль)
                sample_profile_lines.append(line)
            else:
                logger.debug(f"Пропущена пустая строка из файла образца")

        # Записываем объединённый файл
        with open(temp_txt_path, 'w', encoding='utf-8') as f:
            f.writelines(main_profile_lines)
            f.writelines(sample_profile_lines[1:])  # пропускаем дублирование заголовка, если он уже есть

        logger.info(f"Объединённый файл создан: {temp_txt_path}")
        
        # === 4. ЧТЕНИЕ И ФИЛЬТРАЦИЯ ДАННЫХ ИЗ ВРЕМЕННОГО ФАЙЛА (С ИЗМЕНЕНИЕМ) ===
        logger.info(f"Чтение данных из временного файла: {temp_txt_path}")
        # Используем require_suffix=False, чтобы не отфильтровать профиль образца
        profiles_data = read_and_filter_txt_data(temp_txt_path, require_suffix=False)
        logger.info(f"Найдены профили в объединённом файле: {list(profiles_data.keys())}")
        
        # === 5. ПРОВЕРКА НАЛИЧИЯ ПРОФИЛЯ ОБРАЗЦА В ДАННЫХ ===
        if sample_real_profile_name not in profiles_data:
            logger.warning(f"Профиль образца {sample_real_profile_name} не найден в объединенных данных")
            logger.info(f"Доступные профили: {list(profiles_data.keys())}")
            
            # Пытаемся найти альтернативное совпадение
            alternative_found = False
            for profile_name in profiles_data.keys():
                if profile_name.endswith("-26") and profile_name != sample_real_profile_name:
                    logger.info(f"Найден альтернативный профиль для сравнения: {profile_name}")
                    sample_real_profile_name = profile_name
                    alternative_found = True
                    break
            
            if not alternative_found:
                raise ValueError(f"Профиль образца {sample_real_profile_name} не найден в данных и не найдено альтернатив")
        
        sample_profile = profiles_data[sample_real_profile_name]
        
        # === 6. СОЗДАНИЕ КАСТОМНЫХ ЗАГОЛОВКОВ ===
        custom_headers = {}
        if sample_real_profile_name and sample_profile_title:
            custom_headers[sample_real_profile_name] = sample_profile_title
            logger.info(f"Добавлен кастомный заголовок для профиля {sample_real_profile_name}")
        
        # === 7. СРАВНЕНИЕ ПРОФИЛЕЙ ===
        logger.info(f"Начало сравнения профилей. Образец: {sample_profile_title}, номер: {sample_real_profile_name}")
        
        # Создаем словарь локусов и аллелей для образца
        sample_loci = {}
        for line in sample_profile:
            marker = line[3].strip().upper()
            allele1 = line[5].strip() if len(line) > 5 and line[5] else ""
            allele2 = line[6].strip() if len(line) > 6 and line[6] else ""
            
            # Сортируем аллели для consistent сравнения
            alleles = sorted([allele1, allele2])
            sample_loci[marker] = alleles
            logger.debug(f"Образец - Локус {marker}: аллели {alleles}")
        
        # Сравниваем с каждым другим профилем
        profile_match_found = False
        partial_match_found = False
        matching_profile_name = None
        partial_match_info = None
        
        for profile_name, profile_data in profiles_data.items():
            if profile_name == sample_real_profile_name:
                continue
                
            logger.info(f"Сравнение образца с профилем: {profile_name}")
            
            # Создаем словарь для текущего профиля
            current_loci = {}
            for line in profile_data:
                marker = line[3].strip().upper()
                allele1 = line[5].strip() if len(line) > 5 and line[5] else ""
                allele2 = line[6].strip() if len(line) > 6 and line[6] else ""
                
                alleles = sorted([allele1, allele2])
                current_loci[marker] = alleles
                logger.debug(f"Профиль {profile_name} - Локус {marker}: аллели {alleles}")
            
            # Сравниваем общие локусы
            common_loci = set(sample_loci.keys()) & set(current_loci.keys())
            logger.info(f"Общие локусы для сравнения: {common_loci}")
            
            if not common_loci:
                logger.warning("Нет общих локусов для сравнения")
                continue
            
            # Подсчитываем совпадения и несовпадения
            total_loci = len(common_loci)
            matched_loci = 0
            mismatched_loci = 0
            
            for locus in common_loci:
                sample_alleles = sample_loci[locus]
                current_alleles = current_loci[locus]
                
                logger.debug(f"Сравнение локуса {locus}: образец {sample_alleles} vs профиль {current_alleles}")
                
                if sample_alleles == current_alleles:
                    matched_loci += 1
                    logger.debug(f"Локус {locus} совпадает")
                else:
                    mismatched_loci += 1
                    logger.debug(f"Локус {locus} не совпадает")
            
            logger.info(f"Результат сравнения: {matched_loci} совпадений, {mismatched_loci} несовпадений из {total_loci} локусов")
            
            # Проверяем условия
            if mismatched_loci == 0:
                # Полное совпадение
                logger.error(f"Обнаружено полное совпадение профилей! Образец: {sample_real_profile_name}, Профиль: {profile_name}")
                profile_match_found = True
                matching_profile_name = profile_name
                break
            elif 1 <= mismatched_loci <= 2:
                # Почти полное совпадение
                logger.warning(f"Обнаружено почти полное совпадение профилей: {matched_loci} локусов из {total_loci}! Образец: {sample_real_profile_name}, Профиль: {profile_name}")
                partial_match_found = True
                partial_match_info = {
                    'profile_name': profile_name,
                    'matched_loci': matched_loci,
                    'total_loci': total_loci,
                    'sample_name': sample_real_profile_name,
                    'profile_title': sample_profile_title
                }
                # Не прерываем цикл, чтобы проверить другие профили на полное совпадение
                # Но запоминаем только первое частичное совпадение
                if matching_profile_name is None:
                    matching_profile_name = profile_name
        
        # === 8. ОБРАБОТКА РЕЗУЛЬТАТОВ СРАВНЕНИЯ ===
        if profile_match_found:
            logger.error("Обнаружено полное совпадение профилей!")
            # ВОЗВРАЩАЕМ УСПЕХ, НО С ФЛАГОМ ОСТАНОВКИ - таблица будет заполнена перед выводом ошибки
            return True, temp_txt_path, True, profiles_data, custom_headers, matching_profile_name
        elif partial_match_found and partial_match_info:
            # Почти полное совпадение - запрашиваем действие у пользователя
            logger.warning("Обнаружено почти полное совпадение профилей!")
            
            message = (f"Обнаружено почти полное совпадение профилей: {partial_match_info['matched_loci']} локусов из {partial_match_info['total_loci']}! "
                      f"Образец: {partial_match_info['sample_name']}, Профиль: {partial_match_info['profile_name']}. Продолжить работу? (д/н)")
            
            print("\n" + "="*80)
            print(message)
            print("="*80)
            
            # Запрашиваем подтверждение
            while True:
                try:
                    response = input("\nВведите ваш выбор (д/н): ").strip().lower()
                    if response in ['д', 'да', 'y', 'yes']:
                        logger.info("Пользователь выбрал продолжить работу при почти полном совпадении")
                        print("Продолжаем работу...")
                        return True, temp_txt_path, False, profiles_data, custom_headers, None
                    elif response in ['н', 'нет', 'n', 'no']:
                        logger.info("Пользователь выбрал завершить работу при почти полном совпадении")
                        print("Завершение работы программы по требованию пользователя...")
                        import sys
                        sys.exit(0)
                    else:
                        print("Некорректный ввод. Пожалуйста, введите 'д' для продолжения или 'н' для завершения.")
                except KeyboardInterrupt:
                    logger.warning("Пользователь прервал ввод с помощью Ctrl+C")
                    print("\nПрервано пользователем. Завершение работы...")
                    import sys
                    sys.exit(0)
        
        logger.info("Сравнение профилей завершено успешно, совпадений не найдено")
        return True, temp_txt_path, False, profiles_data, custom_headers, None
        
    except Exception as e:
        logger.error(f"Ошибка при обработке ЭКЦ образца: {e}")
        # В случае ошибки возвращаем оригинальный путь и данные
        # Инициализируем profiles_data здесь на случай ошибки
        try:
            profiles_data = read_and_filter_txt_data(original_txt_path)
        except:
            profiles_data = {}
        return False, original_txt_path, False, profiles_data, {}, None

def process_tobolsk_sample(card_data, working_materials_folder, original_txt_path):
    """
    Обрабатывает шаблон "Тобольск_образец_мужчина":
    - объединяет основной профиль (-26) и профиль образца,
    - возвращает профили, кастомные заголовки и имя профиля образца.
    
    Возвращает кортеж (успех_обработки, временный_путь_к_файлу, profiles_data, custom_headers, reference_profile_name)
    """
    logger.info("=== Начало обработки шаблона Тобольск_образец_мужчина ===")
    
    # Список служебных значений
    exclude_values = {"AL", "Al", "al", "K+", "К+", "K-", "К-", "KV", "kv", "KF", "kf"}

    template_value = card_data.get("1", "").strip()
    if template_value != "Тобольск_образец_мужчина":
        logger.info("Шаблон не требует специальной обработки Тобольск – пропускаем")
        return False, original_txt_path, None, {}, None
    
    logger.info(f"Обнаружен шаблон Тобольск: {template_value}")
    
    # Пути
    templates_folder = Path(card_data["3"])
    sample_txt_path = templates_folder / f"{template_value}.txt"
    
    if not sample_txt_path.exists():
        raise FileNotFoundError(f"Файл образца не найден: {sample_txt_path}")
    
    # Проверяем, что в файле образца ровно один профиль
    profile_count = count_profiles_in_txt(sample_txt_path)
    if profile_count != 1:
        error_msg = (f"Файл образца {sample_txt_path} содержит {profile_count} профилей, "
                     f"ожидается ровно 1. Проверьте файл.")
        logger.error(error_msg)
        raise ValueError(error_msg)
    logger.info("Файл образца содержит ровно один профиль")
    
    # Извлекаем заголовок профиля образца (последняя строка с #HEADER:)
    with open(sample_txt_path, 'r', encoding='utf-8') as f:
        sample_lines = f.readlines()
    
    sample_profile_title = None
    sample_data_lines = sample_lines
    if sample_lines and sample_lines[-1].strip().startswith("#HEADER:"):
        sample_profile_title = sample_lines[-1].strip().replace("#HEADER:", "").strip()
        sample_data_lines = sample_lines[:-1]   # удаляем строку заголовка из данных
        logger.info(f"Заголовок образца: '{sample_profile_title}'")
    else:
        logger.warning("В файле образца не найден заголовок #HEADER: – будет использовано стандартное имя")
        sample_profile_title = "Образец"
    
    # Определяем реальное имя профиля образца (теперь любой суффикс)
    sample_real_profile_name = None

    for line_idx, line in enumerate(sample_data_lines, start=1):
        if line_idx == 1:
            continue  # пропускаем строку заголовка колонок
        parts = line.strip().split('\t')
        if not parts:
            continue
        candidate = parts[0].strip()
        if not candidate:
            continue
        if candidate.startswith('#'):
            logger.debug(f"Тобольск: строка {line_idx} — служебная, пропущена")
            continue
        if candidate.upper() in exclude_values:
            logger.debug(f"Тобольск: имя '{candidate}' в списке исключений, пропущена")
            continue
        if not any(ch.isdigit() for ch in candidate):
            logger.debug(f"Тобольск: имя '{candidate}' без цифр, пропущена")
            continue
        # Нашли первое валидное имя
        sample_real_profile_name = candidate
        logger.info(f"Тобольск: извлечено имя образца (строка {line_idx}): '{sample_real_profile_name}'")
        break

    if not sample_real_profile_name:
        raise ValueError("Не удалось определить имя профиля образца Тобольск (нет валидного имени в файле). Проверьте файл образца.")
    
    # Создаём временный объединённый файл
    temp_txt_path = working_materials_folder / f"temp_tobolsk_{card_data['НОМ']}.txt"
    logger.info(f"Создание объединённого файла: {temp_txt_path}")
    
    # Читаем рабочий файл (оригинальный)
    with open(original_txt_path, 'r', encoding='utf-8') as f:
        working_lines = f.readlines()
    if not working_lines:
        raise ValueError("Рабочий файл пуст")
    
    # Отбираем из рабочего файла только профили, оканчивающиеся на -26
    main_profile_lines = [working_lines[0]]   # заголовок
    for line in working_lines[1:]:
        parts = line.strip().split('\t')
        if parts and parts[0].strip().endswith("-26"):
            main_profile_lines.append(line)
        else:
            logger.debug(f"Пропущена строка (не -26): {line[:50]}")
    
    if len(main_profile_lines) <= 1:
        raise ValueError("В рабочем файле нет ни одного профиля с суффиксом -26")
    
    # Отбираем строки образца (исключая служебные)
    exclude_values = {"AL", "Al", "al", "K+", "К+", "K-", "К-", "KV", "kv", "KF", "kf"}
    sample_profile_lines = [sample_lines[0]]   # заголовок
    for line in sample_data_lines[1:]:        # пропускаем первую строку (заголовок колонок)
        line_stripped = line.strip()
        if not line_stripped or line_stripped.startswith("#HEADER:"):
            continue
        parts = line_stripped.split('\t')
        if parts and parts[0].strip():
            name = parts[0].strip()
            if name.upper() in exclude_values or name.startswith('#'):
                continue
            sample_profile_lines.append(line)
    
    # Записываем объединённый файл
    with open(temp_txt_path, 'w', encoding='utf-8') as f:
        f.writelines(main_profile_lines)
        # Добавляем строки образца (без дублирования заголовка)
        f.writelines(sample_profile_lines[1:])
    
    logger.info(f"Объединённый файл создан: {temp_txt_path}")
    
    # Читаем данные из временного файла (require_suffix=False, чтобы включить образец)
    profiles_data = read_and_filter_txt_data(temp_txt_path, require_suffix=False)
    logger.info(f"Найдены профили в объединённом файле: {list(profiles_data.keys())}")
    
    # Создаём кастомный заголовок для профиля образца
    custom_headers = {}
    if sample_real_profile_name in profiles_data:
        custom_headers[sample_real_profile_name] = sample_profile_title
        logger.info(f"Добавлен кастомный заголовок для профиля {sample_real_profile_name}")
    else:
        logger.warning(f"Профиль образца {sample_real_profile_name} не найден в объединённых данных")
    
    logger.info("Обработка Тобольск_образец_мужчина завершена успешно (без сравнения)")
    logger.info("=== Конец обработки Тобольск_образец_мужчина ===")
    
    # Возвращаем дополнительно имя профиля образца
    return True, temp_txt_path, profiles_data, custom_headers, sample_real_profile_name

def split_text_by_length(text, max_length=50):
    """
    Универсальная функция разбиения текста на части.
    Старается разбивать по логическим местам (пробелы, запятые), а не по словам.
    """
    logger.info(f"Разбитие текста: '{text}' на части по {max_length} символов")
    
    # Если текст короткий, возвращаем как есть
    if len(text) <= max_length:
        return [text]
    
    parts = []
    start = 0
    
    while start < len(text):
        # Определяем конец текущего сегмента
        end = start + max_length
        
        # Если дошли до конца текста
        if end >= len(text):
            parts.append(text[start:].strip())
            break
        
        # Ищем лучшее место для разрыва - сначала пробел, затем запятую, затем любой символ
        break_positions = [
            text.rfind(' ', start, end),  # Пробел
            text.rfind(',', start, end),  # Запятая
            text.rfind('.', start, end),  # Точка
            end  # Просто обрезаем
        ]
        
        # Выбираем первую валидную позицию разрыва
        break_pos = -1
        for pos in break_positions:
            if pos > start and pos <= end:
                break_pos = pos
                break
        
        # Если нашли хорошее место для разрыва
        if break_pos > start:
            part = text[start:break_pos].strip()
            if part:  # Проверяем, что часть не пустая
                parts.append(part)
            start = break_pos
        else:
            # Вынужденный разрыв по max_length
            part = text[start:end].strip()
            if part:
                parts.append(part)
            start = end
    
    logger.info(f"Текст разбит на {len(parts)} частей: {parts}")
    return parts

def process_table_cell_wrapping(doc_path, max_line_length=85):
    """
    Обрабатывает таблицы с меткой [WRAP_TEXT], создавая новые строки при необходимости.
    Добавляет нижнюю границу для новых строк.
    Улучшенная версия с детальными проверками, но без записи лог-файла.
    """
    logger.info("Обработка переноса текста в таблицах с маркерами...")
    logger.info("=== Начало обработки переноса текста в таблицах ===")
    
    # Создаем логгер для диагностики (без файла)
    debug_logger, debug_log_path = setup_wrap_text_debug_logger()
    
    # Для важных диагностических сообщений используем основной логгер с уровнем DEBUG
    logger.debug(f"[WRAP_TEXT_DEBUG] Начало обработки файла: {doc_path}")
    logger.debug(f"[WRAP_TEXT_DEBUG] Максимальная длина строки: {max_line_length}")
    
    try:
        # Детальная проверка файла
        logger.debug(f"[WRAP_TEXT_DEBUG] Проверка существования файла: {doc_path}")
        
        if not os.path.exists(doc_path):
            error_msg = f"Файл не существует: {doc_path}"
            logger.error(error_msg)
            raise FileNotFoundError(error_msg)
        
        logger.debug("[WRAP_TEXT_DEBUG] ✓ Файл существует")
        
        # Проверяем размер файла
        try:
            file_size = os.path.getsize(doc_path)
            logger.debug(f"[WRAP_TEXT_DEBUG] Размер файла: {file_size} байт")
        except Exception as size_error:
            logger.warning(f"[WRAP_TEXT_DEBUG] Не удалось получить размер файла: {size_error}")
        
        # Проверяем права доступа
        logger.debug("[WRAP_TEXT_DEBUG] Проверка прав доступа к файлу")
        try:
            logger.debug(f"[WRAP_TEXT_DEBUG] Файл доступен для чтения: {os.access(doc_path, os.R_OK)}")
            logger.debug(f"[WRAP_TEXT_DEBUG] Файл доступен для записи: {os.access(doc_path, os.W_OK)}")
        except Exception as access_error:
            logger.warning(f"[WRAP_TEXT_DEBUG] Не удалось проверить права доступа: {access_error}")
        
        # Попытка открытия документа
        logger.debug("[WRAP_TEXT_DEBUG] Попытка открытия документа...")
        
        try:
            doc = Document(doc_path)
            logger.debug("[WRAP_TEXT_DEBUG] ✓ Документ успешно открыт через python-docx")
        except Exception as doc_error:
            error_msg = f"Ошибка при открытии документа: {doc_error}"
            logger.error(error_msg)
            raise doc_error
        
        tables_processed = 0
        rows_processed = 0
        cells_with_markers = 0
        
        logger.debug(f"[WRAP_TEXT_DEBUG] Общее количество таблиц в документе: {len(doc.tables)}")
        
        # Вложенная функция для установки границ ячейки
        def set_cell_border(cell, border_style="single", border_size=4, border_color="000000"):
            """
            Устанавливает только нижнюю границу для ячейки таблицы.
            """
            try:
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                
                # Удаляем существующие границы, если есть
                for element in tcPr.xpath('.//w:tcBorders'):
                    tcPr.remove(element)
                
                # Создаем новые границы - только нижнюю
                tcBorders = OxmlElement('w:tcBorders')
                
                bottom_border = OxmlElement('w:bottom')
                bottom_border.set(qn('w:val'), border_style)
                bottom_border.set(qn('w:sz'), str(border_size))
                bottom_border.set(qn('w:space'), '0')
                bottom_border.set(qn('w:color'), border_color)
                tcBorders.append(bottom_border)
                
                tcPr.append(tcBorders)
                logger.debug("[WRAP_TEXT_DEBUG] ✓ Границы ячейки успешно установлены")
                return True
            except Exception as e:
                error_msg = f"Ошибка при установке границ ячейки: {e}"
                logger.error(error_msg)
                return False
        
        # Обработка таблиц
        logger.debug("[WRAP_TEXT_DEBUG] Начало обработки таблиц")
        for table_index, table in enumerate(doc.tables):
            logger.info(f"Проверка таблицы {table_index + 1} на наличие маркеров [WRAP_TEXT]")
            
            for row_index, row in enumerate(table.rows):
                for cell_index, cell in enumerate(row.cells):
                    cell_text = cell.text
                    
                    if "[WRAP_TEXT]" in cell_text:
                        cells_with_markers += 1
                        logger.info(f"Найдена ячейка с маркером [WRAP_TEXT]: строка {row_index + 1}, ячейка {cell_index + 1}")
                        logger.debug(f"[WRAP_TEXT_DEBUG] Полный текст ячейки: '{cell_text}'")
                        
                        # Удаляем метку из текста
                        original_text = cell_text.replace("[WRAP_TEXT]", "").strip()
                        logger.debug(f"[WRAP_TEXT_DEBUG] Текст после удаления метки: '{original_text}' (длина: {len(original_text)})")
                        
                        # Проверяем длину текста
                        if len(original_text) > max_line_length:
                            logger.debug(f"[WRAP_TEXT_DEBUG] Текст превышает максимальную длину ({max_line_length}), начинаем разбиение")
                            
                            try:
                                text_parts = split_text_by_length(original_text, max_line_length)
                                logger.debug(f"[WRAP_TEXT_DEBUG] Результат разбиения: {len(text_parts)} частей")
                                
                                logger.info(f"Разбитие текста: '{original_text}' на части по {max_line_length} символов")
                                logger.info(f"Текст разбит на {len(text_parts)} частей: {text_parts}")
                                
                                if len(text_parts) > 1:
                                    logger.info(f"Текст будет разбит на {len(text_parts)} части")
                                    
                                    # Вставляем первую часть в исходную ячейку
                                    logger.debug(f"[WRAP_TEXT_DEBUG] Установка первой части в исходную ячейку: '{text_parts[0]}'")
                                    cell.text = text_parts[0]
                                    
                                    # Для каждой дополнительной части создаем новую строку
                                    for part_index in range(1, len(text_parts)):
                                        logger.debug(f"[WRAP_TEXT_DEBUG] Создание новой строки для части {part_index + 1}")
                                        
                                        try:
                                            new_row = table.add_row()
                                            
                                            # Удаляем лишние ячейки, оставляем только одну объединенную
                                            while len(new_row.cells) > 1:
                                                last_cell = new_row.cells[-1]
                                                last_cell._element.getparent().remove(last_cell._element)
                                            
                                            # Если не осталось ячеек, добавляем одну
                                            if len(new_row.cells) == 0:
                                                new_row.add_cell()
                                            
                                            # Вставляем текст
                                            new_cell = new_row.cells[0]
                                            new_cell.text = text_parts[part_index]
                                            logger.debug(f"[WRAP_TEXT_DEBUG] В новую ячейку вставлен текст: '{text_parts[part_index]}'")
                                            
                                            # Устанавливаем нижнюю границу для новой ячейки
                                            border_result = set_cell_border(new_cell)
                                            logger.debug(f"[WRAP_TEXT_DEBUG] Установка границ: {'успешно' if border_result else 'неудачно'}")
                                            
                                            logger.debug(f"Создана новая строка для части {part_index + 1} с нижней границей")
                                            
                                        except Exception as row_error:
                                            error_msg = f"Ошибка при создании новой строки для части {part_index + 1}: {row_error}"
                                            logger.error(error_msg)
                                            raise row_error
                                    
                                    logger.info(f"Создано {len(text_parts) - 1} дополнительных строк с границами")
                                    logger.debug(f"[WRAP_TEXT_DEBUG] ✓ Успешно создано {len(text_parts) - 1} дополнительных строк")
                                    rows_processed += 1
                                else:
                                    logger.debug("[WRAP_TEXT_DEBUG] Текст не требует разбиения на части")
                                    
                            except Exception as split_error:
                                error_msg = f"Ошибка при разбиении текста: {split_error}"
                                logger.error(error_msg)
                                # Продолжаем работу, просто вставляем текст без метки
                                cell.text = original_text
                                logger.debug("[WRAP_TEXT_DEBUG] Текст вставлен без разбиения из-за ошибки")
                        else:
                            # Если текст короткий, просто вставляем его без метки
                            logger.debug("[WRAP_TEXT_DEBUG] Текст не превышает максимальную длину, вставляем без разбиения")
                            cell.text = original_text
                            logger.debug(f"[WRAP_TEXT_DEBUG] Текст вставлен: '{original_text}'")
                        
                        tables_processed += 1
                        break  # Переходим к следующей строке
        
        logger.debug("[WRAP_TEXT_DEBUG] Попытка сохранения документа")
        
        try:
            # Дополнительная проверка перед сохранением
            logger.debug("[WRAP_TEXT_DEBUG] Проверка возможности записи перед сохранением")
            if not os.access(os.path.dirname(doc_path), os.W_OK):
                logger.warning("[WRAP_TEXT_DEBUG] Папка недоступна для записи!")
            
            doc.save(doc_path)
            logger.debug("[WRAP_TEXT_DEBUG] ✓ Документ успешно сохранен")
            
        except Exception as save_error:
            error_msg = f"КРИТИЧЕСКАЯ ОШИБКА ПРИ СОХРАНЕНИИ: {save_error}"
            logger.error(error_msg)
            logger.error(f"[WRAP_TEXT_DEBUG] Детали ошибки: {str(save_error)}")
            
            # Дополнительная диагностика
            try:
                logger.error(f"[WRAP_TEXT_DEBUG] Файл существует: {os.path.exists(doc_path)}")
                logger.error(f"[WRAP_TEXT_DEBUG] Папка существует: {os.path.exists(os.path.dirname(doc_path))}")
                logger.error(f"[WRAP_TEXT_DEBUG] Путь абсолютный: {os.path.isabs(doc_path)}")
                logger.error(f"[WRAP_TEXT_DEBUG] Длина пути: {len(doc_path)}")
                
            except Exception as diag_error:
                logger.error(f"[WRAP_TEXT_DEBUG] Ошибка диагностики: {diag_error}")
            
            raise save_error
        
        logger.info(f"Обработка завершена: обработано {tables_processed} таблиц, {rows_processed} строк")
        logger.debug(f"[WRAP_TEXT_DEBUG] Итоговая статистика:")
        logger.debug(f"[WRAP_TEXT_DEBUG] Обработано таблиц: {tables_processed}")
        logger.debug(f"[WRAP_TEXT_DEBUG] Обработано строк с разбиением текста: {rows_processed}")
        logger.debug(f"[WRAP_TEXT_DEBUG] Найдено ячеек с маркерами [WRAP_TEXT]: {cells_with_markers}")
        
        logger.info("=== Конец обработки переноса текста в таблицах ===")
        
    except Exception as e:
        error_msg = f"Ошибка при обработке переноса текста: {str(e)}"
        
        logger.error("[WRAP_TEXT_DEBUG] КРИТИЧЕСКАЯ ОШИБКА В ПРОЦЕССЕ ОБРАБОТКИ")
        logger.error(f"[WRAP_TEXT_DEBUG] Тип ошибки: {type(e).__name__}")
        logger.error(f"[WRAP_TEXT_DEBUG] Сообщение об ошибке: {str(e)}")
        logger.error(f"[WRAP_TEXT_DEBUG] Файл: {doc_path}")
        
        # Дополнительная информация об ошибке
        import traceback
        logger.error("[WRAP_TEXT_DEBUG] Полная трассировка стека:")
        logger.error(traceback.format_exc())
        
        logger.error(error_msg)
        
        raise

# Блок проверки консистентности информации в заключении
# Сравнивает номер из Карты, номер с конверта, номер файла txt и номер заключения из фореграммы
# Сравнивает данные по родству из Карты и наличие Y
def preprocess_image_for_ocr(image_path):
    """
    Предварительная обработка изображения для улучшения распознавания OCR.
    """
    # Загружаем изображение с помощью Pillow
    image_pil = Image.open(image_path)
    # Увеличиваем контрастность и яркость
    enhancer = ImageEnhance.Contrast(image_pil)
    image_pil = enhancer.enhance(2.0)
    enhancer = ImageEnhance.Brightness(image_pil)
    image_pil = enhancer.enhance(1.5)
    # Преобразуем в numpy array для работы с OpenCV
    image = np.array(image_pil)
    # Преобразуем из RGB (Pillow) в BGR (OpenCV)
    image = cv2.cvtColor(image, cv2.COLOR_RGB2BGR)
    # Преобразование в оттенки серого
    gray = cv2.cvtColor(image, cv2.COLOR_BGR2GRAY)
    # Удаление шумов
    gray = cv2.GaussianBlur(gray, (5, 5), 0)
    return gray

def determine_gender_from_parents(card_data):
    """
    Определяет пол подэкспертного на основе данных о родителях.
    """
    parent_info = card_data.get("РОД")
    if parent_info in ["матери", "сестры", "дочери"]:
        return "женский"
    elif parent_info in ["отца", "брата", "сына"]:
        return "мужской"
    return None

def determine_gender_from_amel(txt_data):
    """
    Определяет пол подэкспертного на основе данных локуса AMEL.
    """
    for line in txt_data:
        if line[3].strip().upper() == "AMEL":
            alleles = [line[5], line[6]]
            if "Y" in alleles:
                return "мужской"
            elif "X" in alleles and "Y" not in alleles:
                return "женский"
    return None

def check_gender(card_data, txt_data):
    """
    Проверяет пол подэкспертного на основе данных о родителях и локуса AMEL.
    """
    
    if txt_data is None:  # Защита от None
        return
    
    gender_from_parents = determine_gender_from_parents(card_data)
    gender_from_amel = determine_gender_from_amel(txt_data)

    if gender_from_parents and gender_from_amel and gender_from_parents != gender_from_amel:
        raise ValueError(f"Несоответствие данных о поле подэкспертного: родители указывают на {gender_from_parents}, локус AMEL указывает на {gender_from_amel}.")

def extract_case_number_from_image(image_path, working_materials_folder, delete_debug_files=True, use_full_image=False):
    """
    Извлечение номера заключения с комбинированным подходом.
    Использует несколько методов поиска рамки и улучшенную обработку.
    
    Параметры:
        image_path (str): путь к исходному изображению.
        working_materials_folder (str): папка для временных и отладочных файлов.
        delete_debug_files (bool): удалять ли отладочные файлы после работы.
        use_full_image (bool): если True, обрабатывается всё изображение целиком,
                               иначе только верхний левый квадрант (поведение по умолчанию).
    """
    temp_roi_path = None
    temp_cropped_path = None
    temp_debug_files = []
    
    try:
        logger.info(f"Начало обработки изображения (комбинированный метод): {image_path}")
        logger.info(f"Режим анализа: {'ВСЁ ИЗОБРАЖЕНИЕ' if use_full_image else 'ВЕРХНИЙ ЛЕВЫЙ КВАДРАНТ'}")
        
        # 1. Загрузка и предобработка оригинального изображения
        original_image_pil = Image.open(image_path)
        original_image = np.array(original_image_pil)
        original_image = cv2.cvtColor(original_image, cv2.COLOR_RGB2BGR)
        logger.debug(f"Оригинальное изображение загружено. Размер: {original_image.shape}")
        
        # 2. Определяем, обрезать ли до квадранта
        if use_full_image:
            logger.info("Используется полное изображение (use_full_image=True). Пропуск обрезки до квадранта.")
            cropped_image = original_image  # обрабатываем всё изображение
        else:
            # Обрезка до верхнего левого квадранта (стандартное поведение)
            height, width = original_image.shape[:2]
            half_width = width // 2
            half_height = height // 2
            cropped_image = original_image[0:half_height, 0:half_width]
            logger.info(f"Изображение обрезано до верхнего левого квадранта. Новый размер: {cropped_image.shape}")
        
        # Сохраняем изображение для анализа (квадрант или полное)
        cropped_image_pil = Image.fromarray(cv2.cvtColor(cropped_image, cv2.COLOR_BGR2RGB))
        temp_cropped_path = Path(working_materials_folder) / "temp_cropped_analysis_DEBUG.jpg"
        cropped_image_pil.save(temp_cropped_path)
        temp_debug_files.append(temp_cropped_path)
        logger.info(f"Сохранено изображение для анализа: {temp_cropped_path}")
        
        # 3. Подготовка изображения в градациях серого для обработки
        gray_cropped = cv2.cvtColor(cropped_image, cv2.COLOR_BGR2GRAY)
        
        # Список для хранения всех найденных кандидатов из разных методов
        all_candidates = []
        
        # ========== МЕТОД 1: Адаптивная бинаризация ==========
        logger.info("Попытка 1: Адаптивная бинаризация")
        try:
            # Адаптивная бинаризация лучше работает с неравномерным освещением
            adaptive_thresh = cv2.adaptiveThreshold(
                gray_cropped, 255, 
                cv2.ADAPTIVE_THRESH_GAUSSIAN_C, 
                cv2.THRESH_BINARY_INV, 
                blockSize=21,  # Размер окрестности (должен быть нечетным)
                C=10  # Константа, вычитаемая из средневзвешенного
            )
            
            # Морфологические операции для улучшения рамки
            kernel = cv2.getStructuringElement(cv2.MORPH_RECT, (3, 3))
            adaptive_thresh = cv2.morphologyEx(adaptive_thresh, cv2.MORPH_CLOSE, kernel)
            
            # Сохранение для отладки
            adaptive_thresh_pil = Image.fromarray(adaptive_thresh)
            temp_adaptive_path = Path(working_materials_folder) / "temp_adaptive_thresh_DEBUG.jpg"
            adaptive_thresh_pil.save(temp_adaptive_path)
            temp_debug_files.append(temp_adaptive_path)  # Добавляем в список для удаления
            logger.info(f"Сохранено адаптивно бинаризованное изображение: {temp_adaptive_path}")
            
            # Поиск контуров
            contours_adaptive, _ = cv2.findContours(adaptive_thresh, cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_SIMPLE)
            logger.debug(f"Найдено {len(contours_adaptive)} контуров методом адаптивной бинаризации")
            
            # Фильтрация контуров
            candidates_adaptive = filter_rectangle_contours(
                contours_adaptive, cropped_image, 
                min_area=300, method_name="adaptive"
            )
            all_candidates.extend(candidates_adaptive)
            
        except Exception as e:
            logger.warning(f"Ошибка при адаптивной бинаризации: {e}")
        
        # ========== МЕТОД 2: Обычная бинаризация с несколькими порогами ==========
        logger.info("Попытка 2: Обычная бинаризация с разными порогами")
        thresholds_to_try = [60, 80, 100, 120]  # Пробуем разные пороги
        
        for thresh_value in thresholds_to_try:
            try:
                logger.debug(f"Пробуем порог: {thresh_value}")
                _, binary_thresh = cv2.threshold(gray_cropped, thresh_value, 255, cv2.THRESH_BINARY_INV)
                
                # Морфологические операции
                kernel = cv2.getStructuringElement(cv2.MORPH_RECT, (5, 5))
                binary_thresh = cv2.morphologyEx(binary_thresh, cv2.MORPH_CLOSE, kernel)
                
                # Поиск контуров
                contours_binary, _ = cv2.findContours(binary_thresh, cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_SIMPLE)
                
                # Фильтрация контуров
                candidates_binary = filter_rectangle_contours(
                    contours_binary, cropped_image, 
                    min_area=300, method_name=f"binary_thresh_{thresh_value}"
                )
                all_candidates.extend(candidates_binary)
                
            except Exception as e:
                logger.warning(f"Ошибка при бинаризации с порогом {thresh_value}: {e}")
        
        # ========== МЕТОД 3: Детектор Canny с разными параметрами ==========
        logger.info("Попытка 3: Детектор Canny с разными параметрами")
        canny_params = [(30, 80), (50, 100), (40, 120)]  # Разные комбинации порогов
        
        for low_thresh, high_thresh in canny_params:
            try:
                logger.debug(f"Пробуем Canny с порогами: {low_thresh}, {high_thresh}")
                blurred = cv2.GaussianBlur(gray_cropped, (5, 5), 0)
                edges = cv2.Canny(blurred, low_thresh, high_thresh)
                
                # Поиск контуров
                contours_canny, _ = cv2.findContours(edges, cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_SIMPLE)
                
                # Фильтрация контуров
                candidates_canny = filter_rectangle_contours(
                    contours_canny, cropped_image, 
                    min_area=300, method_name=f"canny_{low_thresh}_{high_thresh}"
                )
                all_candidates.extend(candidates_canny)
                
            except Exception as e:
                logger.warning(f"Ошибка при Canny с порогами {low_thresh}, {high_thresh}: {e}")
        
        # ========== Выбор лучшего кандидата ==========
        logger.info(f"Всего найдено кандидатов: {len(all_candidates)}")
        
        if all_candidates:
            # Сортируем по площади и выбираем лучший
            all_candidates.sort(key=lambda x: x[5], reverse=True)
            
            # Визуализация всех кандидатов для отладки
            debug_all_candidates_image = cropped_image.copy()
            for i, (cnt, x, y, w, h, area, method) in enumerate(all_candidates[:5]):  # Показываем топ-5
                color = (0, 255 - i*50, i*50)  # Разные цвета для разных кандидатов
                cv2.rectangle(debug_all_candidates_image, (x, y), (x+w, y+h), color, 2)
                cv2.putText(debug_all_candidates_image, f"{i+1}: {method}", 
                           (x, y-5), cv2.FONT_HERSHEY_SIMPLEX, 0.5, color, 1)
            
            debug_candidates_pil = Image.fromarray(cv2.cvtColor(debug_all_candidates_image, cv2.COLOR_BGR2RGB))
            temp_all_candidates_path = Path(working_materials_folder) / "temp_all_candidates_DEBUG.jpg"
            debug_candidates_pil.save(temp_all_candidates_path)
            temp_debug_files.append(temp_all_candidates_path)  # Добавляем в список для удаления
            logger.info(f"Сохранено изображение со всеми кандидатами: {temp_all_candidates_path}")
            
            # Выбираем лучший кандидат
            best_cnt, x, y, w, h, area, method = all_candidates[0]
            logger.info(f"Выбран лучший кандидат: метод={method}, x={x}, y={y}, w={w}, h={h}, area={area}")
            
            # Извлекаем ROI с небольшим отступом внутрь (чтобы исключить саму рамку)
            padding_out = max(2, int(min(w, h) * 0.02))  # Отступ наружу (минимальный)
            padding_in = max(5, int(min(w, h) * 0.05))   # Отступ внутрь (чтобы убрать рамку)
            
            x_roi = max(0, x + padding_in)
            y_roi = max(0, y + padding_in)
            w_roi = min(cropped_image.shape[1] - x_roi, w - 2 * padding_in)
            h_roi = min(cropped_image.shape[0] - y_roi, h - 2 * padding_in)
            
            # Извлекаем ROI
            roi_image = cropped_image[y_roi:y_roi+h_roi, x_roi:x_roi+w_roi]
            
            if roi_image.size > 0:
                logger.debug(f"ROI извлечен. Размер: {roi_image.shape}")
                
                # Сохраняем ROI
                roi_pil = Image.fromarray(cv2.cvtColor(roi_image, cv2.COLOR_BGR2RGB))
                temp_roi_path = Path(working_materials_folder) / "temp_roi_image_DEBUG.jpg"
                roi_pil.save(temp_roi_path)
                temp_debug_files.append(temp_roi_path)  # Добавляем в список для удаления
                logger.info(f"Сохранено изображение ROI: {temp_roi_path}")
                
                # Применяем улучшенное OCR с несколькими попытками
                case_number = perform_enhanced_ocr(str(temp_roi_path), working_materials_folder)
                
                if case_number:
                    logger.info(f"Номер успешно распознан из ROI: '{case_number}'")
                    return case_number
                else:
                    logger.warning("Не удалось распознать номер из ROI")
            else:
                logger.warning("Извлеченный ROI оказался пустым")
        else:
            logger.info("Не найдено ни одного подходящего кандидата на рамку")
            
    except Exception as e:
        logger.error(f"Критическая ошибка при обработке изображения: {e}")
    finally:
        if delete_debug_files:
            # Удаление файлов как раньше
            for debug_file in temp_debug_files:
                try:
                    if os.path.exists(debug_file):
                        os.remove(debug_file)
                        logger.debug(f"Удалён отладочный файл: {debug_file}")
                except Exception as e:
                    logger.warning(f"Не удалось удалить отладочный файл {debug_file}: {e}")
            logger.info("Все отладочные файлы удалены")
        else:
            logger.info(f"Отладочные файлы сохранены в папке: {working_materials_folder}")
    
    # Fallback: OCR на полном изображении
    logger.info("Используем fallback - OCR на полном изображении")
    try:
        case_number = perform_enhanced_ocr(image_path, working_materials_folder)
        if case_number:
            logger.info(f"Номер распознан из полного изображения: '{case_number}'")
            return case_number
    except Exception as e:
        logger.error(f"Ошибка при fallback OCR: {e}")
    
    return None

def filter_rectangle_contours(contours, reference_image, min_area=300, method_name="unknown"):
    """
    Вспомогательная функция для фильтрации контуров и поиска прямоугольников.
    Возвращает список кортежей (контур, x, y, w, h, площадь, имя_метода).
    """
    candidates = []
    
    for cnt in contours:
        area = cv2.contourArea(cnt)
        
        # Фильтр по минимальной площади
        if area < min_area:
            continue
            
        # Аппроксимация контура
        epsilon = 0.02 * cv2.arcLength(cnt, True)
        approx = cv2.approxPolyDP(cnt, epsilon, True)
        
        # Проверка на 4 вершины (прямоугольник)
        if len(approx) == 4:
            x, y, w, h = cv2.boundingRect(cnt)
            aspect_ratio = float(w) / h
            
            # Фильтры по соотношению сторон и заполненности
            if 0.3 < aspect_ratio < 4.0:  # Расширили диапазон
                rect_area = w * h
                if rect_area > 0:
                    fill_ratio = area / rect_area
                    
                    # Проверка заполненности (прямоугольник должен быть достаточно "полным")
                    if fill_ratio > 0.7:  # Увеличили порог для лучшей фильтрации
                        # Дополнительная проверка: прямоугольник не должен быть слишком большим
                        image_area = reference_image.shape[0] * reference_image.shape[1]
                        if rect_area < image_area * 0.5:  # Не больше 50% от изображения
                            candidates.append((cnt, x, y, w, h, area, method_name))
                            logger.debug(f"Найден кандидат ({method_name}): x={x}, y={y}, w={w}, h={h}, "
                                       f"area={area}, aspect_ratio={aspect_ratio:.2f}, fill_ratio={fill_ratio:.2f}")
    
    return candidates

def perform_enhanced_ocr(image_path, working_materials_folder):
    """
    Улучшенная функция OCR с несколькими попытками и разными настройками.
    """
    try:
        # Предобработка изображения
        gray = preprocess_image_for_ocr(image_path)
        
        # Список конфигураций для Tesseract (разные PSM режимы)
        ocr_configs = [
            '--psm 6 --oem 3 -c tessedit_char_whitelist=0123456789-',  # Оригинальная
            '--psm 8 --oem 3 -c tessedit_char_whitelist=0123456789-',  # Одно слово
            '--psm 7 --oem 3 -c tessedit_char_whitelist=0123456789-',  # Одна строка
            '--psm 11 --oem 3 -c tessedit_char_whitelist=0123456789-', # Разреженный текст
        ]
        
        import re
        
        for i, config in enumerate(ocr_configs):
            try:
                logger.debug(f"OCR попытка {i+1} с конфигурацией: {config}")
                text = pytesseract.image_to_string(gray, config=config)
                logger.debug(f"Распознанный текст (попытка {i+1}): '{text.strip()}'")
                
                # Попробуем несколько паттернов для извлечения номера
                patterns = [
                    r'^(\d+-?\d*)',           # Оригинальный паттерн
                    r'(\d{4,})',               # Просто 4+ цифры подряд
                    r'(\d+)\s*-\s*(\d+)',      # Номер с тире и пробелами
                    r'(\d+)',                  # Любые цифры
                ]
                
                for pattern in patterns:
                    match = re.search(pattern, text)
                    if match:
                        if '-' in match.group(0):
                            case_number = match.group(0).split('-')[0]
                        else:
                            # Если найдено число длиной больше 2, убираем последние 2 цифры (год)
                            number = match.group(0)
                            if len(number) > 2:
                                case_number = number[:-2]
                            else:
                                case_number = number
                        
                        # Проверка валидности номера (должен быть не пустой и разумной длины)
                        if case_number and 2 <= len(case_number) <= 10:
                            logger.info(f"Номер извлечен с конфигурацией {i+1}: '{case_number}'")
                            return case_number
                            
            except Exception as e:
                logger.warning(f"Ошибка при OCR с конфигурацией {i+1}: {e}")
                
    except Exception as e:
        logger.error(f"Ошибка в perform_enhanced_ocr: {e}")
    
    return None

def extract_case_number_from_txt_filename(txt_path):
    """
    Извлекает номер заключения из имени файла .txt.
    """
    import re
    filename = txt_path.name
    match = re.search(r'^(\d+)', filename)
    if match:
        return match.group(1).strip()  # Добавляем strip() для удаления лишних пробелов
    return None

def extract_case_number_from_foregram(foregram_path):
    """
    Извлекает номер заключения из фореграммы с использованием OCR.
    """
    gray = preprocess_image_for_ocr(foregram_path)
    # Применяем OCR для извлечения текста
    text = pytesseract.image_to_string(gray, config='--psm 6 --oem 3 -c tessedit_char_whitelist=0123456789-')
    print(f"Распознанный текст из фореграммы: {text}")  # Отладочный вывод
    
    # Разделяем текст на строки
    lines = text.split('\n')
    
    # Ищем номер заключения в первой строке
    for line in lines:
        line = line.strip()
        if line:
            # Ищем номер заключения (например, число в начале строки)
            import re
            match = re.search(r'^(\d+-?\d*)', line)
            if match:
                case_number_with_year = match.group(1).strip()  # Добавляем strip() для удаления лишних пробелов
                # Убираем последние две цифры года или все символы после тире
                if '-' in case_number_with_year:
                    case_number = case_number_with_year.split('-')[0]
                else:
                    case_number = case_number_with_year[:-2]
                return case_number
            break  # Прерываем цикл после первой строки
    
    return None

def check_case_number(card_data, txt_path, image_path, foregram_path, working_materials_folder):
    """
    Проверяет консистентность номера заключения в Карте, фотографии, фореграмме и файле .txt.
    Для шаблонов ЭКЦ, ЭКЦ_, Тобольск_, СВО_Ростов анализируется всё изображение целиком, 
    а не только верхний левый квадрант.
    """
    # Извлекаем номера
    case_number_from_card_raw = card_data.get("НОМ")
    
    # Определяем, нужно ли анализировать всё изображение
    template_value = card_data.get("1", "")
    # Приводим к строке и удаляем лишние пробелы для надёжности
    template_value_str = str(template_value).strip() if template_value is not None else ""
    
    # Логирование исходного значения шаблона
    logger.info(f"[OCR_MODE] Шаблон заключения (сырое значение): '{template_value}'")
    logger.info(f"[OCR_MODE] Шаблон заключения (очищенный): '{template_value_str}'")
    
    # Новое правило: полное изображение, если шаблон == "ЭКЦ" или начинается с префиксов
    use_full = False
    if template_value_str:
        if (template_value_str == "ЭКЦ" or
            template_value_str.startswith("ЭКЦ_") or
            template_value_str.startswith("Тобольск_") or
            template_value_str.startswith("СВО_Ростов_")):
            use_full = True
            logger.info("[OCR_MODE] Обнаружен шаблон, требующий анализа ВСЕГО изображения.")
        else:
            logger.info("[OCR_MODE] Шаблон не соответствует условиям полного анализа. "
                        "Будет использован только верхний левый квадрант.")
    else:
        logger.warning("[OCR_MODE] Значение шаблона (поле '1') пустое или отсутствует. "
                       "По умолчанию анализируем только верхний левый квадрант.")
    
    logger.info(f"[OCR_MODE] Итоговое решение: анализировать всё изображение? {'ДА' if use_full else 'НЕТ'}")
    
    # Передаем working_materials_folder и новый флаг use_full_image
    case_number_from_image_raw = extract_case_number_from_image(
        image_path, 
        working_materials_folder,
        delete_debug_files=True,      # True - удалять отладку
        use_full_image=use_full
    )
    
    case_number_from_foregram_raw = extract_case_number_from_foregram(foregram_path)
    case_number_from_txt_raw = extract_case_number_from_txt_filename(txt_path)
    
    # Отладочный вывод "сырых" значений
    logger.debug(f"[DEBUG] Номер из Карты (сырой): '{case_number_from_card_raw}' (тип: {type(case_number_from_card_raw)})")
    logger.debug(f"[DEBUG] Номер из Изображения (сырой): '{case_number_from_image_raw}' (тип: {type(case_number_from_image_raw)})")
    logger.debug(f"[DEBUG] Номер из Фореграммы (сырой): '{case_number_from_foregram_raw}' (тип: {type(case_number_from_foregram_raw)})")
    logger.debug(f"[DEBUG] Номер из .txt (сырой): '{case_number_from_txt_raw}' (тип: {type(case_number_from_txt_raw)})")
    
    # Очищаем номера от пробелов и других невидимых символов в начале/конце
    case_number_from_card = str(case_number_from_card_raw).strip() if case_number_from_card_raw is not None else None
    case_number_from_image = str(case_number_from_image_raw).strip() if case_number_from_image_raw is not None else None
    case_number_from_foregram = str(case_number_from_foregram_raw).strip() if case_number_from_foregram_raw is not None else None
    case_number_from_txt = str(case_number_from_txt_raw).strip() if case_number_from_txt_raw is not None else None
    
    # Отладочный вывод "очищенных" значений
    logger.debug(f"[DEBUG] Номер из Карты (очищенный): '{case_number_from_card}'")
    logger.debug(f"[DEBUG] Номер из Изображения (очищенный): '{case_number_from_image}'")
    logger.debug(f"[DEBUG] Номер из Фореграммы (очищенный): '{case_number_from_foregram}'")
    logger.debug(f"[DEBUG] Номер из .txt (очищенный): '{case_number_from_txt}'")
    
    # Сравниваем очищенные номера
    if (case_number_from_card != case_number_from_image or
        case_number_from_card != case_number_from_foregram or
        case_number_from_card != case_number_from_txt):
        error_msg = (f"Несоответствие номеров заключений (после очистки): "
                     f"в Карте - '{case_number_from_card}', "
                     f"в фотографии - '{case_number_from_image}', "
                     f"в фореграмме - '{case_number_from_foregram}', "
                     f"в файле .txt - '{case_number_from_txt}'.")
        logger.error(f"[ERROR] {error_msg}")
        raise ValueError(error_msg)
    
    logger.info("[INFO] Проверка номеров заключений прошла успешно.")

def check_txt_filename_consistency(txt_path, profiles_data, case_number):
    """
    Проверяет, соответствует ли имя файла .txt хотя бы одному профилю в данных.
    Возвращает True, если есть точное или гибкое совпадение, иначе False.
    При наличии нескольких соответствий выводит предупреждение.
    Теперь учитывает только профили, оканчивающиеся на -26 (игнорирует профиль образца).
    """
    logger.info("Начало проверки соответствия названия файла txt и данных профилей")
    
    # Оставляем только профили, оканчивающиеся на -26
    # Это исключает профиль образца (если он добавлен) из проверки
    filtered_profiles = {name: data for name, data in profiles_data.items() if name.endswith("-26")}
    
    if not filtered_profiles:
        error_msg = "Нет ни одного профиля, оканчивающегося на -26, для проверки соответствия имени файла"
        logger.error(error_msg)
        raise ValueError(error_msg)
    
    logger.info(f"Для проверки отобраны профили с суффиксом -26: {list(filtered_profiles.keys())}")
    
    # Получаем оригинальное название файла без расширения и части "Genotypes Table"
    filename = txt_path.name
    base_filename = filename.replace(".txt", "").replace("Genotypes Table", "").strip()
    logger.info(f"Название файла (без расширения и 'Genotypes Table'): {base_filename}")
    
    # Извлекаем номер заключения из имени файла (все цифры в начале строки)
    match = re.match(r'^\d+', base_filename)
    if not match:
        error_msg = "Номер заключения в имени файла должен начинаться с цифр"
        logger.error(error_msg)
        raise ValueError(error_msg)
    
    file_case_number = match.group(0)
    logger.info(f"Номер заключения из имени файла: {file_case_number}")
    
    # Проверяем соответствие номера из файла и номера из карты
    if file_case_number != case_number:
        error_msg = f"Несоответствие номеров: в имени файла '{file_case_number}', в карте '{case_number}'"
        logger.error(error_msg)
        raise ValueError(error_msg)
    
    # Ищем часть строки после номера заключения до конца
    suffix_match = re.search(rf'^\d+(.*)$', base_filename)
    suffix_pattern = suffix_match.group(1) if suffix_match else ""
    
    # Строим регулярное выражение для поиска совпадений
    exact_pattern = re.compile(rf'^{re.escape(base_filename)}$')
    flexible_pattern = re.compile(rf'^{re.escape(file_case_number)}.*?-26$')
    
    # Проверяем все отфильтрованные профили на соответствие
    exact_matches = []
    flexible_matches = []
    
    for profile_name in filtered_profiles.keys():
        if exact_pattern.match(profile_name):
            exact_matches.append(profile_name)
        elif flexible_pattern.match(profile_name):
            flexible_matches.append(profile_name)
    
    # Определяем результат проверки
    if exact_matches:
        logger.info(f"Найдено точное совпадение: {exact_matches}")
        if len(exact_matches) > 1:
            logger.warning(f"Найдено несколько точных совпадений: {exact_matches}")
            return False  # Несколько точных совпадений - требуется подтверждение
        return True  # Одно точное совпадение - проверка пройдена
    
    elif flexible_matches:
        logger.info(f"Найдено гибкое совпадение: {flexible_matches}")
        if len(flexible_matches) > 1:
            logger.warning(f"Найдено несколько гибких совпадений: {flexible_matches}")
            return False  # Несколько гибких совпадений - требуется подтверждение
        return True  # Одно гибкое совпадение - проверка пройдена
    
    else:
        error_msg = "В данных файла txt нет профиля, соответствующего имени файла"
        logger.error(error_msg)
        raise ValueError(error_msg)

def check_consistency(card_data, profiles_data, image_path, foregram_path, txt_path, working_materials_folder):
    """
    Проверяет консистентность данных в карте, данных локусов и изображениях
    с обработкой отсутствующей фореграммы.
    """
    # Получаем номер заключения из карты
    nom_value = card_data.get("НОМ", "")
    case_number = str(nom_value).strip() if nom_value is not None else ""
    logger.info(f"Номер заключения из карты: {case_number}")
    
    # Оставляем только профили, оканчивающиеся на -26
    # Это исключает профиль образца (если он добавлен)
    filtered_profiles = {name: data for name, data in profiles_data.items() if name.endswith("-26")}
    if not filtered_profiles:
        error_msg = f"Нет ни одного профиля, оканчивающегося на -26, для проверки консистентности"
        logger.error(error_msg)
        raise ValueError(error_msg)
    logger.info(f"Для проверки консистентности отобраны профили с суффиксом -26: {list(filtered_profiles.keys())}")
    
    # Находим профили, которые соответствуют номеру заключения (среди отфильтрованных)
    matching_profiles = []
    
    for profile_name in filtered_profiles.keys():
        # Проверяем точное соответствие (профиль должен быть равен номеру заключения)
        if profile_name == case_number:
            matching_profiles.append(profile_name)
            logger.info(f"Найден профиль с точным соответствием: {profile_name}")

    # Если точных соответствий нет, ищем гибкие соответствия
    if not matching_profiles:
        for profile_name in filtered_profiles.keys():
            if profile_name.startswith(case_number) and profile_name.endswith("-26"):
                matching_profiles.append(profile_name)
                logger.info(f"Найден профиль с гибким соответствием: {profile_name}")
    
    # Проверяем результаты
    if not matching_profiles:
        error_msg = f"Не найден профиль, соответствующий номеру заключения {case_number}"
        logger.error(error_msg)
        raise ValueError(error_msg)
    
    # Если найдено несколько соответствий, запрашиваем подтверждение
    if len(matching_profiles) > 1:
        logger.warning(f"Найдено несколько профилей, соответствующих номеру заключения: {matching_profiles}")
        root = tk.Tk()
        root.withdraw()
        response = messagebox.askyesno(
            "Предупреждение", 
            f"Найдено несколько профилей, соответствующих номеру заключения {case_number}:\n" +
            "\n".join(matching_profiles) +
            "\n\nПроверьте корректность данных и нажмите 'Да' для продолжения или 'Нет' для отмены."
        )
        root.destroy()
        
        if not response:
            raise ValueError("Пользователь отменил выполнение из-за несоответствия данных")
        
        # Используем первый профиль из списка соответствий
        selected_profile = matching_profiles[0]
        logger.info(f"Для проверки пола выбран профиль: {selected_profile}")
    else:
        selected_profile = matching_profiles[0]
        logger.info(f"Для проверки пола используется профиль: {selected_profile}")
    
    # Проверка фореграммы
    if foregram_path is None or not foregram_path.exists():
        if (card_data.get("1") != "Кость СВО нет пригодной ДНК") and (card_data.get("1") != "СВО_кость_нет результата"):
            raise FileNotFoundError("Фореграмма отсутствует, но требуется по условиям")
        return  # Пропускаем проверку
    
    # Проверка пола с использованием выбранного профиля
    check_gender(card_data, profiles_data[selected_profile])
    
    # Проверка номера заключения
    check_case_number(card_data, txt_path, image_path, foregram_path, working_materials_folder)

def find_excel_file():
    """
    Ищет единственный Excel-файл в текущей директории.
    Возвращает путь к файлу или вызывает исключение, если файл не найден или их несколько.
    """
    # Получаем список всех Excel-файлов в текущей директории
    excel_files = list(Path('.').glob('*.xlsx')) + list(Path('.').glob('*.xls'))
    
    # Проверяем количество найденных файлов
    if not excel_files:
        raise FileNotFoundError("В текущей директории не найдено Excel-файлов (.xlsx или .xls)")
    if len(excel_files) > 1:
        raise ValueError(f"В директории найдено несколько Excel-файлов: {[f.name for f in excel_files]}. "
                        "Должен быть только один файл.")
    
    return excel_files[0]

# Основная функция
def main():
    try:
        logger.info("Запуск программы")

        # === ТЕСТИРОВАНИЕ ДЕКОРАТОРА ===
        #test_win32com_error_decorator()
        # ================================================

        # Находим Excel-файл
        card_path = find_excel_file()
        logger.info(f"Найден файл: {card_path}")

        # Чтение данных из найденного файла
        card_data = read_card(card_path)
        logger.info("Данные из файла: %s", card_data)

        # Проверяем, что все обязательные поля заполнены
        check_card_data(card_data)

        # Загрузка словаря замен
        templates_folder = Path(card_data["3"])
        switch_dict = load_switch_dictionary(templates_folder)
        if not switch_dict:
            logger.warning("Словарь замен не загружен или пуст")
        else:
            logger.info("Словарь замен успешно загружен")

        # Копирование шаблона Word
        output_path = copy_template(card_data)
        logger.info(f"Шаблон скопирован в: {output_path}")

        # Чтение абзацев из txt-файла
        template_folder = card_data["3"]  # Путь к папке с шаблонами (строка Карты с кодом 3)
        paragraphs = read_paragraphs_from_word(template_folder)
        logger.info(f"Загружены фрагменты для ключей: {list(paragraphs.keys())}")

        # Предварительная обработка абзацев
        preprocess_paragraphs(output_path, paragraphs, card_data)
        logger.info("Абзацы предварительно обработаны.")

        # Обработка таблицы RT
        process_rt_table(output_path, card_data)
        logger.info("Таблица RT обработана.")

        # Создаем словарь со всеми значениями из Карты
        replacements = {
            code: str(value) if value is not None else ''
            for code, value in card_data.items()
        }
        logger.info("Метки для замены: %s", replacements)

        # Генерация динамических блоков ПАН
        try:
            generate_dynamic_pan_blocks(output_path, card_data)
        except Exception as e:
            logger.error(f"Ошибка при генерации динамических блоков ПАН: {e}")
            show_error_message(f"Ошибка при генерации блоков ПАН: {e}")
            sys.exit(1)

        # Обрабатываем switch-метки
        if switch_dict:
            process_switch_tags(output_path, card_data, switch_dict)
        else:
            logger.warning("Пропуск обработки switch-меток: словарь замен не загружен")

        # Заменяем метки в документе
        replace_in_doc(output_path, replacements)
        logger.info("Метки заменены.")

        # Получаем путь к папке с рабочими материалами из Карты (поле с кодом "4")
        base_working_materials_folder = Path(card_data["4"])
        nom_folder_name = f"{card_data['НОМ']}-26"
        working_materials_folder = base_working_materials_folder / nom_folder_name

        # Проверяем, существует ли папка
        if not working_materials_folder.exists():
            raise FileNotFoundError(f"Папка '{working_materials_folder}' не существует.")
        
        # Вставка изображений
        insert_images(output_path, working_materials_folder, card_data)
        logger.info("Изображения вставлены.")

        # Находим ВСЕ обязательные изображения (1, 1_2, 1_3)
        if (card_data.get("1") != "Кость СВО нет пригодной ДНК") and (card_data.get("1") != "СВО_кость_нет результата"):
        #if card_data.get("1") != "СВО_кость_нет результата":
            required_images = ["1", "1_2", "1_3"]
            missing_images = []
            found_images = {}

            for img_name in required_images:
                img_path = None
                for ext in [".jpg", ".jpeg", ".png"]:
                    temp_path = working_materials_folder / f"{img_name}{ext}"
                    if temp_path.exists():
                        img_path = temp_path
                        break
                
                if img_path:
                    found_images[img_name] = img_path
                else:
                    missing_images.append(img_name)

            if missing_images:
                error_msg = f"Не найдены обязательные изображения: {', '.join(missing_images)}"
                logger.error(error_msg)
                show_error_message(error_msg)
                sys.exit(1)
            else:
                logger.info(f"Все специальные изображения найдены: {list(found_images.keys())}")
                foregram_path = found_images["1"]

        # Находим файл .txt в целевой папке
        txt_path = find_txt_file(working_materials_folder)

        # Объявляем переменную для данных профилей
        profiles_data = None

        # Проверяем, нужно ли искать и обрабатывать файл .txt
        if card_data.get("1") in [
            "СВО_кость_нет результата",
            "СВО_Молов_образец_родственники_нет результата_RT", "СВО_Молов_образец_родственники_нет результата_форез",
            "СВО_Ростов_образец_прямая идентификация_нет результата_RT", "СВО_Ростов_образец_прямая идентификация_нет результата_форез",
            "СВО_Ростов_образец_родственники_нет результата_RT", "СВО_Ростов_образец_родственники_нет результата_форез",
            "ЭКЦ_нет результата"
        ]:
            logger.info("В экспертизе нет генотипа - пропускаем обработку файла .txt")
            profiles_data = None
        else:
            if txt_path is None:
                raise FileNotFoundError("Файл .txt с генотипами не найден, и это не соответствует условиям для пропуска.")
            logger.info(f"Найден файл .txt: {txt_path}")
            
            # Обработка специальных шаблонов: ЭКЦ и Тобольск
            temp_txt_path = txt_path
            should_stop = False
            profiles_data = None
            custom_headers = {}
            matching_profile_name = None

            template_value = card_data.get("1", "").strip()
            no_genotype_templates = [
                "СВО_кость_нет результата",
                "СВО_Молов_образец_родственники_нет результата_RT", "СВО_Молов_образец_родственники_нет результата_форез", 
                "СВО_Ростов_образец_прямая идентификация_нет результата_RT", "СВО_Ростов_образец_прямая идентификация_нет результата_форез",
                "СВО_Ростов_образец_родственники_нет результата_RT", "СВО_Ростов_образец_родственники_нет результата_форез",
                "ЭКЦ_нет результата"
            ]

            if template_value in no_genotype_templates:
                logger.info("В экспертизе нет генотипа - пропускаем обработку файла .txt")
                profiles_data = None
            else:
                # Сначала проверяем шаблон Тобольск_образец_мужчина
                if template_value == "Тобольск_образец_мужчина":
                    logger.info("Обнаружен шаблон Тобольск_образец_мужчина – запускаем специальную обработку без сравнения")
                    success, temp_txt_path, profiles_data, custom_headers, ref_profile_name = process_tobolsk_sample(
                        card_data, working_materials_folder, txt_path
                    )
                    if not success:
                        raise ValueError("Ошибка при обработке Тобольск_образец_мужчина")
                    should_stop = False
                    matching_profile_name = None
                    # Определяем данные образца для выделения
                    if ref_profile_name and ref_profile_name in profiles_data:
                        reference_profile_data = profiles_data[ref_profile_name]
                        logger.info(f"Будет выполнено выделение несовпадающих аллелей относительно образца {ref_profile_name}")
                    else:
                        reference_profile_data = None
                        logger.warning("Не удалось определить профиль образца – выделение отключено")
                elif template_value.startswith("ЭКЦ_образец_") or template_value == "Тобольск_образец_женщина":
                    logger.info(f"Проверка на наличие образца для сравнения (шаблон: {template_value})...")
                    success, temp_txt_path, should_stop, profiles_data, custom_headers, matching_profile_name = process_ekc_sample_comparison(
                        card_data, working_materials_folder, txt_path
                    )
                    if not success:
                        raise ValueError("Ошибка при обработке ЭКЦ образца")
                    reference_profile_data = None
                    ref_profile_name = None
                else:
                    # Обычный случай – читаем данные как есть (только профили -26)
                    logger.info("Обычный режим: чтение данных из txt-файла")
                    profiles_data = read_and_filter_txt_data(txt_path)
                    reference_profile_data = None
                    ref_profile_name = None

            logger.info(f"Найдены профили для таблицы: {list(profiles_data.keys())}")
            logger.info(f"Кастомные заголовки: {custom_headers}")

            # Дополнительная проверка для шаблонов, требующих только один профиль
            template_value = card_data.get("1", "").strip()
            single_profile_templates = [
                "СВО_Молов_образец_родственники", 
                "СВО_Ростов_образец_родственники", 
                "СВО_Молов_образец_прямая идентификация",
                "СВО_Молов БВП_образец_прямая идентификация",
                "СВО_Ростов_образец_прямая идентификация", 
                "СВО_кость",
                "ЭКЦ"
            ]

            if template_value in single_profile_templates:
                if len(profiles_data) > 1:
                    error_msg = (f"Для шаблона '{template_value}' допускается только один профиль, "
                                f"но найдено {len(profiles_data)} профилей: {list(profiles_data.keys())}")
                    logger.error(error_msg)
                    show_error_message(error_msg)
                    sys.exit(1)
                elif len(profiles_data) == 0:
                    error_msg = f"Для шаблона '{template_value}' не найдено ни одного профиля"
                    logger.error(error_msg)
                    show_error_message(error_msg)
                    sys.exit(1)

            # Определяем, тобольский ли это шаблон
            is_tobolsk = (card_data.get("1") == "Тобольск_образец_мужчина")
            
            # Проверка на аллели OL в профилях, которые попадут в таблицу
            if profiles_data is not None:
                check_ol_alleles_in_profiles(profiles_data)

            # Вставка данных в таблицу с передачей кастомных заголовков
            # Определяем, нужно ли пропустить комментарий после таблицы
            skip_comment_flag = card_data.get("1", "").strip() in ["Тобольск_образец_мужчина", "Тобольск_образец_женщина"]
            insert_table_data(output_path, profiles_data, card_data, custom_headers, 
                              skip_comment=skip_comment_flag,
                              reference_profile_data=reference_profile_data,
                              reference_profile_name=ref_profile_name,
                              is_tobolsk=is_tobolsk)

            logger.info("Данные вставлены в таблицу.")

            # Удаление временного файла если он был создан
            if temp_txt_path != txt_path and temp_txt_path.exists():
                try:
                    temp_txt_path.unlink()
                    logger.info(f"Временный файл удален: {temp_txt_path}")
                except Exception as e:
                    logger.warning(f"Не удалось удалить временный файл {temp_txt_path}: {e}")

            # Проверяем should_stop после заполнения таблицы
            if should_stop:
                # Выводим сообщение и останавливаем программу после заполнения таблицы
                error_msg = f"ОБНАРУЖЕНО СОВПАДЕНИЕ ПРОФИЛЕЙ! Профиль '{matching_profile_name}' совпадает с образцом ЭКЦ. Поздравляю - дело раскрыто!"
                logger.error(error_msg)
                show_error_message(error_msg)
                # Не завершаем программу сразу, даем пользователю увидеть сообщение
                input("Нажмите Enter для завершения программы...")
                sys.exit(1)

        # Размещение текста по метке [WRAP_TEXT], max_line_length регулируется
        logger.info("Обработка переноса текста в таблицах с маркерами...")
        try:
            process_table_cell_wrapping(output_path, max_line_length=85)
            logger.info("Перенос текста в помеченных ячейках выполнен успешно")
        except Exception as e:
            logger.error(f"Ошибка при обработке переноса текста: {e}")

        # Проверка консистентности данных
        if card_data.get("1") in [
            "СВО_кость_нет результата",
            "СВО_Молов_образец_родственники_нет результата_RT", "СВО_Молов_образец_родственники_нет результата_форез",
            "СВО_Ростов_образец_прямая идентификация_нет результата_RT", "СВО_Ростов_образец_прямая идентификация_нет результата_форез",
            "СВО_Ростов_образец_родственники_нет результата_RT", "СВО_Ростов_образец_родственники_нет результата_форез",
            "ЭКЦ_нет результата"
        ]:
            logger.info("Пропуск проверки консистентности для экспертиз без генотипа")
        else:
            try:
                if txt_path is None:
                    raise FileNotFoundError("Файл .txt с генотипами не найден!")
                if foregram_path is None:
                    raise FileNotFoundError("Фореграмма (изображение 1) не найдена!")
                
                # Поиск обычного изображения
                ordinary_image_path = None
                special_labels = {"1", "1_2", "1_3", "1_4", "1_5", "1_6", "1_7", "1_8", "1_9", "1_10", "1_11", "1_12"}
                for ext in [".jpg", ".jpeg", ".png"]:
                    for img_candidate in working_materials_folder.glob(f"*{ext}"):
                        stem = img_candidate.stem
                        if stem not in special_labels:
                            ordinary_image_path = img_candidate
                            logger.info(f"Найдено обычное изображение: {ordinary_image_path.name}")
                            break
                    if ordinary_image_path is not None:
                        break

                if ordinary_image_path is None:
                    raise FileNotFoundError("Обычное изображение не найдено для проверки OCR.")

                # Проверка соответствия имени файла и данных профилей
                nom_value = card_data.get("НОМ", "")
                case_number = str(nom_value).strip() if nom_value is not None else ""
                consistency_check = check_txt_filename_consistency(txt_path, profiles_data, case_number)
                
                # Если проверка не пройдена (несколько совпадений), запрашиваем подтверждение
                if not consistency_check:
                    logger.warning("Обнаружено несколько профилей, соответствующих имени файла")
                    root = tk.Tk()
                    root.withdraw()
                    response = messagebox.askyesno(
                        "Предупреждение", 
                        "Обнаружено несколько профилей, соответствующих имени файла.\n\n"
                        "Проверьте корректность данных и нажмите 'Да' для продолжения или 'Нет' для отмены."
                    )
                    root.destroy()
                    
                    if not response:
                        raise ValueError("Пользователь отменил выполнение из-за несоответствия данных")

                # Проверка консистентности данных (передаём все профили)
                check_consistency(card_data, profiles_data, ordinary_image_path, foregram_path, txt_path, working_materials_folder)
                print("Проверка консистентности прошла успешно.")
            except Exception as e:
                logger.error(f"ОШИБКА ПРОВЕРКИ: {e}")
                show_error_message(f"ОШИБКА ПРОВЕРКИ: {e}")
                sys.exit(1)

        logger.info("Скрипт успешно завершён!")
        
    except FileNotFoundError as e:
        logger.error(f"Ошибка: {e}")
        show_error_message(f"Ошибка: {e}")
    except Exception as e:
        logger.error(f"Произошла ошибка: {e}")
        show_error_message(f"Произошла ошибка: {e}")
        
# Запуск скрипта
if __name__ == "__main__":
    main()