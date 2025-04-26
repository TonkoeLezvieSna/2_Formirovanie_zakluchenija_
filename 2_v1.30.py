#2
#v1.30

# Импортируем необходимые модули
import shutil
import logging
import cv2
import datetime
import sys
import pytesseract
import re
import json
import docx.text
import numpy as np
import tkinter as tk
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from openpyxl import load_workbook
from PIL import Image, ImageEnhance
from tkinter import messagebox
from copy import deepcopy

# Указываем путь к tesseract
pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'

# Настройка логирования (вывод только в консоль)
logging.basicConfig(level=logging.INFO, format="%(asctime)s - %(levelname)s - %(message)s")
logger = logging.getLogger(__name__)

# Определяем порядок сортировки локусов (в верхнем регистре)
LOCUS_ORDER = [
    "D3S1358", "VWA", "D16S539", "CSF1PO", "TPOX", "YINDEL", "AMEL", 
    "D8S1179", "D21S11", "D18S51", "DYS391", "D2S441", "D19S433", 
    "TH01", "FGA", "D22S1045", "D5S818", "D13S317", "D7S820", "SE33", 
    "D10S1248", "D1S1656", "D12S391", "D2S1338", "SRY", "PENTA E",
    "D6S1043", "RS759551978", "RS771783753", "RS199815934", "PENTA D", "D4S2366"
]

# Показываем сообщение об ошибке в графическом окне и логируем её
def show_error_message(message):
    logger.error(message)
    root = tk.Tk()
    root.withdraw()  # Скрываем главное окно Tkinter
    messagebox.showerror("Ошибка", message)
    root.destroy()

# Загружаем словарь замен из JSON-файла в указанной папке
def load_switch_dictionary(templates_folder):
    """Загружает словарь замен из JSON-файла в указанной папке."""
    dictionary_path = Path(templates_folder) / "switch_dictionary.json"
    try:
        if not dictionary_path.exists():
            logger.warning(f"Файл словаря не найден: {dictionary_path}")
            return None
        
        with open(dictionary_path, 'r', encoding='utf-8') as f:
            data = json.load(f)
            if not data:
                logger.warning("Словарь замен пуст")
            else:
                logger.info(f"Загружен словарь замен. Ключи: {list(data.keys())}")
            return data
            
    except json.JSONDecodeError as e:
        logger.error(f"Ошибка при чтении JSON-файла: {str(e)}")
        return None
    except Exception as e:
        logger.error(f"Неожиданная ошибка при загрузке словаря: {str(e)}")
        return None

# Функции для работы с изображениями
def resize_image(image, max_size=1200):
    """Уменьшает размер изображения до максимального размера с сохранением качества."""
    logger.info(f"Resizing image to max size {max_size}")
    height, width = image.shape[:2]
    if max(height, width) > max_size:
        scale = max_size / max(height, width)
        new_width = int(width * scale)
        new_height = int(height * scale)
        return cv2.resize(image, (new_width, new_height), interpolation=cv2.INTER_LINEAR)
    return image

def increase_contrast(image):
    """Увеличивает контрастность изображения без потери деталей."""
    logger.info("Increasing image contrast")
    lab = cv2.cvtColor(image, cv2.COLOR_BGR2LAB)
    l_channel, a, b = cv2.split(lab)
    clahe = cv2.createCLAHE(clipLimit=3.0, tileGridSize=(8, 8))
    cl = clahe.apply(l_channel)
    enhanced_lab = cv2.merge((cl, a, b))
    return cv2.cvtColor(enhanced_lab, cv2.COLOR_LAB2BGR)

def auto_crop_image(image_path, debug=False):
    """
    Автоматическая обработка и обрезка изображения.
    """
    logger.info(f"Auto cropping image: {image_path}")
    try:
        # Загружаем изображение через Pillow (поддерживает кириллицу в пути)
        image_pil = Image.open(image_path)
        # Конвертируем в numpy array для работы с OpenCV
        image = np.array(image_pil)
        image = cv2.cvtColor(image, cv2.COLOR_RGB2BGR)  # PIL -> OpenCV (RGB -> BGR)
    except Exception as e:
        logger.error(f"Ошибка: не удалось загрузить изображение. Проверьте путь к файлу: {e}")
        return None

    # Увеличение контрастности
    image = increase_contrast(image)
    # Уменьшение размера изображения
    image = resize_image(image)
    original_height, original_width = image.shape[:2]
    logger.info(f"Размер обработанного изображения: {original_width}x{original_height}")

    # Преобразование в градации серого
    gray = cv2.cvtColor(image, cv2.COLOR_BGR2GRAY)
    # Размытие для удаления шума
    gray = cv2.GaussianBlur(gray, (5, 5), 0)

    # Обнаружение границ с помощью Canny
    logger.info("Применение детектора границ Canny...")
    edges = cv2.Canny(gray, 30, 100)

    # Поиск контуров
    logger.info("Поиск контуров...")
    contours, _ = cv2.findContours(edges, cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_SIMPLE)
    if len(contours) == 0:
        logger.error("Ошибка: контуры не найдены.")
        return None

    # Фильтрация маленьких контуров
    filtered_contours = [cnt for cnt in contours if cv2.contourArea(cnt) > 100]
    if len(filtered_contours) == 0:
        logger.error("Ошибка: после фильтрации контуры отсутствуют.")
        return None

    # Выбор ограничивающих прямоугольников
    bounding_rects = [cv2.boundingRect(cnt) for cnt in filtered_contours]
    rects = [(x, y, w, h) for x, y, w, h in bounding_rects if 0.3 < w / h < 3]
    if len(rects) == 0:
        logger.error("Ошибка: не найдено подходящих прямоугольников.")
        return None

    # Отладочная информация
    if debug:
        logger.debug(f"Найдено {len(contours)} контуров.")
        for i, cnt in enumerate(filtered_contours):
            area = cv2.contourArea(cnt)
            x, y, w, h = cv2.boundingRect(cnt)
            logger.debug(f"Контур {i}: площадь={area}, размер=({w}x{h}), соотношение сторон={w/h:.2f}")

    # Объединение контуров
    all_contours = np.concatenate(filtered_contours)
    x, y, w, h = cv2.boundingRect(all_contours)

    # Обрезка изображения
    cropped_image = image[y:y + h, x:x + w]
    logger.info(f"Обрезанное изображение: {w}x{h}")

    if debug:
        cv2.imshow("Автоматическая обрезка", cropped_image)
        cv2.waitKey(0)
        cv2.destroyAllWindows()

    return cropped_image

def interactive_crop(image):
    """Позволяет пользователю выбрать область обрезки."""
    print("Выберите область обрезки (нажмите Enter, когда закончите).")
    roi = cv2.selectROI("Интерактивная обрезка", image, fromCenter=False)
    cv2.destroyWindow("Интерактивная обрезка")
    if sum(roi) == 0:  # Проверка, была ли выбрана область
        print("Область не выбрана. Используется автоматическая обрезка.")
        return None
    x, y, w, h = roi
    cropped_image = image[int(y):int(y+h), int(x):int(x+w)]
    return cropped_image

def adjust_brightness(image, brightness_factor):
    """Увеличивает яркость изображения."""
    enhancer = ImageEnhance.Brightness(image)
    return enhancer.enhance(brightness_factor)

def adjust_contrast(image, contrast_factor):
    """Увеличивает контрастность изображения."""
    enhancer = ImageEnhance.Contrast(image)
    return enhancer.enhance(contrast_factor)

def load_image_with_pillow(image_path):
    """
    Загружает изображение с помощью Pillow и конвертирует его в numpy array.
    """
    try:
        image = Image.open(image_path)
        return np.array(image)  # Конвертируем в numpy array
    except Exception as e:
        raise ValueError(f"Не удалось загрузить изображение: {image_path}. Ошибка: {e}")

def auto_adjust_brightness_contrast(image_path, brightness_factor=1.5, contrast_factor=1.5):
    """
    Автоматически настраивает яркость и контрастность изображения.
    """
    # Загружаем изображение с помощью Pillow
    image_np = load_image_with_pillow(image_path)
    # Преобразуем из RGB (Pillow) в BGR (OpenCV)
    image_bgr = cv2.cvtColor(image_np, cv2.COLOR_RGB2BGR)
    # Разделяем изображение на каналы (B, G, R)
    channels = cv2.split(image_bgr)
    # Нормализуем каждый канал отдельно
    normalized_channels = []
    for channel in channels:
        normalized = cv2.normalize(channel, None, alpha=0, beta=255, norm_type=cv2.NORM_MINMAX)
        normalized_channels.append(normalized)
    # Объединяем каналы обратно в цветное изображение
    result_bgr = cv2.merge(normalized_channels)
    # Преобразуем обратно в RGB
    result_rgb = cv2.cvtColor(result_bgr, cv2.COLOR_BGR2RGB)
    # Конвертируем в формат PIL для совместимости с docx
    result_image = Image.fromarray(result_rgb)
    # Увеличиваем яркость и контрастность
    result_image = adjust_brightness(result_image, brightness_factor)
    result_image = adjust_contrast(result_image, contrast_factor)
    return result_image

# Функция для сортировки данных по локусам
def sort_data_by_locus(data):
    """
    Сортирует данные по локусам в порядке, определённом в LOCUS_ORDER.
    Локусы, не указанные в LOCUS_ORDER, сортируются по алфавиту.
    При этом сохраняется оригинальный регистр названий локусов.
    """
    def get_locus_index(locus):
        # Приводим локус к верхнему регистру для сравнения
        locus_upper = locus.strip().upper()
        try:
            return LOCUS_ORDER.index(locus_upper)
        except ValueError:
            # Если локус не найден в LOCUS_ORDER, возвращаем большое число,
            # чтобы он оказался в конце списка
            return len(LOCUS_ORDER) + ord(locus_upper[0])  # Сортируем по алфавиту
    
    # Сортируем данные по локусам (без учёта регистра), но сохраняем оригинальный регистр
    return sorted(data, key=lambda x: get_locus_index(x[3]))

# Функция для проверки обязательных полей
def check_card_data(card_data):
    required_fields = ["НОМ", "1", "2", "3", "4", "5", "6", "7"]  # Список обязательных полей
    for field in required_fields:
        if field not in card_data or card_data[field] is None:
            raise ValueError(f"Ошибка: в Карте отсутствует значение для кода '{field}'.")

def read_card(file_path):
    """Читает данные из Excel-файла и возвращает их в виде словаря."""
    logger.info(f"Чтение данных из Excel-файла: {file_path}")
    data = {}
    try:
        wb = load_workbook(file_path)
        sheet = wb.active
        for row in sheet.iter_rows(min_row=2, values_only=True):
            code = str(row[0]).strip()  # Преобразуем код в строку и убираем лишние пробелы
            value = row[2]
            if code == "ДБИ" and value is not None and value.strip():
                value = f"с {value.strip()} г."
            data[code] = value
        logger.info("Данные успешно загружены из Excel-файла.")
    except Exception as e:
        logger.error(f"Ошибка при чтении Excel-файла: {e}")
        raise
    return data

def copy_template(card_data):
    """Копирует шаблон Word на основе данных из карты."""
    logger.info("Копирование шаблона Word...")
    template_folder = Path(card_data["3"])  # Путь к папке с шаблонами
    template_name = card_data["1"]  # Название шаблона из строки с кодом 1
    
    # Проверяем, есть ли файл с расширением .doc или .docx
    for ext in [".doc", ".docx"]:
        template_path = template_folder / f"{template_name}{ext}"
        if template_path.exists():
            output_path = Path(card_data["2"]) / f"{card_data['НОМ']}-25{ext}"
            shutil.copy(template_path, output_path)
            logger.info(f"Шаблон скопирован в: {output_path}")
            return output_path
    
    # Если файл не найден
    logger.error(f"Файл шаблона '{template_name}' не найден в папке {template_folder}.")
    raise FileNotFoundError(f"Файл шаблона '{template_name}' не найден в папке {template_folder}.")

# Секция для вставки фрагментов из paragraphs.docx
def read_paragraphs_from_word(template_folder):
    """
    Читает абзацы из Word-файла, находящегося в папке с шаблонами.
    Сохраняет все форматирование исходного документа.
    :param template_folder: Путь к папке с шаблонами (из строки Карты с кодом 3)
    :return: Словарь, где ключ — код (ОБ, ОБ2, ОБ3), а значение — список элементов документа
             с сохраненным форматированием
    """
    paragraphs = {}
    doc_path = Path(template_folder) / "paragraphs.docx"  # Путь к paragraphs.docx
    
    if not doc_path.exists():
        raise FileNotFoundError(f"Файл paragraphs.docx не найден в папке {template_folder}.")
    
    try:
        doc = Document(doc_path)
        
        current_key = None
        current_elements = []
        
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            
            # Проверяем, является ли строка ключом (например, "ОБ:")
            if text.endswith(':') and len(text.split()) == 1:
                # Сохраняем предыдущий ключ, если он был
                if current_key is not None and current_elements:
                    paragraphs[current_key] = current_elements.copy()
                    current_elements.clear()
                
                current_key = text[:-1]  # Убираем двоеточие
            else:
                if current_key is not None:
                    current_elements.append(('paragraph', paragraph))
        
        # Обрабатываем таблицы в документе
        for table in doc.tables:
            if current_key is not None:
                current_elements.append(('table', table))
        
        # Добавляем последний ключ
        if current_key is not None and current_elements:
            paragraphs[current_key] = current_elements.copy()
    
    except Exception as e:
        raise ValueError(f"Ошибка при чтении Word-файла: {e}")
    
    return paragraphs

def insert_word_elements(doc, elements, placeholder_paragraph):
    """
    Вставляет элементы Word (с сохранением форматирования) в указанный параграф-заполнитель.
    :param doc: Документ, в который вставляем
    :param elements: Список элементов для вставки
    :param placeholder_paragraph: Параграф, содержащий метку для замены
    """
    # Находим родительский элемент и индекс заполнителя
    parent = placeholder_paragraph._p.getparent()
    index = parent.index(placeholder_paragraph._p)
    
    # Удаляем заполнитель
    parent.remove(placeholder_paragraph._p)
    
    # Вставляем новые элементы
    for elem_type, elem_content in elements:
        if elem_type == 'paragraph':
            # Создаем новый параграф и копируем содержимое
            new_paragraph = doc.add_paragraph()
            new_paragraph._p = deepcopy(elem_content._p)
            parent.insert(index, new_paragraph._p)
            index += 1
        
        elif elem_type == 'table':
            # Создаем новую таблицу и копируем содержимое
            new_table = doc.add_table(rows=0, cols=0)
            new_table._tbl = deepcopy(elem_content._tbl)
            parent.insert(index, new_table._tbl)
            index += 1

def preprocess_paragraphs(doc_path, paragraphs, card_data):
    """
    Полностью удаляет абзацы с метками INSERT_XXX, если значение пустое
    """
    doc = Document(doc_path)
    
    # Сначала собираем все абзацы для удаления
    paragraphs_to_remove = []
    
    # Обрабатываем все параграфы в документе
    for paragraph in doc.paragraphs:
        for key in card_data:
            placeholder = f"{{INSERT_{key}}}"
            if placeholder in paragraph.text:
                if card_data[key]:  # Если значение не пустое
                    if key in paragraphs:
                        # Вставляем элементы с сохранением форматирования
                        insert_word_elements(doc, paragraphs[key], paragraph)
                    else:
                        # Просто удаляем метку, если нет соответствующего фрагмента
                        paragraph.text = paragraph.text.replace(placeholder, "")
                else:  # Если значение пустое - помечаем абзац для удаления
                    paragraphs_to_remove.append(paragraph)
    
    # Удаляем помеченные абзацы (в обратном порядке)
    for paragraph in reversed(paragraphs_to_remove):
        p = paragraph._element
        p.getparent().remove(p)
    
    # Обрабатываем таблицы
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                cell_paragraphs_to_remove = []
                for paragraph in cell.paragraphs:
                    for key in card_data:
                        placeholder = f"{{INSERT_{key}}}"
                        if placeholder in paragraph.text:
                            if card_data[key]:  # Если значение не пустое
                                if key in paragraphs:
                                    # Очищаем ячейку перед вставкой
                                    cell.text = ""
                                    insert_word_elements(doc, paragraphs[key], paragraph)
                                else:
                                    paragraph.text = paragraph.text.replace(placeholder, "")
                            else:  # Если значение пустое
                                cell_paragraphs_to_remove.append(paragraph)
                
                # Удаляем пустые абзацы в ячейках
                for paragraph in reversed(cell_paragraphs_to_remove):
                    p = paragraph._element
                    p.getparent().remove(p)
    
    doc.save(doc_path)

# Заполнение таблицы RT
def process_rt_table(doc_path, card_data):
    """
    Полная обработка таблицы RT с:
    - сохранением форматирования
    - проверкой наличия маркера
    - безопасной обработкой отсутствующих таблиц
    """
    logger.info("Начало обработки таблицы RT")
    doc = Document(doc_path)
    tables_with_marker = 0  # Счётчик найденных таблиц
    
    # Подготовка данных
    # Проверяем наличие объектов в карте
    ob_exists = bool(card_data.get("ОБ"))
    ob2_exists = bool(card_data.get("ОБ2"))
    ob3_exists = bool(card_data.get("ОБ3"))
    
    # Формируем список объектов, которые нужно оставить
    objects_to_keep = ["ОБ"]
    if ob2_exists:
        objects_to_keep.append("ОБ2")
    if ob3_exists:
        objects_to_keep.append("ОБ3")
    
    # Формируем фразу для KV
    kv_parts = ["KV"]
    if ob2_exists:
        kv_parts.append("KV-2")
    if ob3_exists:
        kv_parts.append("KV-3")
    kv_text = ", ".join(kv_parts) + " (отрицательный контроль при выделении ДНК)"
    
    # Поиск таблиц
    for table in doc.tables:
        marker_found = False
        header_cell = None
        
        # Поиск маркера в таблице
        for row in table.rows:
            for cell in row.cells:
                if "[INSERT_RT_DATA]" in cell.text:
                    header_cell = cell
                    marker_found = True
                    break
            if marker_found:
                break
        
        if not marker_found:
            continue  # Пропускаем таблицы без маркера
        
        tables_with_marker += 1
        logger.info(f"Обработка таблицы #{tables_with_marker} с маркером [INSERT_RT_DATA]")
        
        # Обработка таблицы
        # Удаляем метку из заголовка
        if header_cell.paragraphs:
            for paragraph in header_cell.paragraphs:
                if paragraph.runs:
                    for run in paragraph.runs:
                        run.text = run.text.replace("[INSERT_RT_DATA]", "")
        
        # Поиск ячейки с {KV}
        kv_cell = None
        for row in table.rows:
            for cell in row.cells:
                if "{KV}" in cell.text:
                    kv_cell = cell
                    break
            if kv_cell:
                break
        
        # Сбор информации о строках объектов
        rows_info = []
        for row in table.rows:
            row_text = " ".join(cell.text.strip() for cell in row.cells)
            
            if "Объект {ОБ}" in row_text:
                rows_info.append(("ОБ", row))
            elif "Объект {ОБ2}" in row_text:
                rows_info.append(("ОБ2", row))
            elif "Объект {ОБ3}" in row_text:
                rows_info.append(("ОБ3", row))
        
        # Удаление лишних строк
        for obj_code, row in reversed(rows_info):
            if obj_code not in objects_to_keep:
                try:
                    table._tbl.remove(row._tr)
                    logger.info(f"Удалена строка для {obj_code}")
                except ValueError:
                    logger.warning(f"Строка для {obj_code} уже удалена")
        
        # Заполнение данных
        for obj_code, row in rows_info:
            if obj_code in objects_to_keep:
                obj_num = "" if obj_code == "ОБ" else obj_code[-1]
                
                for cell in row.cells:
                    if cell.paragraphs:
                        for paragraph in cell.paragraphs:
                            if paragraph.runs:
                                for run in paragraph.runs:
                                    original_text = run.text
                                    new_text = original_text.replace(
                                        f"{{{obj_code}}}", 
                                        str(card_data.get(obj_code, "")))
                                    
                                    rt_codes = ["ДМ", "КМ", "YХ", "ИД"]
                                    for rt_code in rt_codes:
                                        full_code = f"{rt_code}{obj_num}"
                                        new_text = new_text.replace(
                                            f"{{{full_code}}}", 
                                            str(card_data.get(full_code, "")))
                                    
                                    run.text = new_text
        
        # Обработка KV
        if kv_cell:
            if kv_cell.paragraphs and kv_cell.paragraphs[0].runs:
                first_run = kv_cell.paragraphs[0].runs[0]
                font = first_run.font
                
                # Сохраняем форматирование
                formatting = {
                    'bold': font.bold,
                    'italic': font.italic,
                    'underline': font.underline,
                    'size': font.size,
                    'name': font.name,
                    'color': font.color.rgb if font.color else None
                }
                
                # Заменяем текст
                first_run.text = kv_text
                
                # Восстанавливаем форматирование
                for attr, value in formatting.items():
                    if value is not None:
                        setattr(font, attr, value)
                
                # Удаляем лишние runs
                for run in kv_cell.paragraphs[0].runs[1:]:
                    kv_cell.paragraphs[0]._p.remove(run._r)
    
    # Финализация
    if tables_with_marker == 0:
        logger.warning("В документе нет таблиц с маркером [INSERT_RT_DATA]")
    else:
        logger.info(f"Успешно обработано таблиц: {tables_with_marker}")
    
    doc.save(doc_path)
    logger.info("Обработка завершена")

# Замена меток по словарю
def process_switch_tags(doc_path, card_data, switch_dict):
    """
    Обрабатывает все метки вида {PREFIX_SWITCH:word} в документе Word.
    Сохраняет форматирование и регистр, добавляет логирование.
    """
    logger.info("Начало обработки switch-меток в документе")
    
    try:
        doc = Document(doc_path)
        total_replacements = 0
        
        def process_text(text):
            nonlocal total_replacements
            new_text = text
            
            # Ищем все метки вида {PREFIX_SWITCH:word}
            pattern = r'\{(\w+)_SWITCH:([^}]+)\}'
            matches = re.findall(pattern, text)
            
            for switch_type, original_word in matches:
                try:
                    # Логируем найденную метку
                    logger.debug(f"Найдена метка: {switch_type}_SWITCH:{original_word}")
                    
                    # Определяем, нужно ли множественное число
                    use_plural = False
                    if switch_type == "OB2":
                        ob2_value = card_data.get('ОБ2')
                        use_plural = bool(ob2_value) if ob2_value is not None else False
                        logger.debug(f"Значение ОБ2: '{ob2_value}', use_plural: {use_plural}")
                    
                    # Получаем варианты замены из словаря
                    replacement_dict = switch_dict.get(f"{switch_type}_SWITCH", {})
                    replacements = replacement_dict.get(original_word.lower())
                    
                    if not replacements or len(replacements) < 2:
                        logger.warning(f"Не найдена замена для слова '{original_word}' в словаре {switch_type}_SWITCH")
                        continue
                    
                    # Выбираем нужную форму слова
                    replacement = replacements[1] if use_plural else replacements[0]
                    
                    # Сохраняем регистр исходного слова
                    if original_word.istitle():
                        replacement = replacement.title()
                    elif original_word.isupper():
                        replacement = replacement.upper()
                    
                    # Заменяем метку в тексте
                    old_text = new_text
                    new_text = new_text.replace(
                        f"{{{switch_type}_SWITCH:{original_word}}}", 
                        replacement
                    )
                    
                    if old_text != new_text:
                        total_replacements += 1
                        logger.debug(f"Замена: '{original_word}' -> '{replacement}'")
                
                except Exception as e:
                    logger.error(f"Ошибка при обработке метки {switch_type}_SWITCH:{original_word}: {str(e)}")
                    continue
            
            return new_text
        
        # Обрабатываем все параграфы
        for paragraph in doc.paragraphs:
            original_text = paragraph.text
            new_text = process_text(original_text)
            if new_text != original_text:
                paragraph.text = new_text
        
        # Обрабатываем таблицы
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        original_text = paragraph.text
                        new_text = process_text(original_text)
                        if new_text != original_text:
                            paragraph.text = new_text
        
        # Обрабатываем колонтитулы
        for section in doc.sections:
            for paragraph in section.header.paragraphs:
                original_text = paragraph.text
                new_text = process_text(original_text)
                if new_text != original_text:
                    paragraph.text = new_text
            
            for paragraph in section.footer.paragraphs:
                original_text = paragraph.text
                new_text = process_text(original_text)
                if new_text != original_text:
                    paragraph.text = new_text
        
        logger.info(f"Обработка switch-меток завершена. Всего замен: {total_replacements}")
        doc.save(doc_path)
    
    except Exception as e:
        logger.error(f"Критическая ошибка при обработке switch-меток: {str(e)}")
        raise

# Функция для замены меток в Word
def replace_in_doc(doc_path, replacements):
    """ Заменяет метки в шаблоне Word """
    doc = Document(doc_path)
    found_keys = set()
    
    def replace_text(element):
        nonlocal found_keys
        if not hasattr(element, "runs"):
            return

        # Собираем полный текст параграфа
        full_text = "".join(run.text for run in element.runs)
        if not full_text:  # Если параграф пустой, пропускаем
            return

        # Создаем карту форматирования для каждого символа
        char_formatting = []
        current_pos = 0
        for run in element.runs:
            for _ in run.text:
                char_formatting.append({
                    'bold': run.bold,
                    'italic': run.italic,
                    'underline': run.underline,
                    'font': run.font.name,
                    'size': run.font.size
                })
            current_pos += len(run.text)

        # Создаем новый текст с заменами
        new_text = full_text
        positions = []  # Список кортежей (позиция, длина_старого_текста, новый_текст)

        # Сначала обрабатываем TIME
        if "{TIME}" in new_text:
            current_time = datetime.datetime.now().strftime("%H.%M")
            start_pos = new_text.find("{TIME}")
            if start_pos != -1:
                positions.append((start_pos, len("{TIME}"), current_time))
                found_keys.add("TIME")

        # Затем обрабатываем остальные метки
        for key, value in replacements.items():
            patterns = [f"{{{key}}}", f"{{ {key} }}", f"{{{key} }}", f"{{ {key}}}", f"{{  {key}}}"]
            for pattern in patterns:
                start_pos = 0
                while True:
                    pos = new_text.find(pattern, start_pos)
                    if pos == -1:
                        break
                    positions.append((pos, len(pattern), str(value)))
                    found_keys.add(key)
                    start_pos = pos + len(pattern)

            # Обработка новых меток {KEY:?, {}}
            pattern = f"{{{key}:?, {{}}}}"
            start_pos = 0
            while True:
                pos = new_text.find(pattern, start_pos)
                if pos == -1:
                    break
                # Если значение есть → подставляем ", значение", иначе пустую строку
                replacement = f", {value}" if value else ""
                positions.append((pos, len(pattern), replacement))
                found_keys.add(key)
                start_pos = pos + len(pattern)

        # Сортируем позиции в обратном порядке, чтобы замены не влияли на последующие позиции
        positions.sort(reverse=True)

        # Выполняем замены и обновляем карту форматирования
        for pos, old_len, new_text_piece in positions:
            new_text = new_text[:pos] + new_text_piece + new_text[pos + old_len:]
            # Обновляем карту форматирования
            format_at_pos = char_formatting[pos]
            char_formatting[pos:pos + old_len] = [format_at_pos] * len(new_text_piece)

        new_text = re.sub(r'\{[^\}]*\}', '', new_text)

        # Очищаем существующие runs
        for run in element.runs:
            run._element.getparent().remove(run._element)

        # Создаем новые runs с правильным форматированием
        current_format = None
        current_text = ""
        
        for i, char in enumerate(new_text):
            char_format = char_formatting[i] if i < len(char_formatting) else char_formatting[-1]
            
            if current_format != char_format:
                if current_text:
                    run = element.add_run(current_text)
                    if current_format:
                        run.bold = current_format['bold']
                        run.italic = current_format['italic']
                        run.underline = current_format['underline']
                        if current_format['font']:
                            run.font.name = current_format['font']
                        if current_format['size']:
                            run.font.size = current_format['size']
                current_text = char
                current_format = char_format
            else:
                current_text += char

        # Добавляем последний run
        if current_text:
            run = element.add_run(current_text)
            if current_format:
                run.bold = current_format['bold']
                run.italic = current_format['italic']
                run.underline = current_format['underline']
                if current_format['font']:
                    run.font.name = current_format['font']
                if current_format['size']:
                    run.font.size = current_format['size']

    # Обрабатываем все элементы документа
    for paragraph in doc.paragraphs:
        replace_text(paragraph)
        
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replace_text(paragraph)
                    
    for section in doc.sections:
        for paragraph in section.header.paragraphs:
            replace_text(paragraph)
        for paragraph in section.footer.paragraphs:
            replace_text(paragraph)

        # Обработка номеров страниц
        footer = section.footer
        for paragraph in footer.paragraphs:
            if "стр. [СТ] из [КОЛ]" in paragraph.text:
                paragraph.text = paragraph.text.replace("стр. [СТ] из [КОЛ]", "")
                run = paragraph.add_run("стр. ")
                
                fldChar1 = OxmlElement('w:fldChar')
                fldChar1.set(qn('w:fldCharType'), 'begin')
                instrText = OxmlElement('w:instrText')
                instrText.set(qn('xml:space'), 'preserve')
                instrText.text = 'PAGE \* Arabic'
                fldChar2 = OxmlElement('w:fldChar')
                fldChar2.set(qn('w:fldCharType'), 'end')
                
                run._r.append(fldChar1)
                run._r.append(instrText)
                run._r.append(fldChar2)
                
                run = paragraph.add_run(" из ")
                
                fldChar3 = OxmlElement('w:fldChar')
                fldChar3.set(qn('w:fldCharType'), 'begin')
                instrText2 = OxmlElement('w:instrText')
                instrText2.set(qn('xml:space'), 'preserve')
                instrText2.text = 'NUMPAGES \* Arabic'
                fldChar4 = OxmlElement('w:fldChar')
                fldChar4.set(qn('w:fldCharType'), 'end')
                
                run._r.append(fldChar3)
                run._r.append(instrText2)
                run._r.append(fldChar4)

    # Проверяем, все ли метки были заменены
    missing_keys = set(replacements.keys()) - found_keys
    if missing_keys:
        logging.warning(f"Предупреждение: следующие метки не найдены в документе: {missing_keys}")

    doc.save(doc_path)

# Функция для вставки изображений
def insert_images(doc_path, images_folder, card_data, brightness_factor=1.5, contrast_factor=1.5):
    """
    Вставляет изображения в документ Word с настройкой яркости и контрастности.
    Яркость и контрастность регулируются только для обычных изображений.
    Обычные изображения получают черную рамку толщиной 0.5 пт.
    """
    doc = Document(doc_path)
    images_folder = Path(images_folder)
    
    # Проверяем наличие обычных изображений
    ordinary_images = [img for img in images_folder.glob("*.*") 
                     if img.stem not in {"1", "1_2", "1_3"} 
                     and img.suffix.lower() in {".jpg", ".jpeg", ".png"}]
    if not ordinary_images:
        raise FileNotFoundError(f"Не найдено ни одного обычного изображения в папке {images_folder}")
    
    # Проверяем наличие специальных изображений, если это требуется
    if card_data.get("1") != "Кость СВО нет пригодной ДНК":
        special_images_found = any(img_path.exists() 
                                for img_name in ["1", "1_2", "1_3"] 
                                for ext in [".jpg", ".jpeg", ".png"] 
                                for img_path in [images_folder / f"{img_name}{ext}"])
        if not special_images_found:
            raise FileNotFoundError(f"Не найдено ни одного специального изображения (1, 1_2, 1_3) в папке {images_folder}")
    
    # Определяем высоту изображения в зависимости от значения в поле "1"
    card_value = card_data.get("1", "").strip()
    if card_value == "Образец СВО Молов фонд":
        image_height = Inches(1.8)  # Высота для "Образец СВО Молов фонд"
    elif card_value == "Образец СВО Молов прямая идентификация":
        image_height = Inches(2.0)  # Высота для "Образец СВО Молов прямая идентификация"
    elif card_value == "Кость СВО":
        image_height = Inches(2.1)  # Высота для "Кость СВО"
    else:
        image_height = Inches(1.8)  # Значение по умолчанию
    
    # Вставляем обычные изображения (не 1, 1_2, 1_3)
    for image_path in images_folder.glob("*.*"):
        if image_path.stem not in {"1", "1_2", "1_3"} and image_path.suffix.lower() in {".jpg", ".jpeg", ".png"}:
            print(f"Обработка обычного изображения: {image_path}")
            try:
                # Применяем автоматическую обрезку
                cropped_image = auto_crop_image(str(image_path))
                if cropped_image is None:
                    raise ValueError("Автоматическая обрезка не удалась.")
                print("Обрезка успешно выполнена.")
                
                # Показываем результат автоматической обрезки
                cv2.imshow("Результат автоматической обрезки", cropped_image)
                cv2.waitKey(0)  # Ждём нажатия клавиши
                cv2.destroyAllWindows()
                
                # Предлагаем пользователю выбрать вариант обрезки
                user_choice = input("Необходимо:\n1. Дополнительно обрезать имеющееся изображение;\n2. Самостоятельно обрезать оригинальное изображение.\nЛюбая клавиша для продолжения.\n")
                
                if user_choice == '1':
                    print("Запуск интерактивного режима...")
                    interactive_cropped_image = interactive_crop(cropped_image)
                    if interactive_cropped_image is not None:
                        cropped_image = interactive_cropped_image
                    print("Интерактивная обрезка завершена.")
                elif user_choice == '2':
                    # Загружаем оригинальное изображение с помощью Pillow
                    print("Загрузка оригинального изображения...")
                    original_image_np = load_image_with_pillow(str(image_path))
                    original_image = cv2.cvtColor(original_image_np, cv2.COLOR_RGB2BGR)  # Pillow -> OpenCV (RGB -> BGR)
                    
                    # Применяем обработку (контрастность, яркость и т.д.)
                    print("Применение обработки к оригинальному изображению...")
                    original_image = increase_contrast(original_image)
                    original_image = resize_image(original_image)
                    
                    # Предоставляем пользователю возможность обрезать изображение
                    print("Запуск интерактивного режима для оригинального изображения...")
                    interactive_cropped_image = interactive_crop(original_image)
                    if interactive_cropped_image is not None:
                        cropped_image = interactive_cropped_image
                    else:
                        cropped_image = original_image  # Если пользователь не выбрал область, используем оригинал
                    print("Интерактивная обрезка оригинального изображения завершена.")
                else:
                    print("Продолжение без изменений.")
                
                # Конвертируем обрезанное изображение в формат Pillow для сохранения
                print("Конвертация изображения в формат Pillow...")
                cropped_image_pil = Image.fromarray(cv2.cvtColor(cropped_image, cv2.COLOR_BGR2RGB))
                
                # Сохраняем временное изображение после обрезки через Pillow
                temp_cropped_path = images_folder / "temp_cropped_image.jpg"
                cropped_image_pil.save(temp_cropped_path, format="JPEG", quality=100)
                if not temp_cropped_path.exists():
                    raise FileNotFoundError(f"Не удалось сохранить временный файл: {temp_cropped_path}")
                print(f"Временное изображение после обрезки сохранено: {temp_cropped_path}")
                
                # Применяем настройку яркости и контрастности
                print("Применение настроек яркости и контрастности...")
                adjusted_image = auto_adjust_brightness_contrast(temp_cropped_path, brightness_factor, contrast_factor)
                
                # Сохраняем временное изображение в формате JPEG
                temp_image_path = images_folder / "temp_adjusted_image.jpg"
                adjusted_image.save(temp_image_path, format="JPEG", quality=100)
                if not temp_image_path.exists():
                    raise FileNotFoundError(f"Не удалось сохранить временный файл: {temp_image_path}")
                print(f"Временное изображение после настройки яркости и контрастности сохранено: {temp_image_path}")
                
                # Вставка в документ
                print("Вставка изображения в документ...")
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for paragraph in cell.paragraphs:
                                if "[IMAGE_OTHER]" in paragraph.text:
                                    paragraph.text = ""
                                    run = paragraph.add_run()
                                    
                                    # Добавляем изображение с высотой, зависящей от значения в поле "1"
                                    picture = run.add_picture(str(temp_image_path), height=image_height)
                                    
                                    # Получаем ширину изображения в точках
                                    image_width = picture.width.pt
                                    
                                    # Устанавливаем ширину ячейки равной ширине изображения плюс отступ
                                    padding = 10  # Отступ в точках
                                    cell.width = Pt(image_width + padding)
                                    
                                    # Добавляем черную границу
                                    props = picture._inline.graphic.graphicData.pic.spPr
                                    ln = OxmlElement('a:ln')
                                    ln.set('w', str(int(0.5 * 12700)))  # толщина 0.5 пт
                                    solidFill = OxmlElement('a:solidFill')
                                    srgbClr = OxmlElement('a:srgbClr')
                                    srgbClr.set('val', '000000')
                                    solidFill.append(srgbClr)
                                    ln.append(solidFill)
                                    props.append(ln)
                                    
                                    # Добавляем границы для всех сторон
                                    for border_name in ['top', 'left', 'bottom', 'right']:
                                        border = OxmlElement(f'a:{border_name}')
                                        border.set(qn('w:val'), 'single')
                                        border.set(qn('w:sz'), '4')
                                        border.set(qn('w:space'), '0')
                                        border.set(qn('w:color'), '000000')
                                        ln.append(border)
                                    
                                    # Устанавливаем отступы внутри ячейки
                                    tcPr = cell._tc.get_or_add_tcPr()
                                    tcMar = OxmlElement('w:tcMar')
                                    for margin_name in ['top', 'left', 'bottom', 'right']:
                                        margin = OxmlElement(f'w:{margin_name}')
                                        margin.set(qn('w:w'), '20')  # Отступ в 18 twips (1 twip = 1/20 pt) (расширение границ таблицы)
                                        margin.set(qn('w:type'), 'dxa')
                                        tcMar.append(margin)
                                    tcPr.append(tcMar)
                                    
                                    print(f"Изображение вставлено в ячейку с меткой [IMAGE_OTHER]")
                                    break
                
                # Удаляем временные изображения после вставки
                print("Удаление временных файлов...")
                if temp_image_path.exists():
                    temp_image_path.unlink()
                if temp_cropped_path.exists():
                    temp_cropped_path.unlink()
                print("Временные файлы удалены.")
            
            except Exception as e:
                print(f"Ошибка при обработке изображения {image_path}: {e}")
    
    # Вставляем специальные изображения (1, 1_2, 1_3) без регулировки яркости и контрастности
    special_images = {
        "1": "[IMAGE_1]",
        "1_2": "[IMAGE_1_2]",
        "1_3": "[IMAGE_1_3]"
    }
    for image_name, placeholder in special_images.items():
        image_path = images_folder / f"{image_name}.jpg"
        if not image_path.exists():  # Если файл не найден, пробуем другие расширения
            for ext in [".jpg", ".jpeg", ".png"]:
                image_path = images_folder / f"{image_name}{ext}"
                if image_path.exists():
                    break
        if image_path.exists():
            print(f"Обработка специального изображения: {image_path}")
            try:
                # Загружаем изображение без регулировки яркости и контрастности
                image = Image.open(image_path)
                
                # Сохраняем временное изображение в формате JPEG
                temp_image_path = images_folder / "temp_special_image.jpg"
                image.save(temp_image_path, format="JPEG", quality=100)
                if not temp_image_path.exists():
                    raise FileNotFoundError(f"Не удалось сохранить временный файл: {temp_image_path}")
                print(f"Временное изображение сохранено: {temp_image_path}")
                
                # Вставка в документ
                for p in doc.paragraphs:
                    if placeholder in p.text:
                        p.text = p.text.replace(placeholder, "")
                        run = p.add_run()
                        run.add_picture(str(temp_image_path), width=Inches(6.4))
                        print(f"Изображение вставлено в метку {placeholder}")
                        break
                
                # Удаляем временное изображение после вставки
                if temp_image_path.exists():
                    temp_image_path.unlink()
            except Exception as e:
                print(f"Ошибка при обработке изображения {image_path}: {e}")
    
    doc.save(doc_path)
    print("Все изображения обработаны и вставлены.")

# Функция для поиска файла .txt в целевой папке
def find_txt_file(folder_path):
    """
    Ищет файл(ы) .txt в указанной папке, которые заканчиваются на "Genotypes Table".
    Если найдено несколько файлов, соответствующих этому правилу, позволяет пользователю выбрать нужный файл.
    Если найден только один файл, соответствующий правилу, он автоматически выбирается.
    Если файлов, соответствующих правилу, нет, возвращает None.
    """
    # Ищем все файлы с расширением .txt в целевой папке
    txt_files = list(Path(folder_path).glob("*.txt"))
    logger.info(f"Найдены файлы .txt: {[f.name for f in txt_files]}")

    # Фильтруем файлы, оставляя только те, которые заканчиваются на "Genotypes Table"
    filtered_files = [file for file in txt_files if file.name.endswith("Genotypes Table.txt")]
    logger.info(f"Отфильтрованные файлы: {[f.name for f in filtered_files]}")

    # Если файлов, соответствующих правилу, нет
    if not filtered_files:
        if txt_files:
            logger.warning(f"Файлы .txt найдены, но ни один не соответствует правилу (не заканчивается на 'Genotypes Table').")
        else:
            logger.warning("Файлы .txt не найдены.")
        return None  # Возвращаем None, если файл не найден или не соответствует правилу
    
    # Если найден только один файл, соответствующий правилу
    if len(filtered_files) == 1:
        logger.info(f"Найден файл, соответствующий правилу: {filtered_files[0].name}")
        return filtered_files[0]  # Возвращаем единственный найденный файл
    
    # Если найдено несколько файлов .txt, соответствующих правилу, спрашиваем пользователя
    logger.info(f"Найдено несколько файлов .txt, заканчивающихся на 'Genotypes Table': {[f.name for f in filtered_files]}")
    print(f"В папке {folder_path} найдено несколько файлов .txt, заканчивающихся на 'Genotypes Table':")
    for i, txt_file in enumerate(filtered_files, start=1):
        print(f"{i}. {txt_file.name}")
    
    while True:
        try:
            choice = int(input("Выберите номер файла для загрузки: "))
            if 1 <= choice <= len(filtered_files):
                selected_file = filtered_files[choice - 1]
                logger.info(f"Выбран файл: {selected_file.name}")
                return selected_file
            else:
                print("Некорректный выбор. Пожалуйста, введите номер из списка.")
        except ValueError:
            print("Пожалуйста, введите числовое значение.")

# Функция для чтения и фильтрации данных из .txt файла
def read_and_filter_txt_data(file_path):
    """ Читает данные из txt файла и исключает ненужные значения """
    # Значения, которые нужно исключить (в разных вариантах регистра)
    exclude_values = {"AL", "Al", "K+", "К+", "K-", "К-"}  # Добавляем все возможные варианты
    with open(file_path, 'r', encoding='utf-8') as file:
        lines = file.readlines()

    # Убираем заголовок и разбиваем строки на столбцы
    data = [line.strip().split('\t') for line in lines[1:]]

    # Получаем все уникальные идентификаторы объектов (первая колонка)
    # исключая значения из exclude_values
    object_ids = set()
    for line in data:
        if line[0] and line[0].strip() not in exclude_values:  # Добавляем strip() для удаления лишних пробелов
            object_ids.add(line[0].strip())

    # Выводим найденные идентификаторы для отладки
    print(f"Найденные идентификаторы объектов: {object_ids}")

    # Проверяем, что есть хотя бы один уникальный идентификатор
    if len(object_ids) == 0:
        raise ValueError("В файле не найдено ни одного допустимого объекта.")

    # Если найдено несколько уникальных идентификаторов, спрашиваем пользователя, какой объект загрузить
    if len(object_ids) > 1:
        print("В файле найдено несколько различных объектов:")
        sorted_object_ids = sorted(object_ids, key=lambda x: (x.isdigit(), x))
        print(f"Отсортированные идентификаторы объектов: {sorted_object_ids}")
        print(f"Типы данных в отсортированных идентификаторах: {[type(obj_id) for obj_id in sorted_object_ids]}")
        for i, obj_id in enumerate(sorted_object_ids, start=1):
            print(f"{i}. {obj_id}")
        while True:
            try:
                choice = int(input("Выберите номер объекта для загрузки: ")) - 1
                print(f"Выбран номер: {choice} (тип: {type(choice)})")
                if 0 <= choice < len(sorted_object_ids):
                    selected_object_id = sorted_object_ids[choice]
                    print(f"Выбранный идентификатор объекта: {selected_object_id} (тип: {type(selected_object_id)})")
                    break
                else:
                    print("Неверный выбор. Пожалуйста, введите номер от 1 до", len(sorted_object_ids))
            except ValueError:
                print("Неверный ввод. Пожалуйста, введите число.")
    else:
        selected_object_id = list(object_ids)[0]
    # Фильтруем данные, оставляя только строки с выбранным идентификатором объекта
    filtered_data = [line for line in data if line[0].strip() == selected_object_id]
    print(f"Отфильтрованные данные: {filtered_data}")
    print(f"Типы данных в отфильтрованных данных: {[type(line[0].strip()) for line in filtered_data]}")

    # Проверка на дублирующиеся локусы
    sample_loci = {}  # Словарь для хранения уникальных локусов для каждого Sample Name
    for line in filtered_data:
        sample_name = line[0].strip()  # Sample Name
        marker = line[3].strip()      # Marker

        # Если Sample Name уже встречался, проверяем локусы
        if sample_name in sample_loci:
            if marker in sample_loci[sample_name]:
                # Если локус уже встречался для этого Sample Name, вызываем ошибку
                raise ValueError(f"Для одного объекта найдены дублирующиеся локусы. Проверьте данные. Sample Name: {sample_name}, Marker: {marker}")
            else:
                # Добавляем локус в список для данного Sample Name
                sample_loci[sample_name].add(marker)
        else:
            # Если Sample Name встречается впервые, создаем новый набор локусов
            sample_loci[sample_name] = {marker}

    # Выводим количество отфильтрованных строк данных для отладки
    print(f"Количество отфильтрованных строк данных: {len(filtered_data)}")
    return filtered_data

def check_y_rs_loci(txt_data):
    """
    Проверяет, пусты ли значения аллелей всех локусов, которые в названии содержат "Y", "y", "RS" или "rs".
    Если хотя бы один локус содержит значение, отличное от пробела или нуля, возвращает False.
    """
    # Локусы, которые нужно проверить
    loci_to_check = ["Y", "y", "RS", "rs"]
    
    for line in txt_data:
        marker = line[3].strip().upper()  # Название локуса
        # Проверяем, содержит ли название локуса "Y", "y", "RS" или "rs"
        if any(substring in marker for substring in loci_to_check):
            # Проверяем значения аллелей (столбцы 5 и 6)
            allele1 = line[5].strip() if line[5] else ""
            allele2 = line[6].strip() if line[6] else ""
            # Если хотя бы одно значение не пустое, возвращаем False
            if allele1 or allele2:
                return False
    return True

# Функция для вставки данных в таблицу Word
def insert_table_data(doc_path, txt_data, card_data):
    """
    Вставляет данные из .txt файла в таблицу Word.
    Добавляет комментарий после таблицы в зависимости от условий.
    """
    logger.info("Начало вставки данных в таблицу")
    doc = Document(doc_path)
    
    # Флаг для проверки наличия буквы Y в Amel/AMEL
    has_y_in_amel = False
    
    # Проверяем, есть ли "Y" в локусе Amel
    for line in txt_data:
        if line[3].strip().upper() == "AMEL":
            alleles = [line[5].strip().upper(), line[6].strip().upper()]
            if "Y" in alleles:
                has_y_in_amel = True
                break
    
    # Если "Y" не обнаружен в Amel, проверяем локусы с "Y", "y", "RS" или "rs"
    if not has_y_in_amel:
        if not check_y_rs_loci(txt_data):
            show_error_message("Обнаружены значения в локусах, содержащих 'Y', 'y', 'RS' или 'rs'. Выполнение остановлено.")
            sys.exit(1)
    
    # Сортируем данные по локусам
    sorted_data = sort_data_by_locus(txt_data)
    logger.info(f"Данные отсортированы по локусам: {sorted_data}")

    # Ищем таблицу с меткой [TABLE_DATA]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if "[TABLE_DATA]" in cell.text:
                    logger.info("Найдена таблица с меткой [TABLE_DATA]")
                    # Очищаем ячейку и вставляем данные
                    cell.text = ""
                    # Вставляем данные из sorted_data (отсортированные данные)
                    for line in sorted_data:
                        marker = line[3].strip()  # Сохраняем оригинальный регистр
                        allele1 = line[5] if line[5] else ""
                        allele2 = line[6] if line[6] else ""
                        alleles = ",".join(filter(None, [allele1, allele2])) if allele1 or allele2 else "-"
                        # Добавляем новую строку в таблицу
                        new_row = table.add_row()
                        new_row.cells[0].text = marker
                        new_row.cells[1].text = alleles
                        logger.debug(f"Добавлен маркер: {marker}, аллели: {alleles}")
                    
                    # Удаляем все пустые строки в таблице
                    rows_to_remove = []
                    for row in table.rows:
                        if all(cell.text.strip() == "" for cell in row.cells):
                            rows_to_remove.append(row)
                    for row in rows_to_remove:
                        table._tbl.remove(row._tr)
                    logger.info("Пустые строки удалены из таблицы")
                    
                    # Убедимся, что границы таблицы видны
                    for row in table.rows:
                        for cell in row.cells:
                            tc = cell._tc
                            tcPr = tc.get_or_add_tcPr()
                            tcBorders = OxmlElement('w:tcBorders')
                            for border_name in ['top', 'left', 'bottom', 'right']:
                                border = OxmlElement(f'w:{border_name}')
                                border.set(qn('w:val'), 'single')
                                border.set(qn('w:sz'), '4')
                                border.set(qn('w:space'), '0')
                                border.set(qn('w:color'), '000000')
                                tcBorders.append(border)
                            tcPr.append(tcBorders)
                    logger.info("Границы таблицы настроены")
                    
                    # Выравниваем текст во втором столбце по центру
                    for row in table.rows:
                        if len(row.cells) > 1:
                            cell = row.cells[1]
                            cell.paragraphs[0].alignment = 1  # Центрирование
                    logger.info("Текст во втором столбце выровнен по центру")
                    
                    # Добавляем примечание, если выполнены условия
                    if not has_y_in_amel:
                        comment = None
                        if card_data.get("1") in ["Образец СВО Молов фонд", "Образец СВО Молов прямая идентификация", "Образец СВО Ростов"]:
                            comment = " (-) - прочерк означает отсутствие продуктов амплификации."
                        elif card_data.get("1") == "Кость СВО":
                            comment = " (-) - отсутствие продуктов амплификации или удовлетворительных результатов."
                        
                        if comment:
                            p = doc.add_paragraph()
                            run = p.add_run("Примечание:")
                            run.underline = True
                            run.font.size = Pt(10)
                            run.font.name = 'Times New Roman'
                            p.add_run(f" {comment}").font.size = Pt(10)
                            # Перемещаем примечание сразу после таблицы
                            table_element = table._tbl
                            parent_element = table_element.getparent()
                            parent_element.insert(parent_element.index(table_element) + 1, p._element)
                            logger.info("Комментарий добавлен после таблицы")
                    break
    
    # Сохраняем документ
    doc.save(doc_path)
    logger.info("Данные успешно вставлены в таблицу и документ сохранён")

# Блок проверки консистентности информации в заключении
# Сравнивает номер из Карты, номер файла txt и номер заключения из фореграммы
# Сравнивает данные по родству из Карты и наличие Y
def preprocess_image_for_ocr(image_path):
    """
    Предварительная обработка изображения для улучшения распознавания OCR.
    """
    # Загружаем изображение с помощью Pillow
    image_pil = Image.open(image_path)
    # Увеличиваем контрастность и яркость
    enhancer = ImageEnhance.Contrast(image_pil)
    image_pil = enhancer.enhance(2.0)
    enhancer = ImageEnhance.Brightness(image_pil)
    image_pil = enhancer.enhance(1.5)
    # Преобразуем в numpy array для работы с OpenCV
    image = np.array(image_pil)
    # Преобразуем из RGB (Pillow) в BGR (OpenCV)
    image = cv2.cvtColor(image, cv2.COLOR_RGB2BGR)
    # Преобразование в оттенки серого
    gray = cv2.cvtColor(image, cv2.COLOR_BGR2GRAY)
    # Удаление шумов
    gray = cv2.GaussianBlur(gray, (5, 5), 0)
    return gray

def determine_gender_from_parents(card_data):
    """
    Определяет пол подэкспертного на основе данных о родителях.
    """
    parent_info = card_data.get("РОД")
    if parent_info in ["матери", "сестры", "дочери"]:
        return "женский"
    elif parent_info in ["отца", "брата", "сына"]:
        return "мужской"
    return None

def determine_gender_from_amel(txt_data):
    """
    Определяет пол подэкспертного на основе данных локуса AMEL.
    """
    for line in txt_data:
        if line[3].strip().upper() == "AMEL":
            alleles = [line[5], line[6]]
            if "Y" in alleles:
                return "мужской"
            elif "X" in alleles and "Y" not in alleles:
                return "женский"
    return None

def check_gender(card_data, txt_data):
    """
    Проверяет пол подэкспертного на основе данных о родителях и локуса AMEL.
    """
    
    if txt_data is None:  # Защита от None
        return
    
    gender_from_parents = determine_gender_from_parents(card_data)
    gender_from_amel = determine_gender_from_amel(txt_data)

    if gender_from_parents and gender_from_amel and gender_from_parents != gender_from_amel:
        raise ValueError(f"Несоответствие данных о поле подэкспертного: родители указывают на {gender_from_parents}, локус AMEL указывает на {gender_from_amel}.")

def extract_case_number_from_txt_filename(txt_path):
    """
    Извлекает номер заключения из имени файла .txt.
    """
    import re
    filename = txt_path.name
    match = re.search(r'^(\d+)', filename)
    if match:
        return match.group(1).strip()  # Добавляем strip() для удаления лишних пробелов
    return None

def extract_case_number_from_foregram(foregram_path):
    """
    Извлекает номер заключения из фореграммы с использованием OCR.
    """
    gray = preprocess_image_for_ocr(foregram_path)
    # Применяем OCR для извлечения текста
    text = pytesseract.image_to_string(gray, config='--psm 6 --oem 3 -c tessedit_char_whitelist=0123456789-')
    print(f"Распознанный текст из фореграммы: {text}")  # Отладочный вывод
    
    # Разделяем текст на строки
    lines = text.split('\n')
    
    # Ищем номер заключения в первой строке
    for line in lines:
        line = line.strip()
        if line:
            # Ищем номер заключения (например, число в начале строки)
            import re
            match = re.search(r'^(\d+-?\d*)', line)
            if match:
                case_number_with_year = match.group(1).strip()  # Добавляем strip() для удаления лишних пробелов
                # Убираем последние две цифры года или все символы после тире
                if '-' in case_number_with_year:
                    case_number = case_number_with_year.split('-')[0]
                else:
                    case_number = case_number_with_year[:-2]
                return case_number
            break  # Прерываем цикл после первой строки
    
    return None

def check_case_number(card_data, txt_path, foregram_path):
    """
    Проверяет консистентность номера заключения в Карте, фотографии, фореграмме и файле .txt.
    """
    case_number_from_card = str(card_data.get("НОМ")).strip()  # Преобразуем к строке и убираем лишние пробелы
    case_number_from_foregram = extract_case_number_from_foregram(foregram_path)
    case_number_from_txt = extract_case_number_from_txt_filename(txt_path)

    if (case_number_from_card != case_number_from_foregram or
        case_number_from_card != case_number_from_txt):
        raise ValueError(f"Несоответствие номеров заключений: "
                         f"в Карте - {case_number_from_card}, "
                         f"в фореграмме - {case_number_from_foregram}, "
                         f"в файле .txt - {case_number_from_txt}.")

def check_consistency(card_data, txt_data, foregram_path, txt_path):
    """
    Проверяет консистентность данных в карте, данных локусов и изображениях
    с обработкой отсутствующей фореграммы.
    """
    
    if txt_data is None:
        return
        
    # Проверка фореграммы
    if foregram_path is None or not foregram_path.exists():
        if card_data.get("1") != "Кость СВО нет пригодной ДНК":
            raise FileNotFoundError("Фореграмма отсутствует, но требуется по условиям")
        return  # Пропускаем проверку
    
    # Проверка пола подэкспертного
    check_gender(card_data, txt_data)

    # Проверка номера заключения
    check_case_number(card_data, txt_path, foregram_path)

# Основные функции
def find_excel_file():
    """
    Ищет единственный Excel-файл в текущей директории.
    Возвращает путь к файлу или вызывает исключение, если файл не найден или их несколько.
    """
    # Получаем список всех Excel-файлов в текущей директории
    excel_files = list(Path('.').glob('*.xlsx')) + list(Path('.').glob('*.xls'))
    
    # Проверяем количество найденных файлов
    if not excel_files:
        raise FileNotFoundError("В текущей директории не найдено Excel-файлов (.xlsx или .xls)")
    if len(excel_files) > 1:
        raise ValueError(f"В директории найдено несколько Excel-файлов: {[f.name for f in excel_files]}. "
                        "Должен быть только один файл.")
    
    return excel_files[0]

# Основная функция
def main():
    try:
        logger.info("Запуск программы")

        # Находим Excel-файл
        card_path = find_excel_file()
        logger.info(f"Найден файл: {card_path}")

        # Чтение данных из найденного файла
        card_data = read_card(card_path)
        logger.info("Данные из файла: %s", card_data)

        # Проверяем, что все обязательные поля заполнены
        check_card_data(card_data)

        # Загрузка словаря замен
        templates_folder = Path(card_data["3"])
        switch_dict = load_switch_dictionary(templates_folder)
        if not switch_dict:
            logger.warning("Словарь замен не загружен или пуст")
        else:
            logger.info("Словарь замен успешно загружен")

        # Копирование шаблона Word
        output_path = copy_template(card_data)
        logger.info(f"Шаблон скопирован в: {output_path}")

        # Чтение абзацев из txt-файла
        template_folder = card_data["3"]  # Путь к папке с шаблонами (строка Карты с кодом 3)
        paragraphs = read_paragraphs_from_word(template_folder)
        logger.info(f"Загружены фрагменты для ключей: {list(paragraphs.keys())}")

        # Предварительная обработка абзацев
        preprocess_paragraphs(output_path, paragraphs, card_data)
        logger.info("Абзацы предварительно обработаны.")

        # Обработка таблицы RT (новая строка)
        process_rt_table(output_path, card_data)
        logger.info("Таблица RT обработана.")

        # Создаем словарь со всеми значениями из Карты
        replacements = {
            code: str(value) if value is not None else ''
            for code, value in card_data.items()
        }
        logger.info("Метки для замены: %s", replacements)

        # Обрабатываем switch-метки
        if switch_dict:
            process_switch_tags(output_path, card_data, switch_dict)
        else:
            logger.warning("Пропуск обработки switch-меток: словарь замен не загружен")

        # Заменяем метки в документе
        replace_in_doc(output_path, replacements)
        logger.info("Метки заменены.")

        # Получаем путь к папке с рабочими материалами из Карты (поле с кодом "4")
        base_working_materials_folder = Path(card_data["4"])
        nom_folder_name = f"{card_data['НОМ']}-25"
        working_materials_folder = base_working_materials_folder / nom_folder_name

        # Проверяем, существует ли папка
        if not working_materials_folder.exists():
            raise FileNotFoundError(f"Папка '{working_materials_folder}' не существует.")
        
        # Вставка изображений
        insert_images(output_path, working_materials_folder, card_data)
        logger.info("Изображения вставлены.")
            
        # Находим ВСЕ обязательные изображения (1, 1_2, 1_3)
        if card_data.get("1") != "Кость СВО нет пригодной ДНК":
            required_images = ["1", "1_2", "1_3"]
            missing_images = []
            found_images = {}

            for img_name in required_images:
                img_path = None
                for ext in [".jpg", ".jpeg", ".png"]:
                    temp_path = working_materials_folder / f"{img_name}{ext}"
                    if temp_path.exists():
                        img_path = temp_path
                        break
                
                if img_path:
                    found_images[img_name] = img_path
                else:
                    missing_images.append(img_name)

            if missing_images:
                error_msg = f"Не найдены обязательные изображения: {', '.join(missing_images)}"
                logger.error(error_msg)
                show_error_message(error_msg)
                sys.exit(1)
            else:
                logger.info(f"Все специальные изображения найдены: {list(found_images.keys())}")
                foregram_path = found_images["1"]

        # Находим файл .txt в целевой папке
        txt_path = find_txt_file(working_materials_folder)

        # Проверяем, нужно ли искать и обрабатывать файл .txt
        if card_data.get("1") in [
            "Кость СВО нет пригодной ДНК",
            "СВО_Молов_образец_родственники_нет результата_форез",
            "СВО_Молов_образец_родственники_нет результата_RT"
        ]:
            logger.info("Значение 'Кость СВО нет пригодной ДНК' в строке Карты с кодом 1 - пропускаем обработку файла .txt")
            txt_data = None
        else:
            if txt_path is None:
                raise FileNotFoundError("Файл .txt с генотипами не найден, и это не соответствует условиям для пропуска.")
            logger.info(f"Найден файл .txt: {txt_path}")
            # Чтение и фильтрация данных из .txt файла
            txt_data = read_and_filter_txt_data(txt_path)
            logger.info("Отфильтрованные данные из .txt файла: %s", txt_data)
            # Вставка данных в таблицу
            insert_table_data(output_path, txt_data, card_data)  # Передаём card_data для проверки условий
            logger.info("Данные вставлены в таблицу.")        

        # Проверка консистентности данных
        if card_data.get("1") in [
            "Кость СВО нет пригодной ДНК",
            "СВО_Молов_образец_родственники_нет результата_форез",
            "СВО_Молов_образец_родственники_нет результата_RT"
        ]:
            logger.info("Пропуск проверки консистентности для особых случаев")
        else:
            try:
                if txt_path is None:
                    raise FileNotFoundError("Файл .txt с генотипами не найден!")
                if foregram_path is None:
                    raise FileNotFoundError("Фореграмма (изображение 1) не найдена!")
                
                check_consistency(card_data, txt_data, foregram_path, txt_path)
                logger.info("Проверка консистентности прошла успешно.")
            except Exception as e:
                logger.error(f"ОШИБКА ПРОВЕРКИ: {e}")
                show_error_message(f"ОШИБКА ПРОВЕРКИ: {e}")
                sys.exit(1)

        logger.info("Скрипт успешно завершён!")
        
    except FileNotFoundError as e:
        logger.error(f"Ошибка: {e}")
        show_error_message(f"Ошибка: {e}")
    except Exception as e:
        logger.error(f"Произошла ошибка: {e}")
        show_error_message(f"Произошла ошибка: {e}")
        
# Запуск скрипта
if __name__ == "__main__":
    main()